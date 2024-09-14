# imports
import pandas as pd
import os
import datetime
import re
from openpyxl import Workbook
from openpyxl import load_workbook
import tkinter as tk
from tkinter import filedialog
from docx import Document

# Ignore Warnings
import warnings
warnings.simplefilter(action='ignore', category=UserWarning)

# author: Javier Andrews
# @jaedango

# -------------- Variables globales -------------- #
# Orden de las paginas
pag_sum = 0
pag_res = 1
pag_ucs = 2
p_ucs_m = 3
pag__hl = 4
pag__ch = 5
pag__pf = 6
p_plt_d = 7
p_plt_a = 8
p_plt_c = 9
pag__ti = 10
pag__tx = 11
pag_sla = 12
pag_adm = 13
pag__eg = 14
pag_drx = 15
pag_mic = 16
pag_tdr = 17

# Nombre de las paginas
name_ucs = "UCS"
name_plt_d = "PLT Diametral"
name_plt_a = "PLT Axial"
name_plt_c = "PLT Colpa"
name_hl = "Hinchamiento Libre"
name_ch = "Corte Hoek"
name_pf = "Propiedades Físicas"
name_ti = "Tracción Indirecta"
name_tx = "TX-M"
name_slake = "Slake"
name_adm = "Azul de Metileno"
name_eg = "Etilenglicol"
name_drx = "DRX"
name_res = "Resumen"
name_ucs_m = "UCS-M"
name_mic = "Microscopía"
name_sum = "Resumen de Cantidades"
name_tdr = "Tipo de Roca"

# Nombre de pagina de errores
error_name = "Informe_errores"

# Diccionario de cantidad de datos por informe
dict_data = {"data_inicial":7}
dict_data.update({"mic":22})
dict_data.update({"pf":4})
dict_data.update({"am":2})
dict_data.update({"hl":1})
dict_data.update({"lim_a":3})
dict_data.update({"plt":3})
dict_data.update({"ucs":2})
dict_data.update({"ucs-m":4})
# dict_data.update({"ucs-s":2})
dict_data.update({"trx":15})
dict_data.update({"ti":1})
dict_data.update({"ch":8})
dict_data.update({"sla":2})
dict_data.update({"eg":12})
dict_data.update({"ph":2})
dict_data.update({"drx":41})

# total de datos
suma_total_data = 0
for key in dict_data:
    suma_total_data += dict_data[key]

# -------------- Extraer info -------------- #
def extraer_info_UCS(df):
    datos = [""] * 21

    try:
        datos[0] = df.iloc[0, 9]    # j1 titulo_informe
        datos[1] = df.iloc[2, 11]   # l3 proyecto
        datos[2] = df.iloc[3, 11]   # l4 n_informe
        # datos[3] = df.iloc[4, 11]   # l5 orden_trabajo
        datos[4] = df.iloc[5, 11]   # l6 fecha_inicio
        datos[5] = df.iloc[6, 11]   # l7 fecha_termino
        datos[3] = df.iloc[9, 11]   # l10 muestra
        datos[6] = df.iloc[10, 11]  # l11 tipo_roca
        datos[7] = df.iloc[11, 11]  # l12 fracturas
        datos[8] = df.iloc[12, 11]  # l13 alteraciones
        datos[9] = df.iloc[13, 11]  # l14 observaciones
        datos[10] = df.iloc[16, 11] # l17 diametro
        datos[11] = df.iloc[17, 11] # l18 altura
        datos[12] = df.iloc[18, 11] # l19 densidad_h
        datos[13] = df.iloc[19, 11] # l20 densidad_s
        datos[14] = df.iloc[20, 11] # l21 contenido_h
        datos[15] = df.iloc[16, 17] # r17 resist_d50
        datos[16] = df.iloc[17, 17] # r18 resist_d64

    except:
        return [1]

    # Revisar si hay algun dato incorrecto
    for dato in datos:
        if pd.isnull(dato):
            return [0]
        
    # formatear datos
    try:
        datos[0] = datos[0].lstrip().rstrip()
        datos[1] = datos[1].lstrip().rstrip()
        datos[2] = datos[2].lstrip().rstrip()
        datos[3] = datos[3].lstrip().rstrip()
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y") 
        datos[6] = datos[6].lstrip().rstrip()
        datos[7] = datos[7].lstrip().rstrip()
        datos[8] = datos[8].lstrip().rstrip()
        datos[9] = datos[9].lstrip().rstrip()
        datos[10] = round(float(datos[10]), 2) if isinstance(datos[10], float) else datos[10]
        datos[11] = round(float(datos[11]), 2) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 2) if isinstance(datos[12], float) else datos[12]
        datos[13] = round(float(datos[13]), 2) if isinstance(datos[13], float) else datos[13]
        datos[14] = round(float(datos[14]), 2) if isinstance(datos[14], float) else datos[14]
        datos[15] = round(float(datos[15]), 1) if isinstance(datos[15], float) else datos[15]
        datos[16] = round(float(datos[16]), 1) if isinstance(datos[16], float) else datos[16]
    except:
        return[1]
    
    # observaciones
    obs = []
    obs_rows = df.iloc[22:, 9]

    # buscar donde se enceuntra la palabra Observaciones
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break
    index += 1  # aumentar index para no seguir directamente desde 'Observaciones'
    val = 17    # valor para seguir en el arreglo de datos

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 9]):
            break
        ob1 = df.iloc[index, 9]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 10]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 10]
            datos[val] = ob1 + ' ' + str(ob2)

        # Aumentar contadores
        index += 1
        val += 1

    for idx, dato in enumerate(datos):
        if dato == "":
            datos[idx] = "-"

    return datos

def extraer_info_UCS_M(df):
    datos = [""] * 24

    try:
        datos[0] = df.iloc[0, 11]	# l1 titulo informe
        datos[1] = df.iloc[2, 12]	# m3 proyecto
        datos[2] = df.iloc[3, 12]	# m4 num informe
        # datos[3] = df.iloc[4, 11]	# m5 orden trabajo
        datos[4] = df.iloc[5, 12]	# m6 fecha inicio
        datos[5] = df.iloc[6, 12]	# m7 fecha termino
        datos[3] = df.iloc[8, 12]	# m9 muestra
        datos[6] = df.iloc[9, 12]	# m10 tipo de roca
        datos[7] = df.iloc[10, 12]	# m11 fracturas
        datos[8] = df.iloc[11, 12]	# m12 alteraciones
        datos[9] = df.iloc[12, 12]	# m13 observaciones
        datos[10] = df.iloc[15, 12]	# m16 diametro
        datos[11] = df.iloc[16, 12]	# m17 altura
        datos[12] = df.iloc[17, 12]	# m18 peso
        datos[13] = df.iloc[18, 12]	# m19 densidad h
        datos[14] = df.iloc[19, 12]	# m20 densidad s
        datos[15] = df.iloc[20, 12]	# m21 humedad
        datos[16] = df.iloc[15, 15]	# p16 resist d50
        datos[17] = df.iloc[16, 15]	# p17 mod def e
        datos[18] = df.iloc[17, 15]	# p18 poisson
        datos[19] = df.iloc[18, 15]	# p19 muestra falla por
    except:
        return [1]

    # Revisar si hay algun dato incorrecto
    for dato in datos:
        if pd.isnull(dato):
            return [0]

    # formatear datos
    try:
        datos[0] = datos[0].lstrip().rstrip()
        datos[1] = datos[1].lstrip().rstrip()
        datos[2] = datos[2].lstrip().rstrip()
        datos[3] = datos[3].lstrip().rstrip()
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y") 
        datos[6] = datos[6].lstrip().rstrip()
        datos[7] = datos[7].lstrip().rstrip()
        datos[8] = datos[8].lstrip().rstrip()
        datos[9] = datos[9].lstrip().rstrip()
        datos[10] = round(float(datos[10]), 2) if isinstance(datos[10], float) else datos[10]
        datos[11] = round(float(datos[11]), 2) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 2) if isinstance(datos[12], float) else datos[12]
        datos[13] = round(float(datos[13]), 2) if isinstance(datos[13], float) else datos[13]
        datos[14] = round(float(datos[14]), 2) if isinstance(datos[14], float) else datos[14]
        datos[15] = round(float(datos[15]), 1) if isinstance(datos[15], float) else datos[15]
        datos[16] = round(float(datos[16]), 1) if isinstance(datos[16], float) else datos[16]
        datos[17] = round(float(datos[17]), 2) if isinstance(datos[17], float) else datos[17]
        datos[18] = round(float(datos[18]), 2) if isinstance(datos[18], float) else datos[18]
        datos[19] = datos[19].lstrip().rstrip()
    except:
        return[1]
    
    # observaciones
    obs = []
    obs_rows = df.iloc[:, 11]

    # buscar donde se enceuntra la palabra Observaciones
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones:':
            break
    index += 1  # aumentar index para no seguir directamente desde 'Observaciones'
    val = 20    # valor para seguir en el arreglo de datos

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 11]):
            break
        ob1 = df.iloc[index, 11]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 12]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 12]
            datos[val] = ob1 + ' ' + str(ob2)

        # Aumentar contadores
        index += 1
        val += 1

    for idx, dato in enumerate(datos):
        if dato == "":
            datos[idx] = "-"

    return datos

# Extrae toda la informacion necesaria para la hoja de hinchamiento libre
def extraer_info_HL(df):
    datos = [''] * 13
    try:
        datos[0] = df.iloc[0, 5] # Titulo informe
        datos[1] = df.iloc[2, 6] # Proyecto
        datos[2] = df.iloc[3, 6] # N° Informe
        # datos[3] = df.iloc[4, 6] # Orden de trabajo
        datos[4] = df.iloc[5, 6] # Fecha Inicio
        datos[5] = df.iloc[6, 6] # Fecha Termino
        datos[3] = df.iloc[9, 6] # Muestra
        datos[6] = df.iloc[12,7] # Vol incial
        datos[7] = df.iloc[13,7] # Vol final
        datos[8] = df.iloc[16,7] # Indice H
    except:
        return[1]

    # Revisar si hay datos incorrectos
    for dato in datos:
        if pd.isnull(dato):
            return [0]
        
    # Formatear datos
    datos[0] = datos[0].lstrip().rstrip()
    datos[1] = datos[1].lstrip().rstrip()
    datos[2] = datos[2].lstrip().rstrip()
    datos[3] = datos[3].lstrip().rstrip()
    datos[4] = datos[4].strftime("%d-%m-%Y") 
    datos[5] = datos[5].strftime("%d-%m-%Y") 
    datos[6] = round(float(datos[6]), 1) if isinstance(datos[6], float) else datos[6]
    datos[7] = round(float(datos[7]), 1) if isinstance(datos[7], float) else datos[7]
    datos[8] = round(float(datos[8]), 1) if isinstance(datos[8], float) else datos[8]

    # observaciones
    obs = []
    obs_rows = df.iloc[:, 5]

    # buscar donde se encuentra la palabra 'Observaciones'
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break
    index += 1  # aumentar index para no seguir directamente desde obs
    val = 9    # valor para seguir el arreglo de datos
    
    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 5]):
            break
        obs1 = df.iloc[index, 5]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 6]):
            datos[val] = obs1
        else:
            obs2 = df.loc[index, 6]
            datos[val] = obs[1] + ' ' + str(obs2)
        # aumentar contadores
        index += 1
        val += 1

    for idx, dato in enumerate(datos):
        if dato == "":
            datos[idx] = "-"

    return datos

# Extrae toda la informacion necesaria para la hoja de Corte Hoek
def extraer_info_CH(df):
    datos =[''] * 21
    try:
        datos[0] = df.iloc[0, 4]    # E1 Titulo Informe
        datos[1] = df.iloc[2, 5]    # F3 Proyecto
        datos[2] = df.iloc[3, 5]    # F4 N Informe
        # datos[3] = df.iloc[4, 5]    # F5 Orden de trabajo
        datos[4] = df.iloc[5, 5]    # f6 Fecha inicio
        datos[5] = df.iloc[6, 5]    # f7 fecha termino
        datos[3] = df.iloc[8, 5]    # f9 muestra
        datos[6] = df.iloc[9, 5]    # f10 tipo roca
        datos[7] = df.iloc[10, 5]   # f11 fracturas
        datos[8] = df.iloc[11, 5]   # f12 alteraciones
        datos[9] = df.iloc[59, 16]	# q60 t normal 1
        datos[10] = df.iloc[59, 17]	# r60 t max 1
        datos[11] = df.iloc[60, 16]	# q61 t normal 2
        datos[12] = df.iloc[60, 17]	# r61 t max 2
        datos[13] = df.iloc[61, 16]	# q62 t normal 3
        datos[14] = df.iloc[61, 17]	# r62 t max 3
        datos[15] = df.iloc[63, 17]	# r64 pendiente
        datos[16] = df.iloc[64, 17]	# r65 fi
    except:
        return 1

    # revisar si hay algun dato incorrecto
    for dato in datos[:-10]:
        if pd.isnull(dato):
            return [0]
        
    # Formatear datos
    try:
        datos[0] = datos[0].lstrip().rstrip()
        datos[1] = datos[1].lstrip().rstrip()
        datos[2] = datos[2].lstrip().rstrip()
        datos[3] = datos[3].lstrip().rstrip()
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos[6] = datos[6].lstrip().rstrip()
        datos[7] = datos[7].lstrip().rstrip()
        datos[8] = datos[8].lstrip().rstrip()
        datos[9] = round(float(datos[9]), 1) if isinstance(datos[9], float) else datos[9]
        datos[10] = round(float(datos[10]), 1) if isinstance(datos[10], float) else datos[10]
        datos[11] = round(float(datos[11]), 1) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 1) if isinstance(datos[12], float) else datos[12]
        datos[13] = round(float(datos[13]), 1) if isinstance(datos[13], float) else datos[13]
        datos[14] = round(float(datos[14]), 1) if isinstance(datos[14], float) else datos[14]
        datos[15] = round(float(datos[15]), 2) if isinstance(datos[15], float) else datos[15]
        datos[16] = round(float(datos[16]), 1) if isinstance(datos[16], float) else datos[16]
    except:
        return[1]
    
    # observaciones
    obs = []
    obs_rows = df.iloc[:, 4]

    # buscar donde se encuentra la palabra 'Observaciones'
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break
    index += 1  # aumentar index para no seguir directamente desde obs
    val = 17    # valor para seguir el arreglo de datos

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 4]):
            break
        obs1 = df.iloc[index, 4]
        # revisa si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 5]):
            datos[val] = obs1
        else:
            obs2 = df.loc[index, 5]
            datos[val] = obs1 + ' ' + obs2
        index += 1
        val += 1

    for idx, dato in enumerate(datos):
        if dato == "":
            datos[idx] = "-"

    return datos

# Extrae toda la informacion necesaria para la hoja de Propiedades Fisicas
# revisar
def extraer_info_PF(df):
    datos_final = []
    datos = [''] * 6
    try:
        datos[0] = df.iloc[10, 10]  # k11 num muestra
        datos[1] = df.iloc[0, 11]   # L1 Titulo Informe
        datos[2] = df.iloc[2, 12]   # M3 Proyecto
        datos[3] = df.iloc[3, 12]   # M4 N Informe
        # datos[4] = df.iloc[4, 12]   # M5 Orden de trabajo
        datos[4] = df.iloc[5, 12]   # M6 Fecha inicio
        datos[5] = df.iloc[6, 12]   # M7 Fecha termino
    except:
        return[1]
    
    datos2 = [''] * 6
    try:
        datos2[0] = df.iloc[13, 11] # L14 Muestra
        datos2[1] = df.iloc[13, 12] # M14 Cota
        datos2[2] = df.iloc[13, 13] # N14 Humedad
        datos2[3] = df.iloc[13, 14] # O14 Absorcion
        datos2[4] = df.iloc[13, 15] # P14 Densidad
        datos2[5] = df.iloc[13, 16] # Q14 Porosidad
    except:
        return[1]
    
    # chequear muestra
    muestra = 14

    if pd.isnull(df.iloc[10, 10]):
        datos[0] = df.iloc[14, 10]
        muestra = 18

    # revisar si hay algun dato incorrecto
    for dato in datos:
        if pd.isnull(dato):
            return [0]
        
    # Formatear datos
    try:
        datos[0] = re.search(r"\d+", str(datos[0])).group(0)
        datos[1] = datos[1].lstrip().rstrip()
        datos[2] = datos[2].lstrip().rstrip()
        datos[3] = datos[3].lstrip().rstrip()
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos2[0] = datos2[0].lstrip().rstrip()
        datos2[1] = datos2[1].lstrip().rstrip()
        datos2[2] = round(float(datos2[2]), 2) if isinstance(datos2[2], float) else datos2[2]
        datos2[3] = round(float(datos2[3]), 2) if isinstance(datos2[3], float) else datos2[3]
        datos2[4] = round(float(datos2[4]), 2) if isinstance(datos2[4], float) else datos2[4]
        datos2[5] = round(float(datos2[5]), 2) if isinstance(datos2[5], float) else datos2[5]
    except:
        return[1]
    
    # observaciones
    obs = [''] * 4
    obs_rows = df.iloc[:, 11]
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break
    
    index += 1  # aumentar index para no seguir directamente desde 'Observaciones'
    val = 0

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 11]):
            break
        ob1 = df.iloc[index, 11]
        # revisar si la celda adyactente tiene contenido
        if pd.isnull(df.iloc[index, 12]):
            obs[val] = ob1
        else:
            obs2 = df.iloc[index, 12]
            obs[val] = ob1 + ' ' + str(obs2)
        
        # aumentar contadores
        index += 1
        val += 1
    
    # Revisar siguientes columnas y agregar datos
    dato = []
    dato.extend(datos)
    dato.extend(datos2)
    dato.extend(obs)

    for idx, dat in enumerate(dato):
        if dat == "":
            dato[idx] = "-"

    datos_final.append(dato)

    index = 14

    # Revisar si hay mas columnas
    while not pd.isnull(df.iloc[index, 11]):
        datos[0] = df.iloc[muestra, 10] # extraer muestras

        datos2 = [''] * 6
        try:
            datos2[0] = df.iloc[index, 11]  # L14 Muestra
            datos2[1] = df.iloc[index, 12]  # M14 Cota
            datos2[2] = df.iloc[index, 13]  # N14 Humedad
            datos2[3] = df.iloc[index, 14]  # O14 Absorcion
            datos2[4] = df.iloc[index, 15]  # P14 Densidad
            datos2[5] = df.iloc[index, 16]  # Q14 Porosidad
        except:
            return[1]
        
        # Formatear datos
        try:
            datos[0] = re.search(r"\d+", str(datos[0])).group(0)
            datos2[0] = datos2[0].lstrip().rstrip()
            datos2[1] = datos2[1].lstrip().rstrip()
            datos2[2] = round(float(datos2[2]), 2) if isinstance(datos2[2], float) else datos2[2]
            datos2[3] = round(float(datos2[3]), 2) if isinstance(datos2[3], float) else datos2[3]
            datos2[4] = round(float(datos2[4]), 2) if isinstance(datos2[4], float) else datos2[4]
            datos2[5] = round(float(datos2[5]), 2) if isinstance(datos2[5], float) else datos2[5]
        except:
            return[1]
        
        # agregar a la lista
        dato = []
        dato.extend(datos)
        dato.extend(datos2)
        dato.extend(obs)

        for idx, dat in enumerate(dato):
            if dat == "":
                dato[idx] = "-"
                
        datos_final.append(dato)

        # sumar a index
        index += 1
        muestra += 4
    
    return datos_final

# Extrae toda la informacion necesaria para la hoja de triaxial
def extraer_info_TX(df):
    datos = [''] * 25
    try:
        datos[0] = df.iloc[0, 10]   # K1  Titulo Informe
        datos[1] = df.iloc[2, 11]   # L3  Proyecto
        datos[2] = df.iloc[3, 11]   # L4  N Informe
        # datos[3] = df.iloc[4, 11]   # L5  Orden de Trabajo
        datos[4] = df.iloc[5, 11]   # L6  Fecha Inicio
        datos[5] = df.iloc[6, 11]   # L7  Fecha termino
        datos[3] = df.iloc[8, 11]   # L9  Muestra
        datos[6] = df.iloc[9, 11]   # L10 Tipo de Roca
        datos[7] = df.iloc[10, 11]  # L11 Fracturas
        datos[8] = df.iloc[11, 11]  # L12 Alteraciones
        datos[9] = df.iloc[12, 11] # L13 Observaciones
        datos[10] = df.iloc[15, 11] # L17 Diametro
        datos[11] = df.iloc[16, 11] # L18 Altura
        datos[12] = df.iloc[17, 11] # L19 Peso
        datos[13] = df.iloc[18, 11] # L20 Densidad H
        datos[14] = df.iloc[19, 11] # L21 Densidad s
        datos[15] = df.iloc[20, 11] # L22 Humedad
        datos[16] = df.iloc[15, 14] # O17 Tension
        datos[17] = df.iloc[16, 14] # O18 Resistencia max
        datos[18] = df.iloc[17, 14] # O19 Modulo de def 
        datos[19] = df.iloc[18, 14] # O20 Razon de Poisson
        datos[20] = df.iloc[19, 14] # O21 Muestra falla por
    except:
        return[1]

    # revisar si hay datos incorrectos
    for dato in datos:
        if pd.isnull(dato):
            return [0]
    
    if datos[3] == 0 or datos[3] == "0":
        return [0]
    
    # Formatear datos
    try:
        datos[0] = datos[0].lstrip().rstrip()
        datos[1] = datos[1].lstrip().rstrip()
        datos[2] = datos[2].lstrip().rstrip()
        datos[3] = datos[3].lstrip().rstrip()
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos[6] = datos[6].lstrip().rstrip()
        datos[7] = datos[7].lstrip().rstrip()
        datos[8] = datos[8].lstrip().rstrip()
        datos[9] = datos[9].lstrip().rstrip()
        datos[10] = round(float(datos[10]), 2) if isinstance(datos[10], float) else datos[10]
        datos[11] = round(float(datos[11]), 2) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 2) if isinstance(datos[12], float) else datos[12]
        datos[13] = round(float(datos[13]), 2) if isinstance(datos[13], float) else datos[13]
        datos[14] = round(float(datos[14]), 2) if isinstance(datos[14], float) else datos[14]
        datos[15] = round(float(datos[15]), 2) if isinstance(datos[15], float) else datos[15]
        datos[16] = round(float(datos[16]), 2) if isinstance(datos[16], float) else datos[16]
        datos[17] = round(float(datos[17]), 2) if isinstance(datos[17], float) else datos[17]
        datos[18] = round(float(datos[18]), 2) if isinstance(datos[18], float) else datos[18]
        datos[19] = round(float(datos[19]), 2) if isinstance(datos[19], float) else datos[19]
        datos[20] = datos[20].lstrip().rstrip()
    except:
        return[1]
    
    # observaciones
    obs_rows = df.iloc[25:, 10]
    
    # buscar donde se encuentra la palabra 'Observaciones'
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones:':
            break
    index += 1  # aumentar index para no seguir directamente desde 'Observaciones'
    val = 21    # valor para seguir en el arreglo de datos

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 10]):
            break
        ob1 = df.iloc[index, 10]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 11]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 11]
            datos[val] = ob1 + ' ' + str(ob2)

        # aumentar contadores
        index += 1
        val +=1

    for idx, dato in enumerate(datos):
        if dato == "":
            datos[idx] = "-"

    return datos

# Extrae toda la informacion necesaria para la hoja de Traccion Indirecta
def extraer_info_TI(df):
    datos = [''] * 30
    try:
        datos[0] = df.iloc[0, 14]	# O1 Titulo Informe
        datos[1] = df.iloc[2, 16]	# Q3 Proyecto
        datos[2] = df.iloc[3, 16]	# Q4 N informe
        # datos[3] = df.iloc[4, 16]	# Q5 Orden de trabao
        datos[4] = df.iloc[5, 16]   # Q6 fecha inicio
        datos[5] = df.iloc[6, 16]	# Q7 fecha termino
        datos[3] = df.iloc[8, 16]	# Q9 muestra
        datos[6] = df.iloc[9, 16]	# Q10 tipo de roca
        datos[7] = df.iloc[10, 16]	# Q11 fracturas
        datos[8] = df.iloc[11, 16]	# Q12 Alteraciones
        datos[9] = df.iloc[16, 14]	# O17 Diametro 1
        datos[10] = df.iloc[16, 15]	# P17 Altura 1
        datos[11] = df.iloc[16, 16]	# Q17 fza resist 1
        datos[12] = df.iloc[16, 17]	# R17 resistencia 1
        datos[13] = df.iloc[16, 18]	# S17 tipo de falla 1
        datos[14] = df.iloc[17, 14]	# O18 diametro 2
        datos[15] = df.iloc[17, 15]	# P18 altura 2
        datos[16] = df.iloc[17, 16]	# Q18 fza resist 2
        datos[17] = df.iloc[17, 17]	# R18 resist 2
        datos[18] = df.iloc[17, 18]	# S18 tipo de falla 2
        datos[19] = df.iloc[18, 14]	# O19 diametro 3 
        datos[20] = df.iloc[18, 15]	# P19 altura 3 
        datos[21] = df.iloc[18, 16]	# Q19 fza resist 3
        datos[22] = df.iloc[18, 17]	# R19 resist 3
        datos[23] = df.iloc[18, 18]	# s19 tipo de falla 3
        datos[24] = df.iloc[16, 19]	# T17 promedio resist
        datos[25] = df.iloc[10, 1]  # B11 Humedad
    except:
        return[1]
    
    # revisar si hay datos incorrectos
    for dato in datos:
        if pd.isnull(dato):
            return [0]
    
    # Formatear datos
    try:
        datos[0] = datos[0].lstrip().rstrip()
        datos[1] = datos[1].lstrip().rstrip()
        datos[2] = datos[2].lstrip().rstrip()
        datos[3] = datos[3].lstrip().rstrip()
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos[6] = datos[6].lstrip().rstrip()
        datos[7] = datos[7].lstrip().rstrip()
        datos[8] = datos[8].lstrip().rstrip()
        datos[9] = round(float(datos[9]), 2) if isinstance(datos[9], float) else datos[9]
        datos[10] = round(float(datos[10]), 2) if isinstance(datos[10], float) else datos[10]
        datos[11] = round(float(datos[11]), 2) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 2) if isinstance(datos[12], float) else datos[12]
        datos[13] = datos[13].lstrip().rstrip()
        datos[14] = round(float(datos[14]), 2) if isinstance(datos[14], float) else datos[14]
        datos[15] = round(float(datos[15]), 2) if isinstance(datos[15], float) else datos[15]
        datos[16] = round(float(datos[16]), 2) if isinstance(datos[16], float) else datos[16]
        datos[17] = round(float(datos[17]), 2) if isinstance(datos[17], float) else datos[17]
        datos[18] = datos[18].lstrip().rstrip()
        datos[19] = round(float(datos[19]), 2) if isinstance(datos[19], float) else datos[19]
        datos[20] = round(float(datos[20]), 2) if isinstance(datos[20], float) else datos[20]
        datos[21] = round(float(datos[21]), 2) if isinstance(datos[21], float) else datos[21]
        datos[22] = round(float(datos[22]), 2) if isinstance(datos[22], float) else datos[22]
        datos[23] = datos[23].lstrip().rstrip()
        datos[24] = round(float(datos[24]), 2) if isinstance(datos[24], float) else datos[24]
        datos[25] = round(float(datos[25]), 2) if isinstance(datos[25], float) else datos[25]
    except:
        return[1]
    
    # observaciones
    obs_rows = df.iloc[:, 14]

    # buscar donde se encuentra la palabra 'Observaciones'
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break
    
    index += 1
    val = 26

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 14]):
            break
        ob1 = df.iloc[index, 14]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 15]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 15]
            datos[val] = ob1 + ' ' + str(ob2)
        
        # aumentar contadores
        index += 1
        val += 1
    
    for idx, dato in enumerate(datos):
        if dato == "":
            datos[idx] = "-"

    return datos

# Extrae toda la informacion necesaria para la hoja de PLT Diametral
def extraer_info_PLT_diametral(df):
    datos = [''] * 37
    try:
        datos[0] = df.iloc[0, 11]	# l1 Titulo informe
        datos[1] = df.iloc[3, 13]	# n4 Proyecto
        datos[2] = df.iloc[4, 13]	# n5 Num informe
        # datos[3] = df.iloc[5, 13]	# n6 Orden de trabajo
        datos[4] = df.iloc[6, 13]	# n7 fecha inicio
        datos[5] = df.iloc[7, 13]	# n8 fecha termino
        datos[3] = df.iloc[10, 13]	# n11 muestra
        datos[6] = df.iloc[11, 13]	# n12 tipo de roca
        datos[7] = df.iloc[12, 13]	# n13 fracturas
        datos[8] = df.iloc[13, 13]	# n14 alteraciones
        datos[9] = df.iloc[20, 11]	# l21 diametro 1
        datos[10] = df.iloc[20, 12]	# m21 largo 1
        datos[11] = df.iloc[20, 13]	# n21 Fza ruptura 1
        datos[12] = df.iloc[20, 14]	# o21 ICP ls 1
        datos[13] = df.iloc[20, 15]	# p21 Factor correccion 1
        datos[14] = df.iloc[20, 16]	# q21 ICP ls50 1
        datos[15] = df.iloc[20, 17]	# r21 Muestra falla por 1
        datos[16] = df.iloc[21, 11]	# l22 Diamtro 2 
        datos[17] = df.iloc[21, 12]	# m22 Largo 2
        datos[18] = df.iloc[21, 13]	# n22 fza ruptura 2
        datos[19] = df.iloc[21, 14]	# o22  ICP ls 2
        datos[20] = df.iloc[21, 15]	# p22 factor correccion 2
        datos[21] = df.iloc[21, 16]	# q22 ICP ls50 2
        datos[22] = df.iloc[21, 17]	# r22 muestra falla por 2
        datos[23] = df.iloc[22, 11]	# l23 Diametro 3
        datos[24] = df.iloc[22, 12]	# m23 Largo 3
        datos[25] = df.iloc[22, 13]	# n23 fza ruptura 3
        datos[26] = df.iloc[22, 14]	# o23 ICP ls 3
        datos[27] = df.iloc[22, 15]	# p23 factor correccion 3
        datos[28] = df.iloc[22, 16]	# q23 ICP ls50 3
        datos[29] = df.iloc[22, 17]	# r23 muestra falla por 3
        datos[30] = df.iloc[25, 15]	# p26 contenido humedad
        datos[31] = df.iloc[26, 15]	# p27 Indice resist diam 50mm
        datos[32] = df.iloc[27, 15]	# p28 clasificacion resist
    except:
        return[1]
    
    # revisar si hay datos incorrectos
    for dato in datos:
        if pd.isnull(dato):
            return [0]
    
    # Formatear datos
    try:
        datos[0] = datos[0].lstrip().rstrip()
        datos[1] = datos[1].lstrip().rstrip()
        datos[2] = datos[2].lstrip().rstrip()
        datos[3] = datos[3].lstrip().rstrip()
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos[6] = datos[6].lstrip().rstrip()
        datos[7] = datos[7].lstrip().rstrip()
        datos[8] = datos[8].lstrip().rstrip()
        datos[9] = round(float(datos[9]), 2) if isinstance(datos[9], float) else datos[9]
        datos[19] = round(float(datos[10]), 2) if isinstance(datos[10], float) else datos[10]
        datos[11] = round(float(datos[11]), 2) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 2) if isinstance(datos[12], float) else datos[12]
        datos[13] = round(float(datos[13]), 2) if isinstance(datos[13], float) else datos[13]
        datos[14] = round(float(datos[14]), 2) if isinstance(datos[14], float) else datos[14]
        datos[15] = datos[15].lstrip().rstrip()
        datos[16] = round(float(datos[16]), 2) if isinstance(datos[16], float) else datos[16]
        datos[17] = round(float(datos[17]), 2) if isinstance(datos[17], float) else datos[17]
        datos[18] = round(float(datos[18]), 2) if isinstance(datos[18], float) else datos[18]
        datos[19] = round(float(datos[19]), 2) if isinstance(datos[19], float) else datos[19]
        datos[20] = round(float(datos[20]), 2) if isinstance(datos[20], float) else datos[20]
        datos[21] = round(float(datos[21]), 2) if isinstance(datos[21], float) else datos[21]
        datos[22] = datos[22].lstrip().rstrip()
        datos[23] = round(float(datos[23]), 2) if isinstance(datos[23], float) else datos[23]
        datos[24] = round(float(datos[24]), 2) if isinstance(datos[24], float) else datos[24]
        datos[25] = round(float(datos[25]), 2) if isinstance(datos[25], float) else datos[25]
        datos[26] = round(float(datos[26]), 2) if isinstance(datos[26], float) else datos[26]
        datos[27] = round(float(datos[27]), 2) if isinstance(datos[27], float) else datos[27]
        datos[28] = round(float(datos[28]), 2) if isinstance(datos[28], float) else datos[28]
        datos[29] = datos[29].lstrip().rstrip()
        datos[30] = round(float(datos[30]), 2) if isinstance(datos[30], float) else datos[30]
        datos[31] = round(float(datos[31]), 2) if isinstance(datos[31], float) else datos[31]
        datos[32] = datos[32].lstrip().rstrip()
    except:
        return[1]
    
    # observaciones
    obs_rows = df.iloc[:, 11]

    # buscar donde se encuentra la palabra Observaciones
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break

    index += 1
    val = 33

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 11]):
            break
        ob1 = df.iloc[index, 11]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 12]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 12]
            datos[val] = ob1 + ' ' + str(ob2)

        # aumentar contadores
        index += 1
        val += 1
    
    for idx, dato in enumerate(datos):
        if dato == "":
            datos[idx] = "-"

    return datos

# Extrae toda la informacion necesaria para la hoja de PLT Axial
def extraer_info_PLT_axial(df):
    try:
        datos = [''] * 43
        datos[0] = df.iloc[0, 12]	# m1 Titulo informe
        datos[1] = df.iloc[3, 13]	# n4 proyecto
        datos[2] = df.iloc[4, 13]	# n5 num muestra
        # datos[3] = df.iloc[5, 13]	# n6 orden de trabajo
        datos[4] = df.iloc[6, 13]	# n7 fecha inicio
        datos[5] = df.iloc[7, 13]	# n8 fecha termino
        datos[3] = df.iloc[10, 13]	# n11 muestra
        datos[6] = df.iloc[11, 13]	# n12 tipo de roca
        datos[7] = df.iloc[12, 13]	# n13 fracturas
        datos[8] = df.iloc[13, 13]	# n14 alteraciones
        datos[9] = df.iloc[19, 12]	# m20 dist entre puntas 1
        datos[10] = df.iloc[19, 13]	# n20 ancho 1
        datos[11] = df.iloc[19, 14]	# o20 fuerza de ruptura 1
        datos[12] = df.iloc[19, 15]	# p20 de 1
        datos[13] = df.iloc[19, 16]	# q20 de 1
        datos[14] = df.iloc[19, 17]	# r20 icp ls 1
        datos[15] = df.iloc[19, 18]	# s20 factor de correccion 1
        datos[16] = df.iloc[19, 19]	# t20 icp ls50 1
        datos[17] = df.iloc[19, 20]	# u20 muestra falla por 1
        datos[18] = df.iloc[20, 12]	# m21 dist entre puntas 2
        datos[19] = df.iloc[20, 13]	# n21 ancho 2
        datos[20] = df.iloc[20, 14]	# o21 fuerza de ruptura 2
        datos[21] = df.iloc[20, 15]	# p21 de 2
        datos[22] = df.iloc[20, 16]	# q21 de 2
        datos[23] = df.iloc[20, 17]	# r21 icp ls 2
        datos[24] = df.iloc[20, 18]	# s21 factor de correccion 2
        datos[25] = df.iloc[20, 19]	# t21 icp ls50 2
        datos[26] = df.iloc[20, 20]	# u21 muestra falla por 2
        datos[27] = df.iloc[21, 12]	# m22 dist entre puntas 3
        datos[28] = df.iloc[21, 13]	# b22 ancho 3
        datos[29] = df.iloc[21, 14]	# o22 fuerza de ruptura 3
        datos[30] = df.iloc[21, 15]	# p22 de 3
        datos[31] = df.iloc[21, 16]	# q22 de 3
        datos[32] = df.iloc[21, 17]	# r22 icp ls 3
        datos[33] = df.iloc[21, 18]	# s22 factor de correccion 3
        datos[34] = df.iloc[21, 19]	# t22 icp ls50 3
        datos[35] = df.iloc[21, 20]	# u22 muestra falla por 3
        datos[36] = df.iloc[24, 15]	# p25 contenido de humedad
        datos[37] = df.iloc[25, 15]	# p26 indice de recist diam 500 mm
        datos[38] = df.iloc[26, 15]	# p27 clasificacion de resistencia
    except:
        return[1]
    
    # revisar si hay datos incorrectos
    for dato in datos:
        if pd.isnull(dato):
            return [0]
        
    # formatear datos
    try:
        datos[0] = datos[0].lstrip().rstrip()
        datos[1] = datos[1].lstrip().rstrip()
        datos[2] = datos[2].lstrip().rstrip()
        datos[3] = datos[3].lstrip().rstrip()
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos[6] = datos[6].lstrip().rstrip()
        datos[7] = datos[7].lstrip().rstrip()
        datos[8] = datos[8].lstrip().rstrip()
        datos[9] = round(float(datos[9]), 1) if isinstance(datos[9], float) else datos[9]
        datos[10] = round(float(datos[10]), 2) if isinstance(datos[10], float) else datos[10]
        datos[11] = round(float(datos[11]), 2) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 2) if isinstance(datos[12], float) else datos[12]
        datos[13] = round(float(datos[13]), 2) if isinstance(datos[13], float) else datos[13]
        datos[14] = round(float(datos[14]), 2) if isinstance(datos[14], float) else datos[14]
        datos[15] = round(float(datos[15]), 2) if isinstance(datos[15], float) else datos[15]
        datos[16] = round(float(datos[16]), 2) if isinstance(datos[16], float) else datos[16]
        datos[17] = datos[17].lstrip().rstrip()
        datos[18] = round(float(datos[18]), 1) if isinstance(datos[18], float) else datos[18]
        datos[19] = round(float(datos[19]), 2) if isinstance(datos[19], float) else datos[19]
        datos[20] = round(float(datos[20]), 2) if isinstance(datos[20], float) else datos[20]
        datos[21] = round(float(datos[21]), 2) if isinstance(datos[21], float) else datos[21]
        datos[22] = round(float(datos[22]), 2) if isinstance(datos[22], float) else datos[22]
        datos[23] = round(float(datos[23]), 2) if isinstance(datos[23], float) else datos[23]
        datos[24] = round(float(datos[24]), 2) if isinstance(datos[24], float) else datos[24]
        datos[25] = round(float(datos[25]), 2) if isinstance(datos[25], float) else datos[25]
        datos[26] = datos[26].lstrip().rstrip()
        datos[27] = round(float(datos[27]), 1) if isinstance(datos[27], float) else datos[27]
        datos[28] = round(float(datos[28]), 2) if isinstance(datos[28], float) else datos[28]
        datos[29] = round(float(datos[29]), 2) if isinstance(datos[29], float) else datos[29]
        datos[30] = round(float(datos[30]), 2) if isinstance(datos[30], float) else datos[30]
        datos[31] = round(float(datos[31]), 2) if isinstance(datos[31], float) else datos[31]
        datos[32] = round(float(datos[32]), 2) if isinstance(datos[32], float) else datos[32]
        datos[33] = round(float(datos[33]), 2) if isinstance(datos[33], float) else datos[33]
        datos[34] = round(float(datos[34]), 2) if isinstance(datos[34], float) else datos[34]
        datos[35] = datos[35].lstrip().rstrip()
        datos[36] = round(float(datos[36]), 2) if isinstance(datos[36], float) else datos[36]
        datos[37] = round(float(datos[37]), 2) if isinstance(datos[37], float) else datos[37]
        datos[38] = datos[38].lstrip().rstrip()
    except:
        return[1]
    
    # observaciones
    obs_rows = df.iloc[:, 12]

    # buscar donde se encuentra la palabra Observaciones
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break

    index += 1
    val = 39

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 12]):
            break
        ob1 = df.iloc[index, 12]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 13]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 13]
            datos[val] = ob1 + ' ' + str(ob2)

        # aumentar contadores
        index += 1
        val += 1

    for idx, dato in enumerate(datos):
        if dato == "":
            datos[idx] = "-"

    return datos

# Extrae toda la informacion necesaria para la hoja de PLT Colpa
def extraer_info_PLT_colpa(df):
    datos_final = [''] * 143
    try:
        datos_final[0] = df.iloc[0, 20]	# u1 Titulo informe
        datos_final[1] = df.iloc[3, 21]	# v4 Proyecto
        datos_final[2] = df.iloc[4, 21]	# v5 num informe
        # datos_final[3] = df.iloc[5, 21]	# v6 Orden de trabajo
        datos_final[4] = df.iloc[6, 21]	# v7 fecha inicio
        datos_final[5] = df.iloc[7, 21]	# v8 fecha termino
        datos_final[3] = df.iloc[8, 21]	# v9 muestra
    except:
        return[1]

    # revisar si hay algun dato incorrecto
    for dato in datos_final:
        if pd.isnull(dato):
            return [0]

    try:
        datos_final[0] = datos_final[0].lstrip().rstrip()
        datos_final[1] = datos_final[1].lstrip().rstrip()
        datos_final[2] = datos_final[2].lstrip().rstrip()
        datos_final[3] = datos_final[3].lstrip().rstrip()
        datos_final[4] = datos_final[4].strftime("%d-%m-%Y") 
        datos_final[5] = datos_final[5].strftime("%d-%m-%Y")
    except:
        return [1]
    
    for index in range(10):
        try:
            datos_final[0 + index * 13 + 6] = df.iloc[14 + index, 20]	# u15 identificacion del fragmento a ensayar
            datos_final[1 + index * 13 + 6] = df.iloc[14 + index, 21]	# v15 distancia entre puntas
            datos_final[2 + index * 13 + 6] = df.iloc[14 + index, 22]	# w15 ancho promedio
            datos_final[3 + index * 13 + 6] = df.iloc[14 + index, 23]	# x15 fuerza de ruptura
            datos_final[4 + index * 13 + 6] = df.iloc[14 + index, 24]	# y15 de2
            datos_final[5 + index * 13 + 6] = df.iloc[14 + index, 25]	# z15 de
            datos_final[6 + index * 13 + 6] = df.iloc[14 + index, 26]	# aa15 icp ls
            datos_final[7 + index * 13 + 6] = df.iloc[14 + index, 27]	# ab15 factor de correccion
            datos_final[8 + index * 13 + 6] = df.iloc[14 + index, 28]	# ac15 icp ls50
            datos_final[9 + index * 13 + 6] = df.iloc[14 + index, 29]	# ad15 tipo de roca
            datos_final[10 + index * 13 + 6] = df.iloc[14 + index, 30]	# ae15 fracturas
            datos_final[11 + index * 13 + 6] = df.iloc[14 + index, 31]	# af15 alteracion
            datos_final[12 + index * 13 + 6] = df.iloc[14 + index, 32]	# ag15 muestra falla por
        except:
            return [1]
        
        # Formatear los datos
        try:
            datos_final[0 + index*13 + 6] = datos_final[0 + index*13 + 6].lstrip().rstrip()
            datos_final[1 + index*13 + 6] = round(float(datos_final[1 + index*13 + 6]), 1) if isinstance(datos_final[1 + index*13 + 6], float) else datos_final[1 + index*13 + 6]
            datos_final[2 + index*13 + 6] = round(float(datos_final[2 + index*13 + 6]), 1) if isinstance(datos_final[2 + index*13 + 6], float) else datos_final[2 + index*13 + 6]
            datos_final[3 + index*13 + 6] = round(float(datos_final[3 + index*13 + 6]), 1) if isinstance(datos_final[3 + index*13 + 6], float) else datos_final[3 + index*13 + 6]
            datos_final[4 + index*13 + 6] = round(float(datos_final[4 + index*13 + 6]), 1) if isinstance(datos_final[4 + index*13 + 6], float) else datos_final[4 + index*13 + 6]
            datos_final[5 + index*13 + 6] = round(float(datos_final[5 + index*13 + 6]), 1) if isinstance(datos_final[5 + index*13 + 6], float) else datos_final[5 + index*13 + 6]
            datos_final[6 + index*13 + 6] = round(float(datos_final[6 + index*13 + 6]), 1) if isinstance(datos_final[6 + index*13 + 6], float) else datos_final[6 + index*13 + 6]
            datos_final[7 + index*13 + 6] = round(float(datos_final[7 + index*13 + 6]), 1) if isinstance(datos_final[7 + index*13 + 6], float) else datos_final[7 + index*13 + 6]
            datos_final[8 + index*13 + 6] = round(float(datos_final[8 + index*13 + 6]), 1) if isinstance(datos_final[8 + index*13 + 6], float) else datos_final[8 + index*13 + 6]
            datos_final[9 + index*13 + 6] = datos_final[9 + index*13 + 6].lstrip().rstrip()
            datos_final[10 + index*13 + 6] = datos_final[10 + index*13 + 6].lstrip().rstrip()
            datos_final[11 + index*13 + 6] = datos_final[11 + index*13 + 6].lstrip().rstrip()
            datos_final[12 + index*13 + 6] = datos_final[12 + index*13 + 6].lstrip().rstrip()
        except:
            return [1]
        
    datos_final[136] = df.iloc[49, 23]	# x50 cont humedad
    datos_final[137] = df.iloc[50, 23]	# x51 ls50
    datos_final[138] = df.iloc[51, 23]	# x52 clasificacion resist

    try:
        datos_final[136] = round(float(datos_final[136]), 1) if isinstance(datos_final[136], float) else datos_final[136]
        datos_final[137] = round(float(datos_final[137]), 2) if isinstance(datos_final[137], float) else datos_final[137]
        datos_final[138] = datos_final[136].lstrip().rstrip()
    except:
        return [1]

    # observaciones
    obs_rows = df.iloc[:, 20]
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break
    index += 1
    val = 139

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 20]):
            break
        ob1 = df.iloc[index, 20]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 21]):
            datos_final[val] = ob1
        else:
            obs2 = df.iloc[index, 21]
            datos_final[val] = ob1 + ' ' + str(obs2)

        # aumentar contadores
        index += 1
        val += 1

    for idx, dato in enumerate(datos_final):
        if dato == "":
            datos_final[idx] = "-"

    return datos_final

# Extrae toda la informacion necesaria para la hoja de Slake
def extraer_info_slake(df):
    datos = [''] * 32
    try:
        datos[0] = df.iloc[0, 7]	# h1 Titulo Informe
        datos[1] = df.iloc[2, 8]	# i3 Proyecto
        datos[2] = df.iloc[3, 8]	# i4 num informe
        # datos[3] = df.iloc[4, 8]	# i5 orden de trabajo
        datos[4] = df.iloc[5, 8]	# i6 fecha inicio
        datos[5] = df.iloc[6, 8]	# i7 fecha termino
        datos[3] = df.iloc[8, 8]	# i9 muestra
        datos[6] = df.iloc[9, 8]	# i10 tipo de roca
        datos[7] = df.iloc[10, 8]	# i11 fracturas
        datos[8] = df.iloc[11, 8]	# i12 alteraciones
        datos[9] = df.iloc[15, 10]	# k16 contenedor
        datos[10] = df.iloc[16, 10]	# k17 contenedor + muestra inicial
        datos[11] = df.iloc[17, 10]	# k18 contenedor 1er ciclo
        datos[12] = df.iloc[18, 10]	# k19 contenedor 2do ciclo
        datos[13] = df.iloc[21, 10]	# k22 tipo de agua utilizada
        datos[14] = df.iloc[22, 10]	# k23 temperatura del agua
        datos[15] = df.iloc[23, 10]	# k24 tiempo de cada ciclo
        datos[16] = df.iloc[24, 10]	# k25 velocidad de giro
        datos[17] = df.iloc[25, 10]	# k26 temperatura de secado
        datos[18] = df.iloc[30, 8]	# i31 desgaste ciclo 0
        datos[19] = df.iloc[30, 10]	# k31 indice de durabilidad ciclo 0
        datos[20] = df.iloc[31, 8]	# i32 desgaste ciclo 1
        datos[21] = df.iloc[31, 10]	# k32 indice de durabilidad ciclo 1
        datos[22] = df.iloc[32, 8]	# i33 desgaste ciclo 2
        datos[23] = df.iloc[32, 10]	# k33 indice de durabilidad ciclo 2
        datos[24] = df.iloc[34, 10]	# k35 contenido humedad inicial
        datos[25] = df.iloc[35, 10]	# k36 slake durability index
        datos[26] = df.iloc[36, 10]	# k37 clasificacion
    except:
        return[1]
    
    # revisar si hay datos incorrectos
    for dato in datos:
        if pd.isnull(dato):
            return [0]
        
    # Formatear datos
    try:
        datos[0] = datos[0].lstrip().rstrip()
        datos[1] = datos[1].lstrip().rstrip()
        datos[2] = datos[2].lstrip().rstrip()
        datos[3] = datos[3].lstrip().rstrip()
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos[6] = datos[6].lstrip().rstrip()
        datos[7] = datos[7].lstrip().rstrip()
        datos[8] = datos[8].lstrip().rstrip()
        datos[9] = round(float(datos[9]), 2) if isinstance(datos[9], float) else datos[9]
        datos[10] = round(float(datos[10]), 2) if isinstance(datos[10], float) else datos[10]
        datos[11] = round(float(datos[11]), 2) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 2) if isinstance(datos[12], float) else datos[12]
        datos[13] = datos[13].lstrip().rstrip()
        datos[18] = round(float(datos[18]), 2) if isinstance(datos[18], float) else datos[18]
        datos[19] = round(float(datos[19]), 2) if isinstance(datos[19], float) else datos[19]
        datos[20] = round(float(datos[20]), 2) if isinstance(datos[20], float) else datos[20]
        datos[21] = round(float(datos[21]), 2) if isinstance(datos[21], float) else datos[21]
        datos[22] = round(float(datos[22]), 2) if isinstance(datos[22], float) else datos[22]
        datos[23] = round(float(datos[23]), 2) if isinstance(datos[23], float) else datos[23]
        datos[24] = round(float(datos[24]), 2) if isinstance(datos[24], float) else datos[24]
        datos[25] = round(float(datos[25]), 2) if isinstance(datos[25], float) else datos[25]
        datos[26] = datos[26].lstrip().rstrip()
    except:
        return[1]
    
    # observaciones
    obs_rows = df.iloc[:, 7]

    # buscar donde se encuentra la palabra Observaciones
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break

    index += 1
    val = 27

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 7]):
            break
        ob1 = df.iloc[index, 7]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 8]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 8]
            datos[val] = ob1 + ' ' + ob2
        
        # aumentar contadores
        index += 1
        val += 1
    
    for idx, dato in enumerate(datos):
        if dato == "":
            datos[idx] = "-"

    return datos

# Extrae toda la informacion necesaria para la hoja de Azul de Metileno
def extraer_info_azul_de_metileno(df):
    datos = [''] * 17
    try:
        datos[0] = df.iloc[0, 7]	# h1 Titulo de informe
        datos[1] = df.iloc[2, 8]	# i3 proyecto
        datos[2] = df.iloc[3, 8]	# i4 num informe
        # datos[3] = df.iloc[4, 8]	# i5 orden de trabajo
        datos[4] = df.iloc[5, 8]	# i6 fecha inicial
        datos[5] = df.iloc[6, 8]	# i7 fecha termino
        datos[3] = df.iloc[9, 8]	# i10 muestra
        datos[6] = df.iloc[16, 7]	# h17 peso polvo adm
        datos[7] = df.iloc[16, 8]	# i17 vol agua destilada sol adm
        datos[8] = df.iloc[16, 10]	# k17 peso polvo de roca
        datos[9] = df.iloc[16, 11]	# l17 vol agua destilada polvo de roca
        datos[10] = df.iloc[19, 10]	# k20 total sol adm adicionada
        datos[11] = df.iloc[20, 10]	# k21 total adm adicionado
        datos[12] = df.iloc[21, 10]	# k22 valor adm
    except:
        return[1]
    
    # revisar si hay datos incorrectos
    for dato in datos:
        if pd.isnull(dato):
            return [0]
        
    # Formatear datos
    try:
        datos[0] = datos[0].lstrip().rstrip()
        datos[1] = datos[1].lstrip().rstrip()
        datos[2] = datos[2].lstrip().rstrip()
        datos[3] = datos[3].lstrip().rstrip()
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos[6] = round(float(datos[6]), 1) if isinstance(datos[6], float) else datos[6]
        datos[7] = round(float(datos[7]), 1) if isinstance(datos[7], float) else datos[7]
        datos[8] = round(float(datos[8]), 1) if isinstance(datos[8], float) else datos[8]
        datos[9] = round(float(datos[9]), 1) if isinstance(datos[9], float) else datos[9]
        datos[10] = round(float(datos[10]), 2) if isinstance(datos[10], float) else datos[10]
        datos[11] = round(float(datos[11]), 2) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 2) if isinstance(datos[12], float) else datos[12]
    except:
        return[1]
    
    # observaciones
    obs_rows = df.iloc[:, 7]

    # buscar donde se encuentra la palabra Observaciones
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break

    index += 1
    val = 13

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 7]):
            break
        ob1 = df.iloc[index, 7]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 8]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 8]
            datos[val] = ob1 + ' ' + ob2
        
        # aumentar contadores
        index += 1
        val += 1
    
    for idx, dato in enumerate(datos):
        if dato == "":
            datos[idx] = "-"

    return datos

# Extrae toda la informacion necesaria para la hoja de Etilenglicol
def extraer_info_EG(df):
    datos = [''] * 25
    try:
        datos[0] = df.iloc[0, 1]	# b1 Titulo informe
        datos[1] = df.iloc[2, 2]	# c3 proyecto
        datos[2] = df.iloc[3, 2]	# c4 num informe
        # datos[3] = df.iloc[4, 2]	# c5 orden de trabajo
        datos[4] = df.iloc[5, 2]	# c6 fecha inicio
        datos[5] = df.iloc[6, 2]	# c7 fecha termino
        datos[3] = df.iloc[9, 2]	# c10 muestra
        datos[6] = df.iloc[10, 2]	# c11 tipo de roca
        datos[7] = df.iloc[11, 2]	# c12 fracturas
        datos[8] = df.iloc[12, 2]	# c13 alteraciones
        datos[9] = df.iloc[57, 2]	# c58 dia 1 promedio
        datos[10] = df.iloc[58, 2]	# c59 dia 1 grado
        datos[11] = df.iloc[57, 6]	# g58 dia 5 promedio
        datos[12] = df.iloc[58, 6]	# g59 dia 5 grado
        datos[13] = df.iloc[57, 10]	# k58 dia 10 promedio
        datos[14] = df.iloc[58, 10]	# k59 dia 10 grado
        datos[15] = df.iloc[57, 14]	# o58 dia 15 promedio
        datos[16] = df.iloc[58, 14]	# o59 dia 15 grado
        datos[17] = df.iloc[57, 18]	# s58 dia 20 promedio
        datos[18] = df.iloc[58, 18]	# s59 dia 20 grado
        datos[19] = df.iloc[57, 22]	# w58 dia 30 promedio
        datos[20] = df.iloc[58, 22]	# w59 dia 30 grado
    except:
        return[1]
    
    # revisar si hay datos incorrectos
    for dato in datos:
        if pd.isnull(dato):
            return[0]
        
    # Formatear datos
    try:
        datos[0] = datos[0].lstrip().rstrip()
        datos[1] = datos[1].lstrip().rstrip()
        datos[2] = datos[2].lstrip().rstrip()
        datos[3] = datos[3].lstrip().rstrip()
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos[6] = datos[6].lstrip().rstrip()
        datos[7] = datos[7].lstrip().rstrip()
        datos[8] = datos[8].lstrip().rstrip()
        datos[9] = str(round(float(datos[9]), 2)) if isinstance(datos[9], float) else datos[9]
        datos[10] = str(datos[10])
        datos[11] = str(round(float(datos[11]), 2)) if isinstance(datos[11], float) else datos[11]
        datos[12] = str(datos[12])
        datos[13] = str(round(float(datos[13]), 2)) if isinstance(datos[13], float) else datos[13]
        datos[14] = str(datos[14])
        datos[15] = str(round(float(datos[15]), 2)) if isinstance(datos[15], float) else datos[15]
        datos[16] = str(datos[16])
        datos[17] = str(round(float(datos[17]), 2)) if isinstance(datos[17], float) else datos[17]
        datos[18] = str(datos[18])
        datos[19] = str(round(float(datos[19]), 2)) if isinstance(datos[19], float) else datos[19]
        datos[20] = str(datos[20])
    except:
        return[1]
    
    # observaciones
    obs_rows = df.iloc[:, 1]

    # buscar donde se encuentra la palabra Observaciones
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break
    index += 1
    val = 21

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 1]):
            break
        ob1 = df.iloc[index, 1]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 2]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 2]
            datos[val] = ob1 + ' ' + ob2

        # aumentar contadores
        index += 1
        val += 1
    
    for idx, dato in enumerate(datos):
        if dato == "":
            datos[idx] = "-"

    return datos


def extraer_info_DRX(df):
    datos_base = [''] * 6
    try:
        datos_base[0] = df.iloc[1, 7]	# h2 Titulo informe
        datos_base[1] = df.iloc[3, 8]	# i4 proyecto
        datos_base[2] = df.iloc[4, 8]	# i5 num informe
        # datos_base[3] = df.iloc[5, 8]	# i6 orden de trabajo
        datos_base[3] = df.iloc[6, 8]	# i7 muestra
        datos_base[4] = df.iloc[7, 8]	# i8 fecha inicio
        datos_base[5] = df.iloc[8, 8]	# i9 fecha termino
    except:
        return[1]
    
    # revisar si hay datos incorrectos
    for dato in datos_base:
        if pd.isnull(dato):
            return [0]
        
    # Formatear datos
    datos_base[0] = datos_base[0].lstrip().rstrip()
    datos_base[1] = datos_base[1].lstrip().rstrip()
    datos_base[2] = datos_base[2].lstrip().rstrip()
    datos_base[3] = datos_base[3].lstrip().rstrip()
    datos_base[4] = datos_base[4].strftime("%d-%m-%Y") 
    datos_base[5] = datos_base[5].strftime("%d-%m-%Y")
    
    datos = [0] * 42
    suma = 0

    index = 10
    while not pd.isnull(df.iloc[index, 2]):
        if (df.iloc[index, 2].lower()) == "lau":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[0] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "qz":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[1] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "pl":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[2] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "fls":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[3] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "mica":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[4] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "chl":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[5] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "anh":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[6] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "cal":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[7] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "ttn":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[8] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "cpx":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[9] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "anl":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[10] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "ep":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[11] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "hm" or df.iloc[index, 2].lower() == "hem":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[12] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "arc-sm":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[13] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "prh":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[14] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "prs":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[15] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "mag":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[16] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "act":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[17] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "crs":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[18] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "heu":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[19] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "stb":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[20] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "py":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[21] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "cnp" or df.iloc[index, 2].lower() == "clp":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[22] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "wai":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[23] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "ilm":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[24] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "phl":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[25] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "opal":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[26] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "cha":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[27] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "dol":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[28] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "bar":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[29] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "gp":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[30] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "dat":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[31] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "thm":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[32] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "px":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[33] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "tre-act":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[34] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "ame":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[35] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "cln":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[36] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "pmp":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[37] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "pot":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[38] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "mgs":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[39] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "ca":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[40] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "eps":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[41] = val
            suma += val
    
        index += 1

    datos_obs = [""] * 4

    # observaciones
    obs_rows = df.iloc[:, 7]

    index = 0

    # buscar donde se encuentra la palabra Observaciones
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break
    
    index += 1
    val = 0

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 7]):
            break
        ob1 = df.iloc[index, 7]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 8]):
            datos_obs[val] = ob1
        else:
            ob2 = df.iloc[index, 8]
            datos_obs[val] = ob1 + ' ' + ob2
        
        # aumentar contadores
        index += 1
        val += 1

    datos_suma = [suma]
    datos_final = datos_base
    datos_final.extend(datos_obs)
    datos_final.extend(datos_suma)
    datos_final.extend(datos)
    
    for idx, dato in enumerate(datos_final):
        if dato == "":
            datos_final[idx] = "-"

    return datos_final

def extraer_cantidad_informes(df):
    datos = [0] * 15
    # try:
    datos[0] = df.iloc[0, 1]	# b2 ucs
    datos[1] = df.iloc[1, 1]	# b3 ucs-m
    datos[2] = df.iloc[2, 1]	# b4 hl
    datos[3] = df.iloc[3, 1]	# b5 ch
    datos[4] = df.iloc[4, 1]	# b6 pf
    datos[5] = df.iloc[5, 1]	# b7 plt d
    datos[6] = df.iloc[6, 1]	# b8 plt a
    datos[7] = df.iloc[7, 1]	# b9 plt c
    datos[8] = df.iloc[8, 1]	# b10 ti
    datos[9] = df.iloc[9, 1]	# b11 tx
    datos[10] = df.iloc[10, 1]	# b12 slake
    datos[11] = df.iloc[11, 1]	# b13 adm
    datos[12] = df.iloc[12, 1]	# b14 eg
    datos[13] = df.iloc[13, 1]	# b15 drx
    datos[14] = df.iloc[14, 1]	# b16 mic
    #except:
    #    return[1]
    
    # Revisar la cantidad
    for dato in range(len(datos)):
        if pd.isnull(datos[dato]):
            datos[dato] = 0
        else:
            datos[dato] = int(datos[dato])

    return datos
# -------------- Extraer info Word -------------- #
def procesar_claves_con_parentesis(data_dict):
    pattern = re.compile(r"(.*)\s+\((.*?)\)")
    new_data = {}
    for key, value in data_dict.items():
        match = pattern.match(key)
        if match and "Grado de meteorización" not in key:
            new_key = match.group(1).strip()
            new_value = f"{value} ({match.group(2)})"
            new_data[new_key] = new_value
        else:
            new_data[key] = value
    return new_data

def extraer_informacion_especial(doc):
    texto_acumulado_1 = ""
    encontrado_1 = False
    info = {"Alteración": None, "Nombre de la Roca": None}
    
    comp1 = {}
    comp2 = {}
    componentes_primera_categoria = ['Cristaloclastos', 'Vitroclastos', 'Litoclastos', 'Matriz', 'Fenocristales', 'Masa fundamental']
    
    componentes_segunda_categoria = ['Juveniles', 'Líticos', 'Piroxenos (Px)', 'Piroxeno (Px)', 'Plagioclasa (Plg)', 
                                     'Plagioclasas (Plg)', 'Cuarzo', 'Feldespato potásico (Fd-k)', 'Feldespatos potásicos (Fd-k)', 
                                     'Opacos', 'Opaco', 'Vidrios', 'Vidrio']
    
    for paragraph in doc.paragraphs:
        if 'Imágenes microscópicas' in paragraph.text:
            encontrado_1 = True
            break
        texto_acumulado_1 += " " + paragraph.text

        words = paragraph.text.split()
        
        if words:
            if (len(words)==2):
                componente = words[0].rstrip(':')
                porcentaje = words[1].rstrip(':%')
                
            elif (len(words)==3):
                componente = words[0] + ' ' + words[1]
                porcentaje = words[2].rstrip(':%')
                
            elif (len(words)==4):
                componente = words[0] + ' ' + words[1] + ' ' + words[2]
                porcentaje = words[3].rstrip(':%')
            else:
                componente = None
            
        if componente in componentes_segunda_categoria:            
                
            componente = 'a_' + componente
            if componente in comp1:
                componente = 'b_' + componente
            comp1[componente] = porcentaje
        
        if (componente in componentes_primera_categoria) and (componente not in comp1):
            comp1[componente] = porcentaje

        if len(words)>0 and ("Ceniza" in words[0]):
            categoria = ' '.join(words[:-1]).strip(":")
            info[categoria] = words[-1]
        elif len(words)>0 and ("Lapilli" in words[0]):
            categoria = ' '.join(words[:-1]).strip(":")
            info[categoria] = words[-1]
        elif len(words)>0 and ("Bloques" in words[0]):
            categoria = ' '.join(words[:-1]).strip(":")
            info[categoria] = words[-1]      
            
    if not encontrado_1:
        return "El texto imágenes microscópicas no fue encontrado"
        
    alteracion_matches = re.findall(r"Alteración:\s*(.*?)\.", texto_acumulado_1)    
    roca_matches = re.findall(r"(Nombre de la roca|Clasificación):\s*(.*?)\.", texto_acumulado_1)
        
    #patron_categoria = re.findall("\b(Ceniza|Lapilli|Bloques|Bombas)\b(?:\s*\(<\d+\s*mm\))?:\s*\d+%", texto_acumulado_2)
    patron_porcentaje = re.compile(r'\d+\s*%')
    #categorias = patron_categoria.findall(texto_acumulado_2)

    if alteracion_matches:
        info["Alteración"] = alteracion_matches[-1].strip()
    if roca_matches:
        info["Nombre de la Roca"] = roca_matches[-1][-1]

    info.update(comp1)
        
    return info

def extraer_informacion_archivo(ruta_archivo):
    doc = Document(ruta_archivo)
    archivo = ruta_archivo.split('\\')[-1]
    titulo_principal = doc.paragraphs[0].text.strip()
    informacion = {
        'Nombre del Archivo': archivo,
        'Ensayo': titulo_principal}
    
    data_dict = {}
    for tabla in doc.tables:
        for row in tabla.rows:
            campo = row.cells[0].text.strip()
            campo = re.sub(r'\s+', ' ', campo).strip()
            if campo == "Info. Ensaye":
                campo = "N° Informe"
            valor = row.cells[1].text.strip()
            data_dict[campo] = valor
            
    data_dict = procesar_claves_con_parentesis(data_dict)    
    informacion.update(data_dict)    
   
    info_especial = extraer_informacion_especial(doc)
    informacion.update(info_especial)

    return informacion

# continue
# -------------- Manejo archivos -------------- #
# Crea o agrega a un txt las hojas con errores
def errores(dir, error_logs, informe):
    dir_errores = dir.split("/")
    dir_errores = "/".join(dir_errores[:-1])
    # crea/abre txt
    file_name = dir_errores + "/" + informe + ".txt"
    f = open(file_name, "a+")

    for line in error_logs:
        f.write(line + "\n")
    
    f.close()

# Crea una lista de archivos por directorio y subdirectorios segun extension
def list_files_by_extension(directory):
    excel_files = []
    # Guarda todos los xlsx
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith('xlsx'):
                excel_files.append(os.path.join(root, file))
    # Guarda todos los xlsm
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith('xlsm'):
                excel_files.append(os.path.join(root, file))
    # Agrega los archivos terminados en xls
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith('xls'):
                excel_files.append(os.path.join(root, file))
    return excel_files

# Crea una lista de archivos .docx por directorio y subdirectorios segun extension
def list_docx_by_extension(directory):
    word_files = []
    # Guarda todos los xlsx
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith('.docx'):
                word_files.append(os.path.join(root, file))
            #if file.endswith('.docx'):
            #    word_files.append(os.path.join(root, file))
    return word_files

# -------------- Formateo -------------- #
# Formatea y agrega las filas correspondiente al archivo final
def format_excel(dir_inicial, dir_final):
    # Chequear si las entradas son validas
    if dir_inicial == "":
        message_label.config(text="Por favor seleccione una carpeta")
    elif dir_final == "":
        message_label.config(text="Por favor seleecione alguna archivo")
    elif not dir_final.endswith(".xlsx"):
        message_label.config(text="Por favor seleccione un archivo válido")
    else:
        # Mostrar mensaje en el programa
        message_label.config(text=f'\t\t\t\tActualizando\t\t\t\t')
        root.update_idletasks()
        
        # Abrir archivo errores
        dir_errores = dir_final.split("/")
        dir_errores = "/".join(dir_errores[:-1])
        dir_errores_final = dir_errores + "/" + error_name + ".txt"
        f = open(dir_errores_final, "a+")

        # crear lista con los archivos
        try:
            archivos = list_files_by_extension(dir_inicial)
        except:
            error = f"Algo inesperado ocurrió tratando de abrir la carpeta [{dir_inicial}]"
            f.write(error + "\n")

        # crear listado archivos existentes
        archivos_existentes = pd.ExcelFile(dir_final)

        # extraer info de archivos existentes
        ucs_existente = pd.read_excel(archivos_existentes, name_ucs)
        ucs_m_existente = pd.read_excel(archivos_existentes, name_ucs_m)
        plt_d_existente = pd.read_excel(archivos_existentes, name_plt_d)
        plt_a_existente = pd.read_excel(archivos_existentes, name_plt_a)
        plt_c_existente = pd.read_excel(archivos_existentes, name_plt_c)
        hl_existente = pd.read_excel(archivos_existentes, name_hl)
        ch_existente = pd.read_excel(archivos_existentes, name_ch)
        pf_existente = pd.read_excel(archivos_existentes, name_pf)
        ti_existente = pd.read_excel(archivos_existentes, name_ti)
        tx_existente = pd.read_excel(archivos_existentes, name_tx)
        slake_existente = pd.read_excel(archivos_existentes, name_slake)
        adm_existente = pd.read_excel(archivos_existentes, name_adm)
        eg_existente = pd.read_excel(archivos_existentes, name_eg)
        drx_existente = pd.read_excel(archivos_existentes, name_drx)
        res_existente = pd.read_excel(archivos_existentes, name_res, skiprows=1)
        doc_existente = pd.read_excel(archivos_existentes, name_mic)
        sum_existente = pd.read_excel(archivos_existentes, name_sum)
        tdr_existente = pd.read_excel(archivos_existentes, name_tdr)

        # cambiar fila
        ucs_existente["Nombre de la Hoja"] = ucs_existente["Nombre de la Hoja"].astype(str)
        ucs_m_existente["Nombre de la Hoja"] = ucs_m_existente["Nombre de la Hoja"].astype(str)
        plt_d_existente["Nombre de la Hoja"] = plt_d_existente["Nombre de la Hoja"].astype(str)
        plt_a_existente["Nombre de la Hoja"] = plt_a_existente["Nombre de la Hoja"].astype(str)
        plt_c_existente["Nombre de la Hoja"] = plt_c_existente["Nombre de la Hoja"].astype(str)
        hl_existente["Nombre de la Hoja"] = hl_existente["Nombre de la Hoja"].astype(str)
        ch_existente["Nombre de la Hoja"] = ch_existente["Nombre de la Hoja"].astype(str)
        pf_existente["Nombre de la Hoja"] = pf_existente["Nombre de la Hoja"].astype(str)
        ti_existente["Nombre de la Hoja"] = ti_existente["Nombre de la Hoja"].astype(str)
        tx_existente["Nombre de la Hoja"] = tx_existente["Nombre de la Hoja"].astype(str)
        slake_existente["Nombre de la Hoja"] = slake_existente["Nombre de la Hoja"].astype(str)
        adm_existente["Nombre de la Hoja"] = adm_existente["Nombre de la Hoja"].astype(str)
        eg_existente["Nombre de la Hoja"] = eg_existente["Nombre de la Hoja"].astype(str)
        drx_existente["Nombre de la Hoja"] = drx_existente["Nombre de la Hoja"].astype(str)
        res_existente["N° Muestra"] = res_existente["N° Muestra"].astype(str)
        res_existente["Ingreso"] = res_existente["Ingreso"].astype(str)
        tdr_existente["N° Muestra"] = tdr_existente["N° Muestra"].astype(str)
        tdr_existente["Ingreso"] = tdr_existente["Ingreso"].astype(str)

        # conteo datos
        count_ucs = 0
        count_ucs_m = 0
        count_plt_a = 0
        count_plt_c = 0
        count_plt_d = 0
        count_hl = 0
        count_ch = 0
        count_pf = 0
        count_ti = 0
        count_tx = 0
        count_sla = 0
        count_adm = 0
        count_eg = 0
        count_drx = 0

        # Abrir archivo final
        try:
            wb = load_workbook(dir_final)
        except FileNotFoundError:
            print("No se ha encontrado el archivo")

        # Abrir las hojas corresponientes
        try:
            ws_res = wb.worksheets[pag_res]
            ws_ucs = wb.worksheets[pag_ucs]
            w_ucsm = wb.worksheets[p_ucs_m]
            ws__hl = wb.worksheets[pag__hl]
            ws__ch = wb.worksheets[pag__ch]
            ws__pf = wb.worksheets[pag__pf]
            w_pltd = wb.worksheets[p_plt_d]
            w_plta = wb.worksheets[p_plt_a]
            w_pltc = wb.worksheets[p_plt_c]
            ws__ti = wb.worksheets[pag__ti]
            ws__tx = wb.worksheets[pag__tx]
            ws_sla = wb.worksheets[pag_sla]
            ws_adm = wb.worksheets[pag_adm]
            ws__eg = wb.worksheets[pag__eg]
            ws_drx = wb.worksheets[pag_drx]
            ws_sum = wb.worksheets[pag_sum]
            ws_tdr = wb.worksheets[pag_tdr]
        except:
            print("Falta una página en el resumen")

        # contador de hojas
        count = 0

        # Iterar por campo correspondiente
        for file in archivos:
            if '~$' in file:
                continue
            
            # Abrir excel y sacar datos de archivos
            nombre_archivo = file.split("\\")[-1]
            try:
                num_ingreso = re.search(r"\d\d\d\d", nombre_archivo).group(0)
            except:
                num_ingreso = '-'
            excel = pd.ExcelFile(file)
            sheets = excel.sheet_names

            check_nombre = nombre_archivo.lower()

            for sheet in sheets:
                print(f'Revisando informe: "{nombre_archivo}\tHoja:"{sheet}"')

                # Hojas de configuracio no contienen informacion importante
                if sheet.lower() == "ingreso datos":
                    continue

                count +=1

                # El numero de muestra se encuentra generalmente en la hoja
                # Revisar num de muestra
                try:
                    num_muestra = sheet.split("-")[-1]
                    num_muestra = '0' + str(num_muestra) if len(re.search(r"\d+", str(num_muestra)).group(0)) == 1 else num_muestra
                except:
                    num_muestra = sheet

                time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                # crear la primera parte de la data final
                final_data = [time, nombre_archivo, sheet, num_ingreso, num_muestra]

                # extraer el resto de informacion de la pagina
                page = pd.read_excel(excel, sheet, header=None)

                # Chequear si es un informe UCS-M
                if "ucs-m" in check_nombre:
                    if nombre_archivo in ucs_m_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = ucs_m_existente[ucs_m_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sheet in df_nombre_archivo[["Nombre de la Hoja"]].values or sh in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    # extraer info ucs-m
                    info = extraer_info_UCS_M(page)

                    # Revisar si hay errores
                    if info[0] == 0:
                        error = f"Datos faltantes [UCS-M]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    elif info[0] == 1:
                        error = f"No se pudo formatear el archivo [UCS-M]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    else:
                        final_data.extend(info)
                        w_ucsm.append(final_data)
                        count_ucs_m += 1

                        # encontrar la ubicacion en la tabla
                        suma = 0
                        for key in dict_data:
                            if key == "ucs-m":
                                break
                            suma += dict_data[key]

                        # chequear si existe en los archivos
                        ingreso = final_data[3]
                        muestra = final_data[4]
                        try: 
                            muestra = re.search(r"\d+", final_data[4]).group(0)
                        except:
                            muestra = final_data[4]
                        muestra2 = str(int(muestra)) if muestra.isnumeric() else muestra
                        nom_muestra = final_data[8]
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & ((res_existente["N° Muestra"]==muestra) | (res_existente["N° Muestra"]==muestra2)) & (res_existente["Nombre muestra"]==nom_muestra)]
                        if len(res.index)==0:
                            line = [''] * suma_total_data
                            line[0] = final_data[3]
                            line[1] = final_data[4]
                            line[2] = final_data[3] + '-' + final_data[4]
                            line[3] = final_data[8]
                            line[4] = final_data[6]
                            cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ''
                            except:
                                line[6] = ''

                            # datos
                            line[suma] = final_data[21] # d50
                            line[suma+1] = final_data[22] # e
                            line[suma+2] = final_data[23] # v
                            line[suma+3] = final_data[24] # tipo falla
                            ws_res.append(line)
                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            ws_res.cell(row=idx, column=suma+1).value = final_data[21]
                            ws_res.cell(row=idx, column=suma+2).value = final_data[22]
                            ws_res.cell(row=idx, column=suma+3).value = final_data[23]
                            ws_res.cell(row=idx, column=suma+4).value = final_data[24]

                        tdr = tdr_existente[(tdr_existente["Ingreso"]==ingreso) & ((tdr_existente["N° Muestra"]==muestra) | (tdr_existente["N° Muestra"]==muestra2)) & (tdr_existente["Nombre muestra"]==nom_muestra)]
                        if len(tdr.index) == 0:
                            line = [''] * 29
                            line[0] = final_data[3]
                            line[1] = final_data[4]
                            line[2] = final_data[3] + '-' + final_data[4]
                            line[3] = final_data[8]
                            line[4] = final_data[6]
                            cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ''
                            except:
                                line[6] = ''
                            
                            # datos
                            line[8] = final_data[11]
                            ws_tdr.append(line)
                            
                        else:
                            # updatear datos
                            idx = res.index[0] + 2
                            ws_tdr.cell(row=idx, column=9).value = final_data[11]

                # Chequear si es un informe UCS
                elif "ucs"  in check_nombre or "compresion uniaxial simple" in check_nombre:
                    if nombre_archivo in ucs_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = ucs_existente[ucs_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sheet in df_nombre_archivo[["Nombre de la Hoja"]].values or sh in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    
                    # extraer info ucs
                    info = extraer_info_UCS(page)

                    # Revisar si hay errores
                    if info[0] == 0:
                        error = f"Datos faltantes [UCS]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    elif info[0] == 1:
                        error = f"No se pudo formatear el archivo [UCS]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    else:
                        final_data.extend(info)
                        ws_ucs.append(final_data)
                        count_ucs += 1

                        # encontrar la ubicacion en la tabla
                        suma = 0
                        for key in dict_data:
                            if key == "ucs":
                                break
                            suma += dict_data[key]

                        # chequear si existe en los archivos
                        ingreso = final_data[3]
                        muestra = final_data[4]
                        try: 
                            muestra = re.search(r"\d+", final_data[4]).group(0)
                        except:
                            muestra = final_data[4]
                        muestra2 = str(int(muestra)) if muestra.isnumeric() else muestra
                        nom_muestra = final_data[8]
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & ((res_existente["N° Muestra"]==muestra) | (res_existente["N° Muestra"]==muestra2)) & (res_existente["Nombre muestra"]==nom_muestra)]
                        if len(res.index)==0:
                            # agregar fila nueva
                            line = [''] * suma_total_data
                            line[0] = final_data[3]
                            line[1] = final_data[4]
                            line[2] = final_data[3] + '-' + final_data[4]
                            line[3] = final_data[8]
                            line[4] = final_data[6]
                            cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ''
                            except:
                                line[6] = ''

                            # datos
                            line[suma] = final_data[20]
                            line[suma+1] = final_data[-3].replace("b. ", "")
                            ws_res.append(line)
                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            ws_res.cell(row=idx, column=suma+1).value = final_data[20]
                            ws_res.cell(row=idx, column=suma+2).value = final_data[-3][2:]

                        tdr = tdr_existente[(tdr_existente["Ingreso"]==ingreso) & ((tdr_existente["N° Muestra"]==muestra) | (tdr_existente["N° Muestra"]==muestra2)) & (tdr_existente["Nombre muestra"]==nom_muestra)]
                        if len(tdr.index) == 0:
                            line = [''] * 29
                            line[0] = final_data[3]
                            line[1] = final_data[4]
                            line[2] = final_data[3] + '-' + final_data[4]
                            line[3] = final_data[8]
                            line[4] = final_data[6]
                            cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ''
                            except:
                                line[6] = ''
                            
                            # datos
                            line[7] = final_data[11]
                            ws_tdr.append(line)
                            
                        else:
                            # updatear datos
                            idx = res.index[0] + 2
                            ws_tdr.cell(row=idx, column=8).value = final_data[11]


                # Chequear si es un informe HL
                elif "hl" in check_nombre or "hinchamiento libre" in check_nombre or "h.l" in check_nombre:
                    if nombre_archivo in hl_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = hl_existente[hl_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sheet in df_nombre_archivo[["Nombre de la Hoja"]].values or sh in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    #extraer info hl
                    info = extraer_info_HL(page)

                    # revisar si hay errores
                    if info[0] == 0:
                        error = f"Datos faltantes [HL]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    elif info[0] == 1:
                        error = f"No se pudo formatear el archivo [HL]\t{nombre_archivo}\t{sheet}"
                    else:
                        final_data.extend(info)
                        ws__hl.append(final_data)
                        count_hl += 1

                        # encontrar la ubicacion en la tabla
                        suma = 0
                        for key in dict_data:
                            if key == "hl":
                                break
                            suma += dict_data[key]

                        # chequear si existe en los archivos
                        ingreso = final_data[3]
                        muestra = final_data[4]
                        try: 
                            muestra = re.search(r"\d+", final_data[4]).group(0)
                        except:
                            muestra = final_data[4]
                        muestra2 = str(int(muestra)) if muestra.isnumeric() else muestra
                        nom_muestra = final_data[8]
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & ((res_existente["N° Muestra"]==muestra) | (res_existente["N° Muestra"]==muestra2)) & (res_existente["Nombre muestra"]==nom_muestra)]
                        if len(res.index)==0:
                            line = [''] * suma_total_data
                            line[0] = final_data[3]
                            line[1] = final_data[4]
                            line[2] = final_data[3] + '-' + final_data[4]
                            line[3] = final_data[8]
                            line[4] = final_data[6]
                            cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ''
                            except:
                                line[6] = ''

                            # datos
                            line[suma] = final_data[13]
                            ws_res.append(line)
                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            ws_res.cell(row=idx, column=suma+1).value = final_data[13]

                # Chequear si es un informe de CH
                elif "ch " in check_nombre or "corte hoek" in check_nombre:
                    if nombre_archivo in ch_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = ch_existente[ch_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheet in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    # extraer info ch
                    info = extraer_info_CH(page)
                    if info[0] == 0:
                        error = f"Datos faltantes [CH]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    elif info[0] == 1:
                        error = f"No se pudo formatear el archivo[CH]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    else:
                        final_data.extend(info)
                        ws__ch.append(final_data)
                        count_ch += 1

                        # encontrar la ubicacion en la tabla
                        suma = 0
                        for key in dict_data:
                            if key == "ch":
                                break
                            suma += dict_data[key]

                        # chequear si existe en los archivos
                        ingreso = final_data[3]
                        muestra = final_data[4]
                        try: 
                            muestra = re.search(r"\d+", final_data[4]).group(0)
                        except:
                            muestra = final_data[4]
                        muestra2 = str(int(muestra)) if muestra.isnumeric() else muestra
                        nom_muestra = final_data[8]
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & ((res_existente["N° Muestra"]==muestra) | (res_existente["N° Muestra"]==muestra2)) & (res_existente["Nombre muestra"]==nom_muestra)]
                        if len(res.index)==0:
                            line = [""] * suma_total_data
                            line[0] = final_data[3]
                            line[1] = final_data[4]
                            line[2] = final_data[3] + '-' + final_data[4]
                            line[3] = final_data[8]
                            line[4] = final_data[6]
                            cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ""
                            except:
                                line[6] = ''

                            # datos
                            for index in range(8):
                                line[suma+index] = final_data[14+index]
                            ws_res.append(line)
                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            for index in range(8):
                                ws_res.cell(row=idx, column=suma+index+1).value = final_data[14+index]

                        tdr = tdr_existente[(tdr_existente["Ingreso"]==ingreso) & ((tdr_existente["N° Muestra"]==muestra) | (tdr_existente["N° Muestra"]==muestra2)) & (tdr_existente["Nombre muestra"]==nom_muestra)]
                        if len(tdr.index) == 0:
                            line = [''] * 29
                            line[0] = final_data[3]
                            line[1] = final_data[4]
                            line[2] = final_data[3] + '-' + final_data[4]
                            line[3] = final_data[8]
                            line[4] = final_data[6]
                            cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ''
                            except:
                                line[6] = ''
                            
                            # datos
                            line[9] = final_data[11]
                            ws_tdr.append(line)
                            
                        else:
                            # updatear datos
                            idx = res.index[0] + 2
                            ws_tdr.cell(row=idx, column=10).value = final_data[11]
                
                # Chequear si es un informe de PF
                elif "pf" in check_nombre or "propiedades fisicas" in check_nombre or "propiedades físicas" in check_nombre:
                    if nombre_archivo in pf_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = pf_existente[pf_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sheet in df_nombre_archivo[["Nombre de la Hoja"]].values or sh in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    # extraer info pf
                    info = extraer_info_PF(page)

                    # revisar si hay errores
                    if info[0] == 0:
                        error = f"Datos faltantes [PF]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    elif info[0] == 1:
                        error = f"No se pudo formatear el archivo [PF]\n{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    else:
                        for row in info:
                            data = []
                            data.extend(final_data)
                            data.extend(row[1:])
                            data[4] = row[0]
                            data[4] = '0' + str(data[4]) if len(re.search(r"\d+", str(data[4])).group(0)) == 1 else data[4]
                            nom_muestra = f"{data[10]} {data[11]}" if "(" in data[11] else f"{data[10]} ({data[11]})"
                            cota = data[11]
                            f_ini = data[8]
                            f_ter = data[9]
                            data[8] = nom_muestra
                            data[9] = f_ini
                            data[10] = f_ter
                            data.remove(data[11])
                            ws__pf.append(data)
                            count_pf += 1

                            # encontrar la ubicacion en la tabla
                            suma = 0
                            for key in dict_data:
                                if key == "pf":
                                    break
                                suma += dict_data[key]

                            # chequear si existe en los archivos
                            ingreso = data[3]
                            muestra = data[4]
                            try: 
                                muestra = re.search(r"\d+", data[4]).group(0)
                            except:
                                muestra = data[4]
                            muestra2 = str(int(muestra)) if muestra.isnumeric() else muestra
                            res = res_existente[(res_existente["Ingreso"]==ingreso) & ((res_existente["N° Muestra"]==muestra) | (res_existente["N° Muestra"]==muestra2)) & (res_existente["Nombre muestra"]==nom_muestra)]
                            if len(res.index)==0:
                                line = [''] * suma_total_data
                                line[0] = data[3]
                                line[1] = data[4]
                                line[2] = data[3] + '-' + data[4]
                                line[3] = nom_muestra
                                line[4] = data[6]
                                cotas = str(cota).split("-")
                                line[5] = cotas[0].replace("(", "")
                                try:
                                    line[6] = cotas[1].replace(")", "")
                                except:
                                    line[6] = ''
                                line[suma] = data[11] # humedad
                                line[suma+1] = data[12] # absorcion
                                line[suma+2] = data[13] # densidad
                                line[suma+3] = data[14]
                                ws_res.append(line)
                            else: # updatear datos
                                idx = res.index[0] + 3
                                ws_res.cell(row=idx, column=suma+1).value = data[11]
                                ws_res.cell(row=idx, column=suma+2).value = data[12]
                                ws_res.cell(row=idx, column=suma+3).value = data[13]
                                ws_res.cell(row=idx, column=suma+4).value = data[14]

                # Chequear si es un informe de PLT
                elif "plt" in check_nombre or "carga puntual" in check_nombre:
                    # Chequear si es un informe PLT Axial
                    if "axial" in check_nombre:
                        if nombre_archivo in plt_a_existente["Nombre del Archivo"].values:
                            df_nombre_archivo = plt_a_existente[plt_a_existente["Nombre del Archivo"]==nombre_archivo]
                            try:
                                sh = str(int(sheet))
                            except:
                                sh = sheet
                            if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheet in df_nombre_archivo[["Nombre de la Hoja"]].values:
                                continue
                        # extraer info plt axial
                        info = extraer_info_PLT_axial(page)

                        # Revisar si hay errores
                        if info[0] == 0:
                            error = f"Datos faltantes [PLT-axial]\t{nombre_archivo}\t{sheet}"
                            f.write(error + "\n")
                        elif info[0] == 1:
                            error = f"No se pudo formatear el archivo [PLT-axial]\t{nombre_archivo}\t{sheet}"
                            f.write(error + "\n")
                        else:
                            final_data.extend(info)
                            w_plta.append(final_data)
                            count_plt_a += 1

                            # encontrar la ubicacion en la tabla
                            suma = 0
                            for key in dict_data:
                                if key == "plt":
                                    break
                                suma += dict_data[key]

                            # chequear si existen los archivos
                            ingreso = final_data[3]
                            muestra = final_data[4]
                            try: 
                                muestra = re.search(r"\d+", final_data[4]).group(0)
                            except:
                                muestra = final_data[4]
                            muestra2 = str(int(muestra)) if muestra.isnumeric() else muestra
                            nom_muestra = final_data[8]
                            res = res_existente[(res_existente["Ingreso"]==ingreso) & ((res_existente["N° Muestra"]==muestra) | (res_existente["N° Muestra"]==muestra2)) & (res_existente["Nombre muestra"]==nom_muestra)]
                            if len(res.index)==0:
                                line = [''] * suma_total_data
                                line[0] = final_data[3]
                                line[1] = final_data[4]
                                line[2] = final_data[3] + '-' + final_data[4]
                                line[3] = final_data[8]
                                line[4] = final_data[6]
                                cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                                cotas = cotas.split("-")
                                line[5] = cotas[0]
                                try:
                                    line[6] = cotas[1].replace(")", "")
                                except:
                                    line[6] = ""
                                # datos
                                line[suma] = final_data[42] # resist a la tension
                                ws_res.append(line)
                            else:
                                # updatear datos
                                idx = res.index[0] + 3
                                ws_res.cell(row=idx, column=suma+1).value = final_data[42]
                            
                            tdr = tdr_existente[(tdr_existente["Ingreso"]==ingreso) & ((tdr_existente["N° Muestra"]==muestra) | (tdr_existente["N° Muestra"]==muestra2)) & (tdr_existente["Nombre muestra"]==nom_muestra)]
                            if len(tdr.index) == 0:
                                line = [''] * 29
                                line[0] = final_data[3]
                                line[1] = final_data[4]
                                line[2] = final_data[3] + '-' + final_data[4]
                                line[3] = final_data[8]
                                line[4] = final_data[6]
                                cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                                cotas = cotas.split("-")
                                line[5] = cotas[0]
                                try:
                                    line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ''
                                except:
                                    line[6] = ''
                                
                                # datos
                                line[10] = final_data[11]
                                ws_tdr.append(line)
                                
                            else:
                                # updatear datos
                                idx = res.index[0] + 2
                                ws_tdr.cell(row=idx, column=11).value = final_data[11]

                    # Chequear si es un informe de PLT Colpa
                    elif "colpa" in check_nombre:
                        if nombre_archivo in plt_c_existente["Nombre del Archivo"].values:
                            df_nombre_archivo = plt_c_existente[plt_c_existente["Nombre del Archivo"]==nombre_archivo]
                            try:
                                sh = str(int(sheet))
                            except:
                                sh = sheet
                            if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheet in df_nombre_archivo[["Nombre de la Hoja"]].values:
                                continue
                        # extraer info plt colpa
                        info = extraer_info_PLT_colpa(page)

                        # Revisar si hay errores
                        if info[0] == 0:
                            error = f"Datos faltantes [PLT-Colpa]\t{nombre_archivo}\t{sheet}"
                            f.write(error + "\n")
                        elif info[0] == 1:
                            error = f"No se pudo formatear el archivo [PLT-Colpa]\t{nombre_archivo}\t{sheet}"
                            f.write(error + "\n")
                        else:
                            final_data.extend(info)
                            w_pltc.append(final_data)
                            count_plt_c += 1

                            # encontrar la ubicacion en la tabla
                            suma = 0
                            for key in dict_data:
                                if key == "plt":
                                    break
                                suma += dict_data[key]

                            # chequear si existe en los archivos
                            ingreso = final_data[3]
                            muestra = final_data[4]
                            try: 
                                muestra = re.search(r"\d+", final_data[4]).group(0)
                            except:
                                muestra = final_data[4]
                            muestra2 = str(int(muestra)) if muestra.isnumeric() else muestra
                            nom_muestra = final_data[8]
                            res = res_existente[(res_existente["Ingreso"]==ingreso) & ((res_existente["N° Muestra"]==muestra) | (res_existente["N° Muestra"]==muestra2)) & (res_existente["Nombre muestra"]==nom_muestra)]
                            if len(res.index)==0:
                                line = [''] * suma_total_data
                                line[0] = final_data[3]
                                line[1] = final_data[4]
                                line[2] = final_data[3] + '-' + final_data[4]
                                line[3] = final_data[8]
                                line[4] = final_data[6]
                                cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                                cotas = cotas.split("-")
                                line[5] = cotas[0]
                                try:
                                    line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ""
                                except:
                                    line[6] = ""

                                # datos
                                line[suma+1] = final_data[142]
                                ws_res.append(line)
                            else:
                                # updatear datos
                                idx = res.index[0] + 3
                                ws_res.cell(row=idx, column=suma+2).value = final_data[142]
                            
                            tdr = tdr_existente[(tdr_existente["Ingreso"]==ingreso) & ((tdr_existente["N° Muestra"]==muestra) | (tdr_existente["N° Muestra"]==muestra2)) & (tdr_existente["Nombre muestra"]==nom_muestra)]
                            if len(tdr.index) == 0:
                                line = [''] * 29
                                line[0] = final_data[3]
                                line[1] = final_data[4]
                                line[2] = final_data[3] + '-' + final_data[4]
                                line[3] = final_data[8]
                                line[4] = final_data[6]
                                cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                                cotas = cotas.split("-")
                                line[5] = cotas[0]
                                try:
                                    line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ''
                                except:
                                    line[6] = ''
                                
                                # datos
                                for index in range(10):
                                    line[11+index] = final_data[20+index*13]
                                ws_tdr.append(line)
                                
                            else:
                                # updatear datos
                                idx = res.index[0] + 2
                                for index in range(10):    
                                    ws_tdr.cell(row=idx, column=12+index).value = final_data[20+index*13]

                    # Chequear si es un informe Diametral
                    elif "diametral" in check_nombre:
                        if nombre_archivo in plt_d_existente["Nombre del Archivo"].values:
                            df_nombre_archivo = plt_d_existente[plt_d_existente["Nombre del Archivo"]==nombre_archivo]
                            try:
                                sh = str(int(sheet))
                            except:
                                sh = sheet
                            if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheet in df_nombre_archivo[["Nombre de la Hoja"]].values:
                                continue
                        # extraer info plt diamentral
                        info = extraer_info_PLT_diametral(page)

                        # Revusar si hay errores
                        if info[0] == 0:
                            error = f"Datos faltantes [PLT-Diametral]\t{nombre_archivo}\t{sheet}"
                            f.write(error + "\n")
                        elif info[0]==1:
                            error = f"No se pudo formatear el archivo [PLT-Diametral]\t{nombre_archivo}\t{sheet}"
                            f.write(error + "\n")
                        else:
                            final_data.extend(info)
                            w_pltd.append(final_data)
                            count_plt_d += 1

                            # encontrar la ubicacion en la tabla
                            suma = 0
                            for key in dict_data:
                                if key == "plt":
                                    break
                                suma += dict_data[key]

                            # chequear si existen los archivos
                            ingreso = final_data[3]
                            muestra = final_data[4]
                            try: 
                                muestra = re.search(r"\d+", final_data[4]).group(0)
                            except:
                                muestra = final_data[4]
                            muestra2 = str(int(muestra)) if muestra.isnumeric() else muestra
                            nom_muestra = final_data[8]
                            res = res_existente[(res_existente["Ingreso"]==ingreso) & ((res_existente["N° Muestra"]==muestra) | (res_existente["N° Muestra"]==muestra2)) & (res_existente["Nombre muestra"]==nom_muestra)]
                            if len(res.index)==0:
                                line = [''] * suma_total_data
                                line[0] = final_data[3]
                                line[1] = final_data[4]
                                line[2] = final_data[3] + '-' + final_data[4]
                                line[3] = final_data[8]
                                line[4] = final_data[6]
                                cotas = nom_muestra.split("(")[1] if "(" in str(nom_muestra) else "-"
                                cotas = cotas.split("-")
                                line[5] = cotas[0]
                                try:
                                    line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ""
                                except:
                                    line[6] = ""
                                
                                # datos
                                line[suma+2] = final_data[36] # resistencia a la tension
                                ws_res.append(line)
                            else:
                                # updatear datos
                                idx = res.index[0] + 3
                                ws_res.cell(row=idx, column=suma+3).value = final_data[36]

                            tdr = tdr_existente[(tdr_existente["Ingreso"]==ingreso) & ((tdr_existente["N° Muestra"]==muestra) | (tdr_existente["N° Muestra"]==muestra2)) & (tdr_existente["Nombre muestra"]==nom_muestra)]
                            if len(tdr.index) == 0:
                                line = [''] * 29
                                line[0] = final_data[3]
                                line[1] = final_data[4]
                                line[2] = final_data[3] + '-' + final_data[4]
                                line[3] = final_data[8]
                                line[4] = final_data[6]
                                cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                                cotas = cotas.split("-")
                                line[5] = cotas[0]
                                try:
                                    line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ''
                                except:
                                    line[6] = ''
                                
                                # datos
                                line[21] = final_data[11]
                                ws_tdr.append(line)
                                
                            else:
                                # updatear datos
                                idx = res.index[0] + 2
                                ws_tdr.cell(row=idx, column=22).value = final_data[11]

                    else:
                        error = f"Informe no identificado en los tipos\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")

                # Chequear si es un informe de TI
                elif "ti " in check_nombre or "traccion indirecta" in check_nombre:
                    if nombre_archivo in ti_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = ti_existente[ti_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheet in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                        # Extraer info TI
                    info = extraer_info_TI(page)

                    # revisar si hay errores
                    if info[0] == 0:
                        error = f"Datos faltantes [TI]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    elif info[0] == 1:
                        error = f"No se pudo formatear el archivo [TI]\t{nombre_archivo}\t{sheet}"
                    else:
                        final_data.extend(info)
                        ws__ti.append(final_data)
                        count_ti += 1

                        # encontrar la ubicacion en la tabla
                        suma = 0
                        for key in dict_data:
                            if key == "ti":
                                break
                            suma += dict_data[key]
                        
                        # chequear si existe en los archivos
                        ingreso = final_data[3]
                        muestra = final_data[4]
                        try: 
                            muestra = re.search(r"\d+", final_data[4]).group(0)
                        except:
                            muestra = final_data[4]
                        muestra2 = str(int(muestra)) if muestra.isnumeric() else muestra
                        nom_muestra = final_data[8]
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & ((res_existente["N° Muestra"]==muestra) | (res_existente["N° Muestra"]==muestra2)) & (res_existente["Nombre muestra"]==nom_muestra)]
                        if len(res.index)==0:
                            line = [''] * suma_total_data
                            line[0] = final_data[3]
                            line[1] = final_data[4] 
                            line[2] = final_data[3] + '-' + final_data[4]
                            line[3] = final_data[8]
                            line[4] = final_data[6]
                            cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ""
                            except:
                                line[6] = ""
                            # datos
                            line[suma] = final_data[29] # resist a la tension
                            ws_res.append(line)
                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            ws_res.cell(row=idx, column=suma+1).value = final_data[29]

                        tdr = tdr_existente[(tdr_existente["Ingreso"]==ingreso) & ((tdr_existente["N° Muestra"]==muestra) | (tdr_existente["N° Muestra"]==muestra2)) & (tdr_existente["Nombre muestra"]==nom_muestra)]
                        if len(tdr.index) == 0:
                            line = [''] * 29
                            line[0] = final_data[3]
                            line[1] = final_data[4]
                            line[2] = final_data[3] + '-' + final_data[4]
                            line[3] = final_data[8]
                            line[4] = final_data[6]
                            cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ''
                            except:
                                line[6] = ''
                            
                            # datos
                            line[22] = final_data[11]
                            ws_tdr.append(line)
                            
                        else:
                            # updatear datos
                            idx = res.index[0] + 2
                            ws_tdr.cell(row=idx, column=23).value = final_data[11]


                # chequear si es un informe de tx
                elif "tx" in check_nombre or "tx-m" in check_nombre or "triaxial" in check_nombre:
                    if nombre_archivo in tx_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = tx_existente[tx_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sheet in df_nombre_archivo[["Nombre de la Hoja"]].values or sh in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    # extraer info tx
                    info = extraer_info_TX(page)

                    # Revisar si hay errores
                    if info[0] == 0:
                        error = f"Datos faltantes [TX]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    elif info[0] == 1:
                        error = f"No se pudo formatear el archivo [TX]\t{nombre_archivo}\t{sheet}"
                    else:
                        final_data.extend(info)
                        ws__tx.append(final_data)
                        count_tx += 1

                        # encontrar la ubicacion en la tabla
                        suma = 0
                        for key in dict_data:
                            if key == "trx":
                                break
                            suma += dict_data[key]

                        # chequear si existe en los archivos
                        ingreso = final_data[3]
                        try: 
                            muestra = re.search(r"\d+", final_data[4]).group(0)
                        except:
                            muestra = final_data[4]
                        muestra2 = str(int(muestra)) if muestra.isnumeric() else muestra
                        nom_muestra = final_data[8]
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & ((res_existente["N° Muestra"]==muestra) | (res_existente["N° Muestra"]==muestra2) & (res_existente["Nombre muestra"]==nom_muestra))]
                        if len(res.index)==0:
                            line = [''] * suma_total_data
                            line[0] = final_data[3]
                            line[1] = final_data[4][:-1]
                            line[2] = final_data[3] + '-' + final_data[4][:-1]
                            line[3] = final_data[8]
                            line[4] = final_data[6]
                            cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ""
                            except:
                                line[6] = ""
                            
                            # datos
                            if sheet[-1] == "A":
                                line[suma] = final_data[21]     # mpa
                                line[suma+1] = final_data[22]   # max d50
                                line[suma+2] = final_data[23]   # e mpa
                                line[suma+3] = final_data[24]   # v
                                line[suma+4] = final_data[25]   # tipo de falla
                                ws_res.append(line)
                            elif sheet[-1] == "B":
                                line[suma+5] = final_data[21]   # mpa
                                line[suma+6] = final_data[22]   # max d50
                                line[suma+7] = final_data[23]   # e mpa
                                line[suma+8] = final_data[24]   # v
                                line[suma+9] = final_data[25]   # tipo de falla
                                ws_res.append(line)
                            elif sheet[-1] == "C":
                                line[suma+10] = final_data[21]   # mpa
                                line[suma+11] = final_data[22]   # max d50
                                line[suma+12] = final_data[23]   # e mpa
                                line[suma+13] = final_data[24]   # v
                                line[suma+14] = final_data[25]   # tipo de falla
                                ws_res.append(line)
                            else:
                                error = f"No se pudo formatear el archivo [TX]\t{nombre_archivo}\t{sheet}"
                                f.write(error + "\n")
                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            if sheet[-1] == "A":
                                ws_res.cell(row=idx, column=suma+1).value = final_data[21]
                                ws_res.cell(row=idx, column=suma+2).value = final_data[22]
                                ws_res.cell(row=idx, column=suma+3).value = final_data[23]
                                ws_res.cell(row=idx, column=suma+4).value = final_data[24]
                                ws_res.cell(row=idx, column=suma+5).value = final_data[25]
                            if sheet[-1] == "B":
                                ws_res.cell(row=idx, column=suma+6).value = final_data[21]
                                ws_res.cell(row=idx, column=suma+7).value = final_data[22]
                                ws_res.cell(row=idx, column=suma+8).value = final_data[23]
                                ws_res.cell(row=idx, column=suma+9).value = final_data[24]
                                ws_res.cell(row=idx, column=suma+10).value = final_data[25]
                            if sheet[-1] == "C":
                                ws_res.cell(row=idx, column=suma+11).value = final_data[21]
                                ws_res.cell(row=idx, column=suma+12).value = final_data[22]
                                ws_res.cell(row=idx, column=suma+13).value = final_data[23]
                                ws_res.cell(row=idx, column=suma+14).value = final_data[24]
                                ws_res.cell(row=idx, column=suma+15).value = final_data[25]
                        
                        tdr = tdr_existente[(tdr_existente["Ingreso"]==ingreso) & ((tdr_existente["N° Muestra"]==muestra) | (tdr_existente["N° Muestra"]==muestra2)) & (tdr_existente["Nombre muestra"]==nom_muestra)]
                        if len(tdr.index) == 0:
                            line = [''] * 29
                            line[0] = final_data[3]
                            line[1] = final_data[4]
                            line[2] = final_data[3] + '-' + final_data[4]
                            line[3] = final_data[8]
                            line[4] = final_data[6]
                            cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ''
                            except:
                                line[6] = ''
                            
                            # datos
                            if sheet[-1] == "A":
                                line[23] = final_data[11]
                                ws_tdr.append(line)
                            if sheet[-1] == "B":
                                line[24] = final_data[11]
                                ws_tdr.append(line)
                            if sheet[-1] == "C":
                                line[25] = final_data[11]
                                ws_tdr.append(line)
                            
                        else:
                            # updatear datos
                            idx = res.index[0] + 2
                            if sheet[-1] == "A":
                                ws_tdr.cell(row=idx, column=24).value = final_data[11]
                            if sheet[-1] == "B":
                                ws_tdr.cell(row=idx, column=25).value = final_data[11]
                            if sheet[-1] == "C":
                                ws_tdr.cell(row=idx, column=26).value = final_data[11]

                # chequear si es un informe Slake
                elif "slake" in check_nombre or "sd" in check_nombre:
                    if nombre_archivo in slake_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = slake_existente[slake_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheet in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    
                    # extraer info Slake
                    info = extraer_info_slake(page)

                    # revisar si hay errores
                    if info[0] == 0:
                        error = f"Datos faltantes [Slake]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    elif info[0] == 1:
                        error = f"No se pudo formatear el archivo [Slake]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    else:
                        final_data.extend(info)
                        ws_sla.append(final_data)
                        count_sla += 1

                        # encontrar la ubicacion en la tabla
                        suma = 0
                        for key in dict_data:
                            if key == "sla":
                                break
                            suma += dict_data[key]
                        
                        # chequear si existe en los archivos
                        ingreso = final_data[3]
                        try: 
                            muestra = re.search(r"\d+", final_data[4]).group(0)
                        except:
                            muestra = final_data[4]
                        muestra2 = str(int(muestra)) if muestra.isnumeric() else muestra
                        nom_muestra = final_data[8]
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & ((res_existente["N° Muestra"]==muestra) | (res_existente["N° Muestra"]==muestra2)) & (res_existente["Nombre muestra"] == nom_muestra)]
                        if len(res.index)==0:
                            # agreagar fila nueva
                            line = [''] * suma_total_data
                            line[0] = final_data[3]
                            line[1] = final_data[4]
                            line[2] = final_data[3] + '-' + final_data[4]
                            line[3] = final_data[8]
                            line[4] = final_data[6]
                            cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ""
                            except:
                                line[6] = ""
                            
                            # datos
                            line[suma] = final_data[26]     # ciclo 1
                            line[suma+1] = final_data[28]   # ciclo 1
                            ws_res.append(line)
                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            ws_res.cell(row=idx, column=suma+1).value = final_data[26]
                            ws_res.cell(row=idx, column=suma+2).value = final_data[28]
                        
                        tdr = tdr_existente[(tdr_existente["Ingreso"]==ingreso) & ((tdr_existente["N° Muestra"]==muestra) | (tdr_existente["N° Muestra"]==muestra2)) & (tdr_existente["Nombre muestra"]==nom_muestra)]
                        if len(tdr.index) == 0:
                            line = [''] * 29
                            line[0] = final_data[3]
                            line[1] = final_data[4]
                            line[2] = final_data[3] + '-' + final_data[4]
                            line[3] = final_data[8]
                            line[4] = final_data[6]
                            cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ''
                            except:
                                line[6] = ''
                            
                            # datos
                            line[26] = final_data[11]
                            ws_tdr.append(line)
                            
                        else:
                            # updatear datos
                            idx = res.index[0] + 2
                            ws_tdr.cell(row=idx, column=27).value = final_data[11]

                # chequear si es un informe de azul de metileno
                elif "azul de metileno" in check_nombre or "am " in check_nombre or "azul" in check_nombre or "metileno" in check_nombre:
                    if nombre_archivo in adm_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = adm_existente[adm_existente["Nombre del Archivo"]==nombre_archivo]
                        try: 
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheet in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    # extraer info azul de metileno
                    info = extraer_info_azul_de_metileno(page)

                    # Revisar si hay errores
                    if info[0] == 0:
                        error = f"Datos faltantes [Azul de Metileno]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    elif info[0] == 1:
                        error = f"No se pudo formatear el archivo [Azul de Metileno]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    else:
                        final_data.extend(info)
                        ws_adm.append(final_data)
                        count_adm += 1

                        # encontrar la ubicacion en la tabla
                        suma = 0
                        for key in dict_data:
                            if key == "am":
                                break
                            suma += dict_data[key]
                        
                        # chequear si existe en los archivos
                        ingreso = final_data[3]
                        try: 
                            muestra = re.search(r"\d+", final_data[4]).group(0)
                        except:
                            muestra = final_data[4]
                        muestra2 = str(int(muestra)) if muestra.isnumeric() else muestra
                        nom_muestra = final_data[8]
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & ((res_existente["N° Muestra"]==muestra) | (res_existente["N° Muestra"]==muestra2)) & (res_existente["Nombre muestra"]==nom_muestra)]
                        if len(res.index)==0:
                            line = [''] * suma_total_data
                            line[0] = final_data[3]
                            line[1] = final_data[4]
                            line[2] = final_data[3] + '-' + final_data[4]
                            line[3] = final_data[8]
                            line[4] = final_data[6]
                            cotas = final_data[8].split("(") [1] if "(" in str(final_data[8]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ""
                            except:
                                line[6] = ""

                            # datos
                            line[suma] = final_data[15]     # adicionado
                            line[suma+1] = final_data[17]   # mbv
                            ws_res.append(line)
                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            ws_res.cell(row=idx, column=suma+1).value = final_data[15]
                            ws_res.cell(row=idx, column=suma+2).value = final_data[17]
                
                # chequear si es un informe de etilenglicol
                elif "etilenglicol" in check_nombre or "eg" in check_nombre:
                    if nombre_archivo in eg_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = eg_existente[eg_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sheet in df_nombre_archivo[["Nombre de la Hoja"]].values or sh in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    # extraer info etilenglicol
                    info = extraer_info_EG(page)
                    if info[0] == 0:
                        error = f"Datos faltantes [EG]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    elif info[0] == 1:
                        error = f"No se pudo formatear el archivo [EG]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    else:
                        final_data.extend(info)
                        ws__eg.append(final_data)
                        count_eg += 1

                        # encontrar la ubicacion en la tabla
                        suma = 0
                        for key in dict_data:
                            if key == "eg":
                                break
                            suma += dict_data[key]
                        
                        # chequear si existe en los archivos
                        ingreso = final_data[3]
                        try: 
                            muestra = re.search(r"\d+", final_data[4]).group(0)
                        except:
                            muestra = final_data[4]
                        muestra2 = str(int(muestra)) if muestra.isnumeric() else muestra
                        nom_muestra = final_data[8]
                        res = res_existente[((res_existente["Ingreso"]==ingreso) & ((res_existente["N° Muestra"]==muestra) | (res_existente["N° Muestra"]==muestra2)) & (res_existente["Nombre muestra"] == nom_muestra))]
                        if len(res.index)==0:
                            line = [''] * suma_total_data
                            line[0] = final_data[3]
                            line[1] = final_data[4]
                            line[2] = final_data[3] + '-' + final_data[4]
                            line[3] = final_data[8]
                            line[4] = final_data[6]
                            cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ""
                            except:
                                line[6] = ""
                            
                            # datos
                            for index in range(12):
                                line[suma+index] = final_data[14+index]
                            ws_res.append(line)
                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            for index in range(12):
                                ws_res.cell(row=idx, column=suma+index+1).value = final_data[14+index]
                        
                        tdr = tdr_existente[(tdr_existente["Ingreso"]==ingreso) & ((tdr_existente["N° Muestra"]==muestra) | (tdr_existente["N° Muestra"]==muestra2)) & (tdr_existente["Nombre muestra"]==nom_muestra)]
                        if len(tdr.index) == 0:
                            line = [''] * 29
                            line[0] = final_data[3]
                            line[1] = final_data[4]
                            line[2] = final_data[3] + '-' + final_data[4]
                            line[3] = final_data[8]
                            line[4] = final_data[6]
                            cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ''
                            except:
                                line[6] = ''
                            
                            # datos
                            line[27] = final_data[11]
                            ws_tdr.append(line)
                            
                        else:
                            # updatear datos
                            idx = res.index[0] + 2
                            ws_tdr.cell(row=idx, column=28).value = final_data[11]

                # chequear si es un informe de drx
                elif "drx" in check_nombre or "difracciones rayos x" in check_nombre:
                    if sheet.lower() == "datos iniciales":
                        continue
                    if nombre_archivo in drx_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = drx_existente[drx_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheets in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    # extraer info drx
                    info = extraer_info_DRX(page)

                    # Revisar si hay errores
                    if info[0] == 0:
                        error = f"Datos faltantes [DRX]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    elif info[0] == 1:
                        error = f"No se pudo formatear el archivo [DRX]\t{nombre_archivo}\t{sheet}"
                        f.write(error + "\n")
                    else:
                        final_data.extend(info)
                        ws_drx.append(final_data)
                        count_drx += 1

                        if info[10] != 100:
                            error = f"DRX no suma 100 [DRX]\t{nombre_archivo}\t{sheet}"
                            f.write(error + "\n")

                        # encontrar ubicacion en la tabla
                        suma = 0
                        for key in dict_data:
                            if key == "drx":
                                break
                            suma += dict_data[key]

                        # chequear si existen los archivos
                        ingreso = final_data[3]
                        muestra = final_data[4]
                        muestra2 = str(int(muestra)) if muestra.isnumeric() else muestra
                        nom_muestra = final_data[8]
                        muestra = muestra.lower().replace("v", "") if not "infilling" in nom_muestra.lower() else muestra
                        muestra2 = muestra2[1:] if muestra2[0] == '0' else muestra2
                        muestra2 = muestra2.lower().replace("v", "") if not "infilling" in nom_muestra.lower() else muestra2
                        res = res_existente[((res_existente["Ingreso"]==ingreso) & ((res_existente["N° Muestra"]==muestra) | (res_existente["N° Muestra"]==muestra2) & (res_existente["Nombre muestra"]==nom_muestra)))]
                        if len(res.index)==0:
                            line = [''] * suma_total_data
                            line[0] = final_data[3]
                            line[1] = final_data[4]
                            line[2] = final_data[3] + '-' + final_data[4]
                            line[3] = final_data[8]
                            line[4] = final_data[6]
                            cotas = final_data[8].split("(")[1] if "(" in str(final_data[8]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ""
                            except:
                                line[6] = ""
                            
                            # datos
                            for index in range(dict_data["drx"]):
                                line[suma+index] = final_data[16+index]
                            ws_res.append(line)
                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            for index in range(dict_data["drx"]):
                                ws_res.cell(row=idx, column=suma+index+1).value = final_data[16+index]

                else:
                    error = f"Informe no identificado en los tipos\t{nombre_archivo}\t{sheet}"
                    f.write(error + "\n")

        print(f"Se han actualizado {count} hojas")
        
        # Agregar sumas
        info = extraer_cantidad_informes(sum_existente)
        info[0] += count_ucs
        info[1] += count_ucs_m
        info[2] += count_hl
        info[3] += count_ch
        info[4] += count_pf
        info[5] += count_plt_d
        info[6] += count_plt_a
        info[7] += count_plt_c
        info[8] += count_ti
        info[9] += count_tx
        info[10] += count_sla
        info[11] += count_adm
        info[12] += count_eg
        info[13] += count_drx
        
        for idx in range(14):
            ws_sum.cell(row=idx+2, column=2).value = info[idx]
        
        wb.save(dir_final)


        # Formatear archivos docx
        try:
            archivos = list_docx_by_extension(dir_inicial)
            # archivos = []
        except:
            error = ["Algo inesperado ocurrio tratando de abrir la carpeta"]
            for line in error:
                f.write(line + "\n")

        try:
            wb = load_workbook(dir_final)
            ws_mic = wb.worksheets[pag_mic]
            ws_res = wb.worksheets[pag_res]
            ws_sum = wb.worksheets[pag_sum]
            ws_tdr = wb.worksheets[pag_tdr]
        except FileNotFoundError:
            print("No se ha encontrado el archivo")

        for file in archivos:
            count += 1
            if file[:2] == '~$':
                continue
            

            nombre_archivo = file.split("\\")[-1]

            if nombre_archivo in doc_existente["Nombre del Archivo"].values:
                continue

            print(f"Revisando Informe {nombre_archivo}")
            

            time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            ingreso = re.search(r"\d\d\d\d", nombre_archivo).group(0).rstrip().lstrip()
            muestra = muestra = re.search(r"-\d([A-Za-z]|\d)*", nombre_archivo).group(0).replace("-", "")
            muestra = str(int(muestra)) if isinstance(muestra, int) else muestra
            muestra = muestra[1:] if muestra[0] == '0' else muestra
            muestra = muestra.lstrip().rstrip()
            muestra2 = '0' + str(muestra) if len(re.search(r"\d+", str(muestra)).group(0))==1 else muestra

            info_archivos = [""] * 34
            info_archivos[0] = time     # hora
            info_archivos[2] = ingreso  # ingreso
            info_archivos[3] = muestra2 # num muestra


            line = extraer_informacion_archivo(file)
            for key in line.keys():
                if key == "Nombre del Archivo":
                    info_archivos[1] = line[key]
                if key == "Ensayo":
                    info_archivos[4] = line[key]
                if key == "Nombre Proyecto":
                    info_archivos[5] = line[key]
                if key == "N° Informe":
                    info_archivos[6] = line[key]
                if key == "Fecha Inicio":
                    info_archivos[8] = line[key]
                if key == "Fecha Término":
                    info_archivos[9] = line[key]
                if key == "Muestra":
                    info_archivos[7] = line[key]
                if key == "1. Textura":
                    info_archivos[10] = line[key]
                if key == "2. Grado de cristalinidad":
                    info_archivos[11] = line[key]
                if key == "3. Tamaño relativo de cristales":
                    info_archivos[12] = line[key]
                if key == "4. Tamaño absoluto de los cristales":
                    info_archivos[13] = line[key]
                if key == "5. Tamaño del grano":
                    info_archivos[14] = line[key]
                if key == "6. Forma de los cristales":
                    info_archivos[15] = line[key]
                if key == "7. Estructura":
                    info_archivos[16] = line[key]
                if key == "8. Fábrica":
                    info_archivos[17] = line[key]
                if key == "9. Morfología especial":
                    info_archivos[18] = line[key]
                if key == "10. Índice de color":
                    info_archivos[19] = line[key]
                if key == "11. Grado de meteorización (ISRM, 1981)":
                    info_archivos[20] = line[key]
                if key == "Alteración":
                    info_archivos[21] = line[key]
                if key == "Nombre de la Roca":
                    info_archivos[22] = line[key]
                if key == "Ceniza (<2 mm)":
                    info_archivos[23] = line[key]
                if key == "Lapilli (2 - 64 mm)":
                    info_archivos[24] = line[key]
                if key == "Bloques y Bombas (>64 mm)":
                    info_archivos[25] = line[key]
                if key == "Cristaloclastos":
                    info_archivos[26] = line[key]
                if key == "a_Plagioclasas (Plg)":
                    info_archivos[27] = line[key]
                if key == "a_Piroxenos (Px)":
                    info_archivos[28] = line[key]
                if key == "Litoclastos":
                    info_archivos[29] = line[key]
                if key == "Vitroclastos":
                    info_archivos[30] = line[key]
                if key == "Matriz":
                    info_archivos[31] = line[key]
                if key == "Fenocristales":
                    info_archivos[32] = line[key]
                if key == "Masa fundamental":
                    info_archivos[33] = line[key]

            for idx, data in enumerate(info_archivos):
                if data == "":
                    info_archivos[idx] = '-'
                    
            ws_mic.append(info_archivos)
            info[14] += 1

            # encontrar ubicacion en la tabla
            suma = 0
            for key in dict_data:
                if key == "mic":
                    break
                suma += dict_data[key]

            nom_muestra = info_archivos[7]

            # chequear si existen los archivos
            res = res_existente[(res_existente["Ingreso"]==ingreso) & ((res_existente["N° Muestra"]==muestra) | (res_existente["N° Muestra"]==muestra2)) & (res_existente["Nombre muestra"]==nom_muestra)]
            if len(res.index)==0:
                # agregar fila nueva
                line = [''] * suma_total_data
                line[0] = ingreso
                line[1] = muestra
                line[2] = ingreso + '-' + muestra
                line[3] = info_archivos[7]
                line[4] = info_archivos[5]
                cotas = info_archivos[7].split("(")[1] if "(" in info_archivos[7] else "-"
                cotas = cotas.split("-")
                line[5] = cotas[0]
                try:
                    line[6] = cotas[1].replace(')', '')
                except:
                    line[6] = ''

                final_data = info_archivos
                # datos
                line[suma] = final_data[22]
                line[suma+1] = final_data[21]
                line[suma+2] = final_data[10]
                line[suma+3] = final_data[11]
                line[suma+4] = final_data[12]
                line[suma+5] = final_data[13]
                line[suma+6] = final_data[14]
                line[suma+7] = final_data[15]
                line[suma+8] = final_data[16]
                line[suma+9] = final_data[17]
                line[suma+10] = final_data[18]
                line[suma+11] = final_data[19]
                line[suma+12] = final_data[20]
                line[suma+13] = final_data[23]
                line[suma+14] = final_data[24]
                line[suma+15] = final_data[25]
                line[suma+16] = final_data[26]
                line[suma+17] = final_data[29]
                line[suma+18] = final_data[30]
                line[suma+19] = final_data[31]
                line[suma+20] = final_data[32]
                line[suma+21] = final_data[33]
                ws_res.append(line)

            else:
                # updatear datos
                final_data = info_archivos
                idx = res.index[0] + 3
                ws_res.cell(row=idx, column=suma+1).value = final_data[22]
                ws_res.cell(row=idx, column=suma+2).value = final_data[21]
                ws_res.cell(row=idx, column=suma+3).value = final_data[10]
                ws_res.cell(row=idx, column=suma+4).value = final_data[11]
                ws_res.cell(row=idx, column=suma+5).value = final_data[12]
                ws_res.cell(row=idx, column=suma+6).value = final_data[13]
                ws_res.cell(row=idx, column=suma+7).value = final_data[14]
                ws_res.cell(row=idx, column=suma+8).value = final_data[15]
                ws_res.cell(row=idx, column=suma+9).value = final_data[16]
                ws_res.cell(row=idx, column=suma+10).value = final_data[17]
                ws_res.cell(row=idx, column=suma+11).value = final_data[18]
                ws_res.cell(row=idx, column=suma+12).value = final_data[19]
                ws_res.cell(row=idx, column=suma+13).value = final_data[20]
                ws_res.cell(row=idx, column=suma+14).value = final_data[23]
                ws_res.cell(row=idx, column=suma+15).value = final_data[24]
                ws_res.cell(row=idx, column=suma+16).value = final_data[25]
                ws_res.cell(row=idx, column=suma+17).value = final_data[26]
                ws_res.cell(row=idx, column=suma+18).value = final_data[29]
                ws_res.cell(row=idx, column=suma+19).value = final_data[30]
                ws_res.cell(row=idx, column=suma+20).value = final_data[31]
                ws_res.cell(row=idx, column=suma+21).value = final_data[32]
                ws_res.cell(row=idx, column=suma+22).value = final_data[33]

            tdr = tdr_existente[(tdr_existente["Ingreso"]==ingreso) & ((tdr_existente["N° Muestra"]==muestra) | (tdr_existente["N° Muestra"]==muestra2)) & (tdr_existente["Nombre muestra"]==nom_muestra)]
            if len(tdr.index) == 0:
                line = [''] * 29
                line[0] = ingreso
                line[1] = muestra
                line[2] = ingreso + '-' + muestra
                line[3] = info_archivos[7]
                line[4] = info_archivos[5]
                cotas = info_archivos[7].split("(")[1] if "(" in str(info_archivos[7]) else "-"
                cotas = cotas.split("-")
                line[5] = cotas[0]
                try:
                    line[6] = cotas[1].replace(")", "") if ")" in cotas[1] else ''
                except:
                    line[6] = ''
                
                # datos
                line[28] = final_data[22]
                ws_tdr.append(line)
                
            else:
                # updatear datos
                idx = res.index[0] + 2
                ws_tdr.cell(row=idx, column=29).value = final_data[22]

        # revisar
        ws_sum.cell(row=16, column=2).value = info[14]
        wb.save(dir_final)

        error = [f'---------------Fecha {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}---------------']
        for line in error:
            f.write(line + "\n")
        message_label.config(text="Se ha terminado la actualizacion")
        print(f"Se han analizado {count} hojas")

# -------------- Funciones Programa -------------- #
# Seleccionar carpeta
def select_folder():
    folder_path = filedialog.askdirectory()
    if folder_path:
        print(f"Selected folder: {folder_path}")
        folder_selected.set(folder_path)

# Seleccionar archivo
def select_file():
    file_path = filedialog.askopenfilename()
    if file_path:
        print(f"Selected file: {file_path}")
        file_selected.set(file_path)

# Funcion para actualizar
def update():
    folder_path = select_folder()
    file_path = select_file()
    if folder_path and file_path:
        format_excel(folder_path, file_path)


# -------------- Programa -------------- #
root = tk.Tk()
root.title("Actualización de Informes")

# Variables para guardar los paths
folder_selected = tk.StringVar()
file_selected = tk.StringVar()

# Mostrar status de la funcion
message_label = tk.Label(root, text="Seleccionar carpeta y archivo, luego presionar 'ejecutar datos'.")
message_label.grid(row=3, columnspan=2, padx=10, pady=10)

# Boton de carpeta
folder_button = tk.Button(root, text="Seleccionar Carpeta", command=select_folder)
folder_button.grid(row=0, column=0, padx=10, pady=10)

# Entry to display selected folder path
folder_label = tk.Label(root, text="Seleccione la carpeta donde se encuentran los archivos que se van a actualizar")
folder_label.grid(row=0, column=1, padx=10, pady=10)

# Boton para archivo
file_button = tk.Button(root, text="Seleccionar archivo", command=select_file)
file_button.grid(row=1, column=0, padx=10, pady=10)

# Entry to display selected file path
file_label = tk.Label(root, text="Seleccione el archivo de Excel donde se encuentran los resumenes de informes")
file_label.grid(row=1, column=1, padx=10, pady=10) 

# Boton para la funcion
run_button = tk.Button(root, text="Actualizar datos", command=lambda: format_excel(folder_selected.get(), file_selected.get()))
run_button.grid(row=2, columnspan=2, padx=10, pady=10)

# Ejecutar Tkinter
root.mainloop()
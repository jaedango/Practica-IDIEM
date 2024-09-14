# imports 
import pandas as pd
import os
import datetime
import re
from openpyxl import Workbook
from openpyxl import load_workbook
import tkinter as tk
from tkinter import filedialog
from docx import Document

# Ignore Warnings
import warnings
warnings.simplefilter(action='ignore', category=UserWarning)

# -------------- Variables globales -------------- #
# Orden de las hojas
# para cambiar el orden de las hojas cambiar los numeros
pag_ucs = 1
pag_plt_d = 6
pag_plt_a = 7
pag_plt_c = 8
pag_hl = 3
pag_ch = 4
pag_pf = 5
pag_ti = 9
pag_tx = 10
pag_slake = 11
pag_adm = 12
pag_eg = 13
pag_drx = 14
pag_res = 0
pag_ucs_m = 2
pag_mic = 15

# Nombres de las hojas
# en caso que se quiera cambiar los nombres de las hojas cambiar aca
name_ucs = "UCS"
name_plt_d = "PLT Diametral"
name_plt_a = "PLT Axial"
name_plt_c = "PLT Colpa"
name_hl = "Hinchamiento Libre"
name_ch = "Corte Hoek"
name_pf = "Propiedades Físicas"
name_ti = "Tracción Indirecta"
name_tx = "TX-M"
name_slake = "Slake"
name_adm = "Azul de Metileno"
name_eg = "Etilenglicol"
name_drx = "DRX"
name_res = "Resumen"
name_ucs_m = "UCS-M"
name_mic = "Microscopía"

# nombre de la pagina de errores
error_name = "Informe_errores"

# -------------- Extraer info -------------- #
# Extrae toda la informacion necesaria para la hoja de UCS
def extraer_info_UCS(df):
    datos = [""] * 22

    try:
        datos[0] = df.iloc[0, 9]    # j1 titulo_informe
        datos[1] = df.iloc[2, 11]   # l3 proyecto
        datos[2] = df.iloc[3, 11]   # l4 n_informe
        datos[3] = df.iloc[4, 11]   # l5 orden_trabajo
        datos[4] = df.iloc[5, 11]   # l6 fecha_inicio
        datos[5] = df.iloc[6, 11]   # l7 fecha_termino
        datos[6] = df.iloc[9, 11]   # l10 muestra
        datos[7] = df.iloc[10, 11]  # l11 tipo_roca
        datos[8] = df.iloc[11, 11]  # l12 fracturas
        datos[9] = df.iloc[12, 11]  # l13 alteraciones
        datos[10] = df.iloc[13, 11] # l14 observaciones
        datos[11] = df.iloc[16, 11] # l17 diametro
        datos[12] = df.iloc[17, 11] # l18 altura
        datos[13] = df.iloc[18, 11] # l19 densidad_h
        datos[14] = df.iloc[19, 11] # l20 densidad_s
        datos[15] = df.iloc[20, 11] # l21 contenido_h
        datos[16] = df.iloc[16, 17] # r17 resist_d50
        datos[17] = df.iloc[17, 17] # r18 resist_d64

    except:
        return [1]
    
    # Revisar si hay algun dato incorrecto
    for dato in datos:
        if pd.isnull(dato):
            # return manejo_errores(datos, name_ucs)
            return [0]
        
    # formatear datos
    try:
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y") 
        datos[6] = datos[6].lstrip(" ").rstrip(" ")
        datos[7] = datos[7].lstrip(" ").rstrip(" ")
        datos[8] = datos[8].lstrip(" ").rstrip(" ")
        datos[9] = datos[9].lstrip(" ").rstrip(" ")
        datos[10] = datos[10].lstrip(" ").rstrip(" ")
        datos[11] = round(float(datos[11]), 2) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 2) if isinstance(datos[12], float) else datos[12]
        datos[13] = round(float(datos[13]), 2) if isinstance(datos[13], float) else datos[13]
        datos[14] = round(float(datos[14]), 2) if isinstance(datos[14], float) else datos[14]
        datos[15] = round(float(datos[15]), 2) if isinstance(datos[15], float) else datos[15]
        datos[16] = round(float(datos[16]), 1) if isinstance(datos[16], float) else datos[16]
        datos[17] = round(float(datos[17]), 1) if isinstance(datos[17], float) else datos[17]
    except:
        return[1]
    
    # observaciones
    obs = []
    obs_rows = df.iloc[22:, 9]

    # buscar donde se enceuntra la palabra Observaciones
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break
    index += 1  # aumentar index para no seguir directamente desde 'Observaciones'
    val = 18    # valor para seguir en el arreglo de datos

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 9]):
            break
        ob1 = df.iloc[index, 9]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 10]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 10]
            datos[val] = ob1 + ' ' + str(ob2)

        # Aumentar contadores
        index += 1
        val += 1

    return datos

def extraer_info_UCS_M(df):
    datos = [""] * 24

    try:
        datos[0] = df.iloc[0, 10]	# k1 titulo informe
        datos[1] = df.iloc[2, 11]	# l3 proyecto
        datos[2] = df.iloc[3, 11]	# l4 num informe
        datos[3] = df.iloc[4, 11]	# l5 orden trabajo
        datos[4] = df.iloc[5, 11]	# l6 fecha inicio
        datos[5] = df.iloc[6, 11]	# l7 fecha termino
        datos[6] = df.iloc[8, 11]	# l9 muestra
        datos[7] = df.iloc[9, 11]	# l10 tipo de roca
        datos[8] = df.iloc[10, 11]	# l11 fracturas
        datos[9] = df.iloc[11, 11]	# l12 alteraciones
        datos[10] = df.iloc[14, 11]	# l15 diametro
        datos[11] = df.iloc[15, 11]	# l16 altura
        datos[12] = df.iloc[16, 11]	# l17 peso
        datos[13] = df.iloc[17, 11]	# l18 densidad h
        datos[14] = df.iloc[18, 11]	# l19 densidad s
        datos[15] = df.iloc[19, 11]	# l20 humedad
        datos[16] = df.iloc[14, 14]	# o15 resist d50
        datos[17] = df.iloc[15, 14]	# o16 mod def e
        datos[18] = df.iloc[16, 14]	# o17 poisson
        datos[19] = df.iloc[17, 14]	# o18 muestra falla por
    except:
        return [1]

    # Revisar si hay algun dato incorrecto
    for dato in datos:
        if pd.isnull(dato):
            # return manejo_errores(datos, name_ucs)
            return [0]
        
    # formatear datos
    try:
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y") 
        datos[6] = datos[6].lstrip(" ").rstrip(" ")
        datos[7] = datos[7].lstrip(" ").rstrip(" ")
        datos[8] = datos[8].lstrip(" ").rstrip(" ")
        datos[9] = datos[9].lstrip(" ").rstrip(" ")
        datos[10] = round(float(datos[10]), 2) if isinstance(datos[10], float) else datos[10]
        datos[11] = round(float(datos[11]), 2) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 2) if isinstance(datos[12], float) else datos[12]
        datos[13] = round(float(datos[13]), 2) if isinstance(datos[13], float) else datos[13]
        datos[14] = round(float(datos[14]), 2) if isinstance(datos[14], float) else datos[14]
        datos[15] = round(float(datos[15]), 2) if isinstance(datos[15], float) else datos[15]
        datos[16] = round(float(datos[16]), 1) if isinstance(datos[16], float) else datos[16]
        datos[17] = round(float(datos[17]), 1) if isinstance(datos[17], float) else datos[17]
        datos[18] = round(float(datos[18]), 1) if isinstance(datos[18], float) else datos[18]
    except:
        return[1]
    
    # observaciones
    obs = []
    obs_rows = df.iloc[:, 10]

    # buscar donde se enceuntra la palabra Observaciones
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones:':
            break
    index += 1  # aumentar index para no seguir directamente desde 'Observaciones'
    val = 20    # valor para seguir en el arreglo de datos

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 10]):
            break
        ob1 = df.iloc[index, 10]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 11]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 11]
            datos[val] = ob1 + ' ' + str(ob2)

        # Aumentar contadores
        index += 1
        val += 1

    return datos

# Extrae toda la informacion necesaria para la hoja de hinchamiento libre
def extraer_info_HL(df):
    datos = [''] * 14
    try:
        datos[0] = df.iloc[0, 5] # Titulo informe
        datos[1] = df.iloc[2, 6] # Proyecto
        datos[2] = df.iloc[3, 6] # N° Informe
        datos[3] = df.iloc[4, 6] # Orden de trabajo
        datos[4] = df.iloc[5, 6] # Fecha Inicio
        datos[5] = df.iloc[6, 6] # Fecha Termino
        datos[6] = df.iloc[9, 6] # Muestra
        datos[7] = df.iloc[12,7] # Vol incial
        datos[8] = df.iloc[13,7] # Vol final
        datos[9] = df.iloc[16,7] # Indice H
    except:
        return[1]

    # Revisar si hay datos incorrectos
    for dato in datos:
        if pd.isnull(dato):
            # return manejo_errores(datos, name_hl)
            return [0]
        
    # Formatear datos
    datos[4] = datos[4].strftime("%d-%m-%Y") 
    datos[5] = datos[5].strftime("%d-%m-%Y") 

    # observaciones
    obs = []
    obs_rows = df.iloc[:, 5]

    # buscar donde se encuentra la palabra 'Observaciones'
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break
    index += 1  # aumentar index para no seguir directamente desde obs
    val = 10    # valor para seguir el arreglo de datos
    
    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 5]):
            break
        obs1 = df.iloc[index, 5]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 6]):
            datos[val] = obs1
        else:
            obs2 = df.loc[index, 6]
            datos[val] = obs[1] + ' ' + str(obs2)
        # aumentar contadores
        index += 1
        val += 1
    return datos

def extraer_info_CH(df):
    datos =[''] * 22
    try:
        datos[0] = df.iloc[0, 4]    # E1 Titulo Informe
        datos[1] = df.iloc[2, 5]    # F3 Proyecto
        datos[2] = df.iloc[3, 5]    # F4 N Informe
        datos[3] = df.iloc[4, 5]    # F5 Orden de trabajo
        datos[4] = df.iloc[5, 5]    # f6 Fecha inicio
        datos[5] = df.iloc[6, 5]    # f7 fecha termino
        datos[6] = df.iloc[8, 5]    # f9 muestra
        datos[7] = df.iloc[9, 5]    # f10 tipo roca
        datos[8] = df.iloc[10, 5]   # f11 fracturas
        datos[9] = df.iloc[11, 5]   # f12 alteraciones
        datos[10] = df.iloc[59, 16]	# q60 t normal 1
        datos[11] = df.iloc[59, 17]	# r60 t max 1
        datos[12] = df.iloc[60, 16]	# q61 t normal 2
        datos[13] = df.iloc[60, 17]	# r61 t max 2
        datos[14] = df.iloc[61, 16]	# q62 t normal 3
        datos[15] = df.iloc[61, 17]	# r62 t max 3
        datos[16] = df.iloc[63, 17]	# r64 pendiente
        datos[17] = df.iloc[64, 17]	# r65 fi
    except:
        return 1

    # revisar si hay algun dato incorrecto
    for dato in datos[:-10]:
        if pd.isnull(dato):
            # return manejo_errores(datos, name_ch)
            return [0]
        
    # Formatear datos
    try:
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos[10] = round(float(datos[10]), 1) if isinstance(datos[10], float) else datos[10]
        datos[11] = round(float(datos[11]), 1) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 1) if isinstance(datos[12], float) else datos[12]
        datos[13] = round(float(datos[13]), 1) if isinstance(datos[13], float) else datos[13]
        datos[14] = round(float(datos[14]), 1) if isinstance(datos[14], float) else datos[14]
        datos[15] = round(float(datos[15]), 1) if isinstance(datos[15], float) else datos[15]
        datos[16] = round(float(datos[16]), 2) if isinstance(datos[16], float) else datos[16]
        datos[17] = round(float(datos[17]), 1) if isinstance(datos[17], float) else datos[17]
    except:
        return[1]
    
    # observaciones
    obs = []
    obs_rows = df.iloc[:, 4]

    # buscar donde se encuentra la palabra 'Observaciones'
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break
    index += 1  # aumentar index para no seguir directamente desde obs
    val = 18    # valor para seguir el arreglo de datos

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 4]):
            break
        obs1 = df.iloc[index, 4]
        # revisa si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 5]):
            datos[val] = obs1
        else:
            obs2 = df.loc[index, 5]
            datos[val] = obs1 + ' ' + obs2
        index += 1
        val += 1
    return datos 

def extraer_info_PF(df):
    datos_final = []
    datos = [''] * 7
    try:
        datos[0] = df.iloc[10, 10]  # k11 num muestra
        datos[1] = df.iloc[0, 11]   # L1 Titulo Informe
        datos[2] = df.iloc[2, 12]   # M3 Proyecto
        datos[3] = df.iloc[3, 12]   # M4 N Informe
        datos[4] = df.iloc[4, 12]   # M5 Orden de trabajo
        datos[5] = df.iloc[5, 12]   # M6 Fecha inicio
        datos[6] = df.iloc[6, 12]   # M7 Fecha termino
    except:
        return[1]
    
    datos2 = [''] * 6
    try:
        datos2[0] = df.iloc[13, 11] # L14 Muestra
        datos2[1] = df.iloc[13, 12] # M14 Cota
        datos2[2] = df.iloc[13, 13] # N14 Humedad
        datos2[3] = df.iloc[13, 14] # O14 Absorcion
        datos2[4] = df.iloc[13, 15] # P14 Densidad
        datos2[5] = df.iloc[13, 16] # Q14 Porosidad
    except:
        return[1]
    
    # chequear muestra
    muestra = 14

    if pd.isnull(df.iloc[10, 10]):
        datos[0] = df.iloc[14, 10]
        muestra = 18

    # revisar si hay algun dato incorrecto
    for dato in datos:
        if pd.isnull(dato):
            # return manejo_errores(datos, name_pf)
            return [0]
        
    # Formatear datos
    try:
        datos[0] = re.search(r"\d+", str(datos[0])).group(0)
        datos[5] = datos[5].strftime("%d-%m-%Y") 
        datos[6] = datos[6].strftime("%d-%m-%Y")
        datos2[2] = round(float(datos2[2]), 2) if isinstance(datos2[2], float) else datos2[2]
        datos2[3] = round(float(datos2[3]), 2) if isinstance(datos2[3], float) else datos2[3]
        datos2[4] = round(float(datos2[4]), 2) if isinstance(datos2[4], float) else datos2[4]
        datos2[5] = round(float(datos2[5]), 2) if isinstance(datos2[5], float) else datos2[5]
    except:
        return[1]
    
    # observaciones
    obs = [''] * 4
    obs_rows = df.iloc[:, 11]
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break
    
    index += 1  # aumentar index para no seguir directamente desde 'Observaciones'
    val = 0

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 11]):
            break
        ob1 = df.iloc[index, 11]
        # revisar si la celda adyactente tiene contenido
        if pd.isnull(df.iloc[index, 12]):
            obs[val] = ob1
        else:
            obs2 = df.iloc[index, 12]
            obs[val] = ob1 + ' ' + str(obs2)
        
        # aumentar contadores
        index += 1
        val += 1
    
    # Revisar siguientes columnas y agregar datos
    dato = []
    dato.extend(datos)
    dato.extend(datos2)
    dato.extend(obs)
    datos_final.append(dato)

    index = 14
    # muestra = 14

    # Revisar si hay mas columnas
    while not pd.isnull(df.iloc[index, 11]):
        datos[0] = df.iloc[muestra, 10] # extraer muestras

        datos2 = [''] * 6
        try:
            datos2[0] = df.iloc[index, 11]  # L14 Muestra
            datos2[1] = df.iloc[index, 12]  # M14 Cota
            datos2[2] = df.iloc[index, 13]  # N14 Humedad
            datos2[3] = df.iloc[index, 14]  # O14 Absorcion
            datos2[4] = df.iloc[index, 15]  # P14 Densidad
            datos2[5] = df.iloc[index, 16]  # Q14 Porosidad
        except:
            return[1]
        
        # Formatear datos
        try:
            datos[0] = re.search(r"\d+", str(datos[0])).group(0)
            datos2[2] = round(float(datos2[2]), 2) if isinstance(datos2[2], float) else datos2[2]
            datos2[3] = round(float(datos2[3]), 2) if isinstance(datos2[3], float) else datos2[3]
            datos2[4] = round(float(datos2[4]), 2) if isinstance(datos2[4], float) else datos2[4]
            datos2[5] = round(float(datos2[5]), 2) if isinstance(datos2[5], float) else datos2[5]
        except:
            return[1]
        
        # agregar a la lista
        dato = []
        dato.extend(datos)
        dato.extend(datos2)
        dato.extend(obs)
        datos_final.append(dato)

        # sumar a index
        index += 1
        muestra += 4
    
    return datos_final

# Extrae toda la informacion necesaria para la hoja de hinchamiento libre
def extraer_info_TX(df):
    datos = [''] * 26
    try:
        datos[0] = df.iloc[0, 10]   # K1  Titulo Informe
        datos[1] = df.iloc[2, 11]   # L3  Proyecto
        datos[2] = df.iloc[3, 11]   # L4  N Informe
        datos[3] = df.iloc[4, 11]   # L5  Orden de Trabajo
        datos[4] = df.iloc[5, 11]   # L6  Fecha Inicio
        datos[5] = df.iloc[6, 11]   # L7  Fecha termino
        datos[6] = df.iloc[8, 11]   # L9  Muestra
        datos[7] = df.iloc[9, 11]   # L10 Tipo de Roca
        datos[8] = df.iloc[10, 11]  # L11 Fracturas
        datos[9] = df.iloc[11, 11]  # L12 Alteraciones
        datos[10] = df.iloc[12, 11] # L13 Observaciones
        datos[11] = df.iloc[15, 11] # L17 Diametro
        datos[12] = df.iloc[16, 11] # L18 Altura
        datos[13] = df.iloc[17, 11] # L19 Peso
        datos[14] = df.iloc[18, 11] # L20 Densidad H
        datos[15] = df.iloc[19, 11] # L21 Densidad s
        datos[16] = df.iloc[20, 11] # L22 Humedad
        datos[17] = df.iloc[15, 14] # O17 Tension
        datos[18] = df.iloc[16, 14] # O18 Resistencia max
        datos[19] = df.iloc[17, 14] # O19 Modulo de def 
        datos[20] = df.iloc[18, 14] # O20 Razon de Poisson
        datos[21] = df.iloc[19, 14] # O21 Muestra falla por
    except:
        return[1]

    # revisar si hay datos incorrectos
    for dato in datos:
        if pd.isnull(dato):
            # return manejo_errores(datos, name_tx)
            return [0]
    
    if datos[6] == 0 or datos[6] == "0":
        return [0]
    # Formatear datos
    try:
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos[11] = round(float(datos[11]), 2) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 2) if isinstance(datos[12], float) else datos[12]
        datos[13] = round(float(datos[13]), 2) if isinstance(datos[13], float) else datos[13]
        datos[14] = round(float(datos[14]), 2) if isinstance(datos[14], float) else datos[14]
        datos[15] = round(float(datos[15]), 2) if isinstance(datos[15], float) else datos[15]
        datos[16] = round(float(datos[16]), 2) if isinstance(datos[16], float) else datos[16]
        datos[17] = round(float(datos[17]), 2) if isinstance(datos[17], float) else datos[17]
        datos[18] = round(float(datos[18]), 2) if isinstance(datos[18], float) else datos[18]
        datos[19] = round(float(datos[19]), 2) if isinstance(datos[19], float) else datos[19]
        datos[20] = round(float(datos[20]), 2) if isinstance(datos[20], float) else datos[20]
    except:
        return[1]
    # observaciones
    obs_rows = df.iloc[25:, 10]
    
    # buscar donde se encuentra la palabra 'Observaciones'
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones:':
            break
    index += 1  # aumentar index para no seguir directamente desde 'Observaciones'
    val = 22    # valor para seguir en el arreglo de datos

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 10]):
            break
        ob1 = df.iloc[index, 10]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 11]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 11]
            datos[val] = ob1 + ' ' + str(ob2)

        # aumentar contadores
        index += 1
        val +=1

    return datos

def extraer_info_TI(df):
    datos = [''] * 30
    try:
        datos[0] = df.iloc[0, 14]	# O1 Titulo Informe
        datos[1] = df.iloc[2, 16]	# Q3 Proyecto
        datos[2] = df.iloc[3, 16]	# Q4 N informe
        datos[3] = df.iloc[4, 16]	# Q5 Orden de trabao
        datos[4] = df.iloc[5, 16]   # Q6 fecha inicio
        datos[5] = df.iloc[6, 16]	# Q7 fecha termino
        datos[6] = df.iloc[8, 16]	# Q9 muestra
        datos[7] = df.iloc[9, 16]	# Q10 tipo de roca
        datos[8] = df.iloc[10, 16]	# Q11 fracturas
        datos[9] = df.iloc[11, 16]	# Q12 Alteraciones
        datos[10] = df.iloc[16, 14]	# O17 Diametro 1
        datos[11] = df.iloc[16, 15]	# P17 Altura 1
        datos[12] = df.iloc[16, 16]	# Q17 fza resist 1
        datos[13] = df.iloc[16, 17]	# R17 resistencia 1
        datos[14] = df.iloc[16, 18]	# S17 tipo de falla 1
        datos[15] = df.iloc[17, 14]	# O18 diametro 2
        datos[16] = df.iloc[17, 15]	# P18 altura 2
        datos[17] = df.iloc[17, 16]	# Q18 fza resist 2
        datos[18] = df.iloc[17, 17]	# R18 resist 2
        datos[19] = df.iloc[17, 18]	# S18 tipo de falla 2
        datos[20] = df.iloc[18, 14]	# O19 diametro 3 
        datos[21] = df.iloc[18, 15]	# P19 altura 3 
        datos[22] = df.iloc[18, 16]	# Q19 fza resist 3
        datos[23] = df.iloc[18, 17]	# R19 resist 3
        datos[24] = df.iloc[18, 18]	# s19 tipo de falla 3
        datos[25] = df.iloc[16, 19]	# T17 promedio resist
    except:
        return[1]
    
    # revisar si hay datos incorrectos
    for dato in datos:
        if pd.isnull(dato):
            # return manejo_errores(datos, name_ti)
            return [0]
    
    # Formatear datos
    try:
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos[10] = round(float(datos[10]), 2) if isinstance(datos[10], float) else datos[10]
        datos[11] = round(float(datos[11]), 2) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 2) if isinstance(datos[12], float) else datos[12]
        datos[13] = round(float(datos[13]), 2) if isinstance(datos[13], float) else datos[13]
        datos[15] = round(float(datos[15]), 2) if isinstance(datos[15], float) else datos[15]
        datos[16] = round(float(datos[16]), 2) if isinstance(datos[16], float) else datos[16]
        datos[17] = round(float(datos[17]), 2) if isinstance(datos[17], float) else datos[17]
        datos[18] = round(float(datos[18]), 2) if isinstance(datos[18], float) else datos[18]
        datos[20] = round(float(datos[20]), 2) if isinstance(datos[20], float) else datos[20]
        datos[21] = round(float(datos[21]), 2) if isinstance(datos[21], float) else datos[21]
        datos[22] = round(float(datos[22]), 2) if isinstance(datos[22], float) else datos[22]
        datos[23] = round(float(datos[23]), 2) if isinstance(datos[23], float) else datos[23]
        datos[25] = round(float(datos[25]), 2) if isinstance(datos[25], float) else datos[25]
    except:
        return[1]
    
    # observaciones
    obs_rows = df.iloc[:, 14]

    # buscar donde se encuentra la palabra 'Observaciones'
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break
    
    index += 1
    val = 26

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 14]):
            break
        ob1 = df.iloc[index, 14]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 15]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 15]
            datos[val] = ob1 + ' ' + str(ob2)
        
        # aumentar contadores
        index += 1
        val += 1
    
    return datos

def extraer_info_PLT_diametral(df):
    datos = [''] * 38
    try:
        datos[0] = df.iloc[0, 11]	# l1 Titulo informe
        datos[1] = df.iloc[3, 13]	# n4 Proyecto
        datos[2] = df.iloc[4, 13]	# n5 Num informe
        datos[3] = df.iloc[5, 13]	# n6 Orden de trabajo
        datos[4] = df.iloc[6, 13]	# n7 fecha inicio
        datos[5] = df.iloc[7, 13]	# n8 fecha termino
        datos[6] = df.iloc[10, 13]	# n11 muestra
        datos[7] = df.iloc[11, 13]	# n12 tipo de roca
        datos[8] = df.iloc[12, 13]	# n13 fracturas
        datos[9] = df.iloc[13, 13]	# n14 alteraciones
        datos[10] = df.iloc[20, 11]	# l21 diametro 1
        datos[11] = df.iloc[20, 12]	# m21 largo 1
        datos[12] = df.iloc[20, 13]	# n21 Fza ruptura 1
        datos[13] = df.iloc[20, 14]	# o21 ICP ls 1
        datos[14] = df.iloc[20, 15]	# p21 Factor correccion 1
        datos[15] = df.iloc[20, 16]	# q21 ICP ls50 1
        datos[16] = df.iloc[20, 17]	# r21 Muestra falla por 1
        datos[17] = df.iloc[21, 11]	# l22 Diamtro 2 
        datos[18] = df.iloc[21, 12]	# m22 Largo 2
        datos[19] = df.iloc[21, 13]	# n22 fza ruptura 2
        datos[20] = df.iloc[21, 14]	# o22  ICP ls 2
        datos[21] = df.iloc[21, 15]	# p22 factor correccion 2
        datos[22] = df.iloc[21, 16]	# q22 ICP ls50 2
        datos[23] = df.iloc[21, 17]	# r22 muestra falla por 2
        datos[24] = df.iloc[22, 11]	# l23 Diametro 3
        datos[25] = df.iloc[22, 12]	# m23 Largo 3
        datos[26] = df.iloc[22, 13]	# n23 fza ruptura 3
        datos[27] = df.iloc[22, 14]	# o23 ICP ls 3
        datos[28] = df.iloc[22, 15]	# p23 factor correccion 3
        datos[29] = df.iloc[22, 16]	# q23 ICP ls50 3
        datos[30] = df.iloc[22, 17]	# r23 muestra falla por 3
        datos[31] = df.iloc[25, 15]	# p26 contenido humedad
        datos[32] = df.iloc[26, 15]	# p27 Indice resist diam 50mm
        datos[33] = df.iloc[27, 15]	# p28 clasificacion resist
    except:
        return[1]
    
    # revisar si hay datos incorrectos
    for dato in datos:
        if pd.isnull(dato):
            # return manejo_errores(datos, name_plt_d)
            return [0]
    
    # Formatear datos
    try:
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos[10] = round(float(datos[10]), 2) if isinstance(datos[10], float) else datos[10]
        datos[11] = round(float(datos[11]), 2) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 2) if isinstance(datos[12], float) else datos[12]
        datos[13] = round(float(datos[13]), 2) if isinstance(datos[13], float) else datos[13]
        datos[14] = round(float(datos[14]), 2) if isinstance(datos[14], float) else datos[14]
        datos[15] = round(float(datos[15]), 2) if isinstance(datos[15], float) else datos[15]
        datos[17] = round(float(datos[17]), 2) if isinstance(datos[17], float) else datos[17]
        datos[18] = round(float(datos[18]), 2) if isinstance(datos[18], float) else datos[18]
        datos[19] = round(float(datos[19]), 2) if isinstance(datos[19], float) else datos[19]
        datos[20] = round(float(datos[20]), 2) if isinstance(datos[20], float) else datos[20]
        datos[21] = round(float(datos[21]), 2) if isinstance(datos[21], float) else datos[21]
        datos[22] = round(float(datos[22]), 2) if isinstance(datos[22], float) else datos[22]
        datos[24] = round(float(datos[24]), 2) if isinstance(datos[24], float) else datos[24]
        datos[25] = round(float(datos[25]), 2) if isinstance(datos[25], float) else datos[25]
        datos[26] = round(float(datos[26]), 2) if isinstance(datos[26], float) else datos[26]
        datos[27] = round(float(datos[27]), 2) if isinstance(datos[27], float) else datos[27]
        datos[28] = round(float(datos[28]), 2) if isinstance(datos[28], float) else datos[28]
        datos[29] = round(float(datos[29]), 2) if isinstance(datos[29], float) else datos[29]
        datos[31] = round(float(datos[31]), 2) if isinstance(datos[31], float) else datos[31]
        datos[32] = round(float(datos[32]), 2) if isinstance(datos[32], float) else datos[32]
    except:
        return[1]
    
    # observaciones
    obs_rows = df.iloc[:, 11]

    # buscar donde se encuentra la palabra Observaciones
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break

    index += 1
    val = 34

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 11]):
            break
        ob1 = df.iloc[index, 11]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 12]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 12]
            datos[val] = ob1 + ' ' + str(ob2)

        # aumentar contadores
        index += 1
        val += 1
    
    return datos

def extraer_info_PLT_axial(df):
    try:
        datos = [''] * 44
        datos[0] = df.iloc[0, 12]	# m1 Titulo informe
        datos[1] = df.iloc[3, 13]	# n4 proyecto
        datos[2] = df.iloc[4, 13]	# n5 num muestra
        datos[3] = df.iloc[5, 13]	# n6 orden de trabajo
        datos[4] = df.iloc[6, 13]	# n7 fecha inicio
        datos[5] = df.iloc[7, 13]	# n8 fecha termino
        datos[6] = df.iloc[10, 13]	# n11 muestra
        datos[7] = df.iloc[11, 13]	# n12 tipo de roca
        datos[8] = df.iloc[12, 13]	# n13 fracturas
        datos[9] = df.iloc[13, 13]	# n14 alteraciones
        datos[10] = df.iloc[19, 12]	# m20 dist entre puntas 1
        datos[11] = df.iloc[19, 13]	# n20 ancho 1
        datos[12] = df.iloc[19, 14]	# o20 fuerza de ruptura 1
        datos[13] = df.iloc[19, 15]	# p20 de 1
        datos[14] = df.iloc[19, 16]	# q20 de 1
        datos[15] = df.iloc[19, 17]	# r20 icp ls 1
        datos[16] = df.iloc[19, 18]	# s20 factor de correccion 1
        datos[17] = df.iloc[19, 19]	# t20 icp ls50 1
        datos[18] = df.iloc[19, 20]	# u20 muestra falla por 1
        datos[19] = df.iloc[20, 12]	# m21 dist entre puntas 2
        datos[20] = df.iloc[20, 13]	# n21 ancho 2
        datos[21] = df.iloc[20, 14]	# o21 fuerza de ruptura 2
        datos[22] = df.iloc[20, 15]	# p21 de 2
        datos[23] = df.iloc[20, 16]	# q21 de 2
        datos[24] = df.iloc[20, 17]	# r21 icp ls 2
        datos[25] = df.iloc[20, 18]	# s21 factor de correccion 2
        datos[26] = df.iloc[20, 19]	# t21 icp ls50 2
        datos[27] = df.iloc[20, 20]	# u21 muestra falla por 2
        datos[28] = df.iloc[21, 12]	# m22 dist entre puntas 3
        datos[29] = df.iloc[21, 13]	# b22 ancho 3
        datos[30] = df.iloc[21, 14]	# o22 fuerza de ruptura 3
        datos[31] = df.iloc[21, 15]	# p22 de 3
        datos[32] = df.iloc[21, 16]	# q22 de 3
        datos[33] = df.iloc[21, 17]	# r22 icp ls 3
        datos[34] = df.iloc[21, 18]	# s22 factor de correccion 3
        datos[35] = df.iloc[21, 19]	# t22 icp ls50 3
        datos[36] = df.iloc[21, 20]	# u22 muestra falla por 3
        datos[37] = df.iloc[24, 15]	# p25 contenido de humedad
        datos[38] = df.iloc[25, 15]	# p26 indice de recist diam 500 mm
        datos[39] = df.iloc[26, 15]	# p27 clasificacion de resistencia
    except:
        return[1]
    
    # revisar si hay datos incorrectos
    for dato in datos:
        if pd.isnull(dato):
            return [0]
        
    # formatear datos
    try:
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos[10] = round(float(datos[10]), 1) if isinstance(datos[10], float) else datos[10]
        datos[11] = round(float(datos[11]), 2) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 2) if isinstance(datos[12], float) else datos[12]
        datos[13] = round(float(datos[13]), 2) if isinstance(datos[13], float) else datos[13]
        datos[14] = round(float(datos[14]), 2) if isinstance(datos[14], float) else datos[14]
        datos[15] = round(float(datos[15]), 2) if isinstance(datos[15], float) else datos[15]
        datos[16] = round(float(datos[16]), 2) if isinstance(datos[16], float) else datos[16]
        datos[17] = round(float(datos[17]), 2) if isinstance(datos[17], float) else datos[17]
        datos[19] = round(float(datos[19]), 1) if isinstance(datos[19], float) else datos[19]
        datos[20] = round(float(datos[20]), 2) if isinstance(datos[20], float) else datos[20]
        datos[21] = round(float(datos[21]), 2) if isinstance(datos[21], float) else datos[21]
        datos[22] = round(float(datos[22]), 2) if isinstance(datos[22], float) else datos[22]
        datos[23] = round(float(datos[23]), 2) if isinstance(datos[23], float) else datos[23]
        datos[24] = round(float(datos[24]), 2) if isinstance(datos[24], float) else datos[24]
        datos[25] = round(float(datos[25]), 2) if isinstance(datos[25], float) else datos[25]
        datos[26] = round(float(datos[26]), 2) if isinstance(datos[26], float) else datos[26]
        datos[28] = round(float(datos[28]), 1) if isinstance(datos[28], float) else datos[28]
        datos[29] = round(float(datos[29]), 2) if isinstance(datos[29], float) else datos[29]
        datos[30] = round(float(datos[30]), 2) if isinstance(datos[30], float) else datos[30]
        datos[31] = round(float(datos[31]), 2) if isinstance(datos[31], float) else datos[31]
        datos[32] = round(float(datos[32]), 2) if isinstance(datos[32], float) else datos[32]
        datos[33] = round(float(datos[33]), 2) if isinstance(datos[33], float) else datos[33]
        datos[34] = round(float(datos[34]), 2) if isinstance(datos[34], float) else datos[34]
        datos[35] = round(float(datos[35]), 2) if isinstance(datos[35], float) else datos[35]
        datos[37] = round(float(datos[37]), 2) if isinstance(datos[37], float) else datos[37]
        datos[38] = round(float(datos[38]), 2) if isinstance(datos[38], float) else datos[38]
    except:
        return[1]
    
    # observaciones
    obs_rows = df.iloc[:, 12]

    # buscar donde se encuentra la palabra Observaciones
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break

    index += 1
    val = 40

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 12]):
            break
        ob1 = df.iloc[index, 12]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 13]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 13]
            datos[val] = ob1 + ' ' + str(ob2)

        # aumentar contadores
        index += 1
        val += 1

    return datos

def extraer_info_PLT_colpa(df):
    datos_final = [''] * 144
    try:
        datos_final[0] = df.iloc[0, 20]	# u1 Titulo informe
        datos_final[1] = df.iloc[3, 21]	# v4 Proyecto
        datos_final[2] = df.iloc[4, 21]	# v5 num informe
        datos_final[3] = df.iloc[5, 21]	# v6 Orden de trabajo
        datos_final[4] = df.iloc[6, 21]	# v7 fecha inicio
        datos_final[5] = df.iloc[7, 21]	# v8 fecha termino
        datos_final[6] = df.iloc[8, 21]	# v9 muestra
    except:
        return[1]

    # revisar si hay algun dato incorrecto
    for dato in datos_final:
        if pd.isnull(dato):
            return [0]

    try:
        datos_final[4] = datos_final[4].strftime("%d-%m-%Y") 
        datos_final[5] = datos_final[5].strftime("%d-%m-%Y")
    except:
        return [1]
    
    for index in range(10):
        try:
            # datos2 = [''] * 13
            datos_final[0 + index * 13 + 7] = df.iloc[14 + index, 20]	# u15 identificacion del fragmento a ensayar
            datos_final[1 + index * 13 + 7] = df.iloc[14 + index, 21]	# v15 distancia entre puntas
            datos_final[2 + index * 13 + 7] = df.iloc[14 + index, 22]	# w15 ancho promedio
            datos_final[3 + index * 13 + 7] = df.iloc[14 + index, 23]	# x15 fuerza de ruptura
            datos_final[4 + index * 13 + 7] = df.iloc[14 + index, 24]	# y15 de2
            datos_final[5 + index * 13 + 7] = df.iloc[14 + index, 25]	# z15 de
            datos_final[6 + index * 13 + 7] = df.iloc[14 + index, 26]	# aa15 icp ls
            datos_final[7 + index * 13 + 7] = df.iloc[14 + index, 27]	# ab15 factor de correccion
            datos_final[8 + index * 13 + 7] = df.iloc[14 + index, 28]	# ac15 icp ls50
            datos_final[9 + index * 13 + 7] = df.iloc[14 + index, 29]	# ad15 tipo de roca
            datos_final[10 + index * 13 + 7] = df.iloc[14 + index, 30]	# ae15 fracturas
            datos_final[11 + index * 13 + 7] = df.iloc[14 + index, 31]	# af15 alteracion
            datos_final[12 + index * 13 + 7] = df.iloc[14 + index, 32]	# ag15 muestra falla por
            # datos_final.extend(datos2)
        except:
            return [1]
        
        # Formatear los datos
        try:
            datos_final[1 + index*13 + 7] = round(float(datos_final[1 + index*13 + 7]), 1) if isinstance(datos_final[1 + index*13 + 7], float) else datos_final[1 + index*13 + 7]
            datos_final[2 + index*13 + 7] = round(float(datos_final[2 + index*13 + 7]), 1) if isinstance(datos_final[2 + index*13 + 7], float) else datos_final[2 + index*13 + 7]
            datos_final[3 + index*13 + 7] = round(float(datos_final[3 + index*13 + 7]), 1) if isinstance(datos_final[3 + index*13 + 7], float) else datos_final[3 + index*13 + 7]
            datos_final[4 + index*13 + 7] = round(float(datos_final[4 + index*13 + 7]), 1) if isinstance(datos_final[4 + index*13 + 7], float) else datos_final[4 + index*13 + 7]
            datos_final[5 + index*13 + 7] = round(float(datos_final[5 + index*13 + 7]), 1) if isinstance(datos_final[5 + index*13 + 7], float) else datos_final[5 + index*13 + 7]
            datos_final[6 + index*13 + 7] = round(float(datos_final[6 + index*13 + 7]), 1) if isinstance(datos_final[6 + index*13 + 7], float) else datos_final[6 + index*13 + 7]
            datos_final[7 + index*13 + 7] = round(float(datos_final[7 + index*13 + 7]), 1) if isinstance(datos_final[7 + index*13 + 7], float) else datos_final[7 + index*13 + 7]
            datos_final[8 + index*13 + 7] = round(float(datos_final[8 + index*13 + 7]), 1) if isinstance(datos_final[8 + index*13 + 7], float) else datos_final[8 + index*13 + 7]
        except:
            return [1]
        
    datos_final[137] = df.iloc[49, 23]	# x50 cont humedad
    datos_final[138] = df.iloc[50, 23]	# x51 ls50
    datos_final[139] = df.iloc[51, 23]	# x52 clasificacion resist

    try:
        datos_final[137] = round(float(datos_final[137]), 1) if isinstance(datos_final[137], float) else datos_final[137]
        datos_final[138] = round(float(datos_final[138]), 2) if isinstance(datos_final[138], float) else datos_final[138]
    except:
        return [1]

    # observaciones
    obs_rows = df.iloc[:, 20]
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break
    index += 1
    val = 140

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 20]):
            break
        ob1 = df.iloc[index, 20]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 21]):
            datos_final[val] = ob1
        else:
            obs2 = df.iloc[index, 21]
            datos_final[val] = ob1 + ' ' + str(obs2)

        # aumentar contadores
        index += 1
        val += 1

    return datos_final


def extraer_info_PLT_colpa2(df):
    datos_final = []
    datos = [''] * 7
    try:
        datos[0] = df.iloc[0, 20]	# u1 Titulo informe
        datos[1] = df.iloc[3, 21]	# v4 Proyecto
        datos[2] = df.iloc[4, 21]	# v5 num informe
        datos[3] = df.iloc[5, 21]	# v6 Orden de trabajo
        datos[4] = df.iloc[6, 21]	# v7 fecha inicio
        datos[5] = df.iloc[7, 21]	# v8 fecha termino
        datos[6] = df.iloc[8, 21]	# v9 muestra
    except:
        return[1]
    
    # Datos que se repiten en varias filas
    datos2 = [''] * 13
    try:
        datos2[0] = df.iloc[14, 20]	# u15 identificacion del fragmento a ensayar
        datos2[1] = df.iloc[14, 21]	# v15 distancia entre puntas
        datos2[2] = df.iloc[14, 22]	# w15 ancho promedio
        datos2[3] = df.iloc[14, 23]	# x15 fuerza de ruptura
        datos2[4] = df.iloc[14, 24]	# y15 de2
        datos2[5] = df.iloc[14, 25]	# z15 de
        datos2[6] = df.iloc[14, 26]	# aa15 icp ls
        datos2[7] = df.iloc[14, 27]	# ab15 factor de correccion
        datos2[8] = df.iloc[14, 28]	# ac15 icp ls50
        datos2[9] = df.iloc[14, 29]	# ad15 tipo de roca
        datos2[10] = df.iloc[14, 30]	# ae15 fracturas
        datos2[11] = df.iloc[14, 31]	# af15 alteracion
        datos2[12] = df.iloc[14, 32]	# ag15 muestra falla por
    except:
        return[1]
    
    # revisar si hay algun dato incorrecto
    for dato in datos:
        if pd.isnull(dato):
            # return manejo_errores(datos, name_plt_c)
            return [0]
    
    # Formatear datos
    try:
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos2[1] = round(float(datos2[1]), 1) if isinstance(datos2[1], float) else datos2[1]    
        datos2[2] = round(float(datos2[2]), 1) if isinstance(datos2[2], float) else datos2[2]
        datos2[3] = round(float(datos2[3]), 1) if isinstance(datos2[3], float) else datos2[3]
        datos2[4] = round(float(datos2[4]), 1) if isinstance(datos2[4], float) else datos2[4]
        datos2[5] = round(float(datos2[5]), 1) if isinstance(datos2[5], float) else datos2[5]
        datos2[6] = round(float(datos2[6]), 1) if isinstance(datos2[6], float) else datos2[6]
        datos2[7] = round(float(datos2[7]), 1) if isinstance(datos2[7], float) else datos2[7]
        datos2[8] = round(float(datos2[8]), 1) if isinstance(datos2[8], float) else datos2[8]
    except:
        return[1]
    
    # observaciones
    obs = [''] * 4
    obs_rows = df.iloc[:, 20]
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break

    index += 1
    val = 0

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 20]):
            break
        ob1 = df.iloc[index, 20]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 21]):
            obs[val] = ob1
        else:
            obs2 = df.iloc[index, 21]
            obs[val] = ob1 + ' ' + str(obs2)

        # aumentar contadores
        index += 1
        val += 1

    # Revisar siguientes columnas y agregar datos
    dato = []
    dato.extend(datos)
    dato.extend(datos2)
    dato.extend(obs)
    datos_final.append(dato)

    # Revisar si hay mas columnas
    index = 15
    while not pd.isnull(df.iloc[index, 21]) and not df.iloc[index, 21] == '-':
        datos2 = [''] * 13
        try:
            datos2[0] = df.iloc[index, 20]	# u15 identificacion del fragmento a ensayar
            datos2[1] = df.iloc[index, 21]	# v15 distancia entre puntas
            datos2[2] = df.iloc[index, 22]	# w15 ancho promedio
            datos2[3] = df.iloc[index, 23]	# x15 fuerza de ruptura
            datos2[4] = df.iloc[index, 24]	# y15 de2
            datos2[5] = df.iloc[index, 25]	# z15 de
            datos2[6] = df.iloc[index, 26]	# aa15 icp ls
            datos2[7] = df.iloc[index, 27]	# ab15 factor de correccion
            datos2[8] = df.iloc[index, 28]	# ac15 icp ls50
            datos2[9] = df.iloc[index, 29]	# ad15 tipo de roca
            datos2[10] = df.iloc[index, 30]	# ae15 fracturas
            datos2[11] = df.iloc[index, 31]	# af15 alteracion
            datos2[12] = df.iloc[index, 32]	# ag15 muestra falla por
        except:
            return[1]

        # Formatear datos
        try:
            datos2[1] = round(float(datos2[1]), 1) if isinstance(datos2[1], float) else datos2[1]    
            datos2[2] = round(float(datos2[2]), 1) if isinstance(datos2[2], float) else datos2[2]
            datos2[3] = round(float(datos2[3]), 1) if isinstance(datos2[3], float) else datos2[3]
            datos2[4] = round(float(datos2[4]), 1) if isinstance(datos2[4], float) else datos2[4]
            datos2[5] = round(float(datos2[5]), 1) if isinstance(datos2[5], float) else datos2[5]
            datos2[6] = round(float(datos2[6]), 1) if isinstance(datos2[6], float) else datos2[6]
            datos2[7] = round(float(datos2[7]), 1) if isinstance(datos2[7], float) else datos2[7]
            datos2[8] = round(float(datos2[8]), 1) if isinstance(datos2[8], float) else datos2[8]
        except:
            return[1]
        
        # agregar a la lista
        dato = []
        dato.extend(datos)
        dato.extend(datos2)
        dato.extend(obs)
        datos_final.append(dato)

        # sumar a index 
        index += 1
    
    return datos_final

def extraer_info_slake(df):
    datos = [''] * 32
    try:
        datos[0] = df.iloc[0, 7]	# h1 Titulo Informe
        datos[1] = df.iloc[2, 8]	# i3 Proyecto
        datos[2] = df.iloc[3, 8]	# i4 num informe
        datos[3] = df.iloc[4, 8]	# i5 orden de trabajo
        datos[4] = df.iloc[5, 8]	# i6 fecha inicio
        datos[5] = df.iloc[6, 8]	# i7 fecha termino
        datos[6] = df.iloc[8, 8]	# i9 muestra
        datos[7] = df.iloc[9, 8]	# i10 tipo de roca
        datos[8] = df.iloc[10, 8]	# i11 fracturas
        datos[9] = df.iloc[11, 8]	# i12 alteraciones
        datos[10] = df.iloc[15, 10]	# k16 contenedor
        datos[11] = df.iloc[16, 10]	# k17 contenedor + muestra inicial
        datos[12] = df.iloc[17, 10]	# k18 contenedor 1er ciclo
        datos[13] = df.iloc[18, 10]	# k19 contenedor 2do ciclo
        datos[14] = df.iloc[21, 10]	# k22 tipo de agua utilizada
        datos[15] = df.iloc[22, 10]	# k23 temperatura del agua
        datos[16] = df.iloc[23, 10]	# k24 tiempo de cada ciclo
        datos[17] = df.iloc[24, 10]	# k25 velocidad de giro
        datos[18] = df.iloc[25, 10]	# k26 temperatura de secado
        datos[19] = df.iloc[30, 8]	# i31 desgaste ciclo 0
        datos[20] = df.iloc[30, 10]	# k31 indice de durabilidad ciclo 0
        datos[21] = df.iloc[31, 8]	# i32 desgaste ciclo 1
        datos[22] = df.iloc[31, 10]	# k32 indice de durabilidad ciclo 1
        datos[23] = df.iloc[32, 8]	# i33 desgaste ciclo 2
        datos[24] = df.iloc[32, 10]	# k33 indice de durabilidad ciclo 2
        datos[25] = df.iloc[34, 10]	# k35 contenido humedad inicial
        datos[26] = df.iloc[35, 10]	# k36 slake durability index
        datos[27] = df.iloc[36, 10]	# k37 clasificacion
    except:
        return[1]
    
    # revisar si hay datos incorrectos
    for dato in datos:
        if pd.isnull(dato):
            # return manejo_errores(datos, name_slake)
            return [0]
        
    # Formatear datos
    try:
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos[10] = round(float(datos[10]), 2) if isinstance(datos[10], float) else datos[10]
        datos[11] = round(float(datos[11]), 2) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 2) if isinstance(datos[12], float) else datos[12]
        datos[13] = round(float(datos[13]), 2) if isinstance(datos[13], float) else datos[13]
        datos[19] = round(float(datos[19]), 2) if isinstance(datos[19], float) else datos[19]
        datos[20] = round(float(datos[20]), 2) if isinstance(datos[20], float) else datos[20]
        datos[21] = round(float(datos[21]), 2) if isinstance(datos[21], float) else datos[21]
        datos[22] = round(float(datos[22]), 2) if isinstance(datos[22], float) else datos[22]
        datos[23] = round(float(datos[23]), 2) if isinstance(datos[23], float) else datos[23]
        datos[24] = round(float(datos[24]), 2) if isinstance(datos[24], float) else datos[24]
        datos[25] = round(float(datos[25]), 2) if isinstance(datos[25], float) else datos[25]
        datos[26] = round(float(datos[26]), 2) if isinstance(datos[26], float) else datos[26]
    except:
        return[1]
    
    # observaciones
    obs_rows = df.iloc[:, 7]

    # buscar donde se encuentra la palabra Observaciones
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break

    index += 1
    val = 28

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 7]):
            break
        ob1 = df.iloc[index, 7]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 8]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 8]
            datos[val] = ob1 + ' ' + ob2
        
        # aumentar contadores
        index += 1
        val += 1
    
    return datos

def extraer_info_azul_de_metileno(df):
    datos = [''] * 18
    try:
        datos[0] = df.iloc[0, 7]	# h1 Titulo de informe
        datos[1] = df.iloc[2, 8]	# i3 proyecto
        datos[2] = df.iloc[3, 8]	# i4 num informe
        datos[3] = df.iloc[4, 8]	# i5 orden de trabajo
        datos[4] = df.iloc[5, 8]	# i6 fecha inicial
        datos[5] = df.iloc[6, 8]	# i7 fecha termino
        datos[6] = df.iloc[9, 8]	# i10 muestra
        datos[7] = df.iloc[16, 7]	# h17 peso polvo adm
        datos[8] = df.iloc[16, 8]	# i17 vol agua destilada sol adm
        datos[9] = df.iloc[16, 10]	# k17 peso polvo de roca
        datos[10] = df.iloc[16, 11]	# l17 vol agua destilada polvo de roca
        datos[11] = df.iloc[19, 10]	# k20 total sol adm adicionada
        datos[12] = df.iloc[20, 10]	# k21 total adm adicionado
        datos[13] = df.iloc[21, 10]	# k22 valor adm
    except:
        return[1]
    
    # revisar si hay datos incorrectos
    for dato in datos:
        if pd.isnull(dato):
            # return manejo_errores(datos, name_adm)
            return [0]
        
    # Formatear datos
    try:
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos[7] = round(float(datos[7]), 1) if isinstance(datos[7], float) else datos[7]
        datos[8] = round(float(datos[8]), 1) if isinstance(datos[8], float) else datos[8]
        datos[9] = round(float(datos[9]), 1) if isinstance(datos[9], float) else datos[9]
        datos[10] = round(float(datos[10]), 1) if isinstance(datos[10], float) else datos[10]
        datos[11] = round(float(datos[11]), 2) if isinstance(datos[11], float) else datos[11]
        datos[12] = round(float(datos[12]), 2) if isinstance(datos[12], float) else datos[12]
        datos[13] = round(float(datos[13]), 2) if isinstance(datos[13], float) else datos[13]
    except:
        return[1]
    
    # observaciones
    obs_rows = df.iloc[:, 7]

    # buscar donde se encuentra la palabra Observaciones
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break

    index += 1
    val = 14

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 7]):
            break
        ob1 = df.iloc[index, 7]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 8]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 8]
            datos[val] = ob1 + ' ' + ob2
        
        # aumentar contadores
        index += 1
        val += 1
    
    return datos

def extraer_info_EG(df):
    datos = [''] * 26
    try:
        datos[0] = df.iloc[0, 1]	# b1 Titulo informe
        datos[1] = df.iloc[2, 2]	# c3 proyecto
        datos[2] = df.iloc[3, 2]	# c4 num informe
        datos[3] = df.iloc[4, 2]	# c5 orden de trabajo
        datos[4] = df.iloc[5, 2]	# c6 fecha inicio
        datos[5] = df.iloc[6, 2]	# c7 fecha termino
        datos[6] = df.iloc[9, 2]	# c10 muestra
        datos[7] = df.iloc[10, 2]	# c11 tipo de roca
        datos[8] = df.iloc[11, 2]	# c12 fracturas
        datos[9] = df.iloc[12, 2]	# c13 alteraciones
        datos[10] = df.iloc[57, 2]	# c58 dia 1 promedio
        datos[11] = df.iloc[58, 2]	# c59 dia 1 grado
        datos[12] = df.iloc[57, 6]	# g58 dia 5 promedio
        datos[13] = df.iloc[58, 6]	# g59 dia 5 grado
        datos[14] = df.iloc[57, 10]	# k58 dia 10 promedio
        datos[15] = df.iloc[58, 10]	# k59 dia 10 grado
        datos[16] = df.iloc[57, 14]	# o58 dia 15 promedio
        datos[17] = df.iloc[58, 14]	# o59 dia 15 grado
        datos[18] = df.iloc[57, 18]	# s58 dia 20 promedio
        datos[19] = df.iloc[58, 18]	# s59 dia 20 grado
        datos[20] = df.iloc[57, 22]	# w58 dia 30 promedio
        datos[21] = df.iloc[58, 22]	# w59 dia 30 grado
    except:
        return[1]
    
    # revisar si hay datos incorrectos
    for dato in datos:
        if pd.isnull(dato):
            return[0]
        
    # Formatear datos
    try:
        datos[4] = datos[4].strftime("%d-%m-%Y") 
        datos[5] = datos[5].strftime("%d-%m-%Y")
        datos[10] = round(float(datos[10]), 2) if isinstance(datos[10], float) else datos[10]
        datos[12] = round(float(datos[12]), 2) if isinstance(datos[12], float) else datos[12]
        datos[14] = round(float(datos[14]), 2) if isinstance(datos[14], float) else datos[14]
        datos[16] = round(float(datos[16]), 2) if isinstance(datos[16], float) else datos[16]
        datos[18] = round(float(datos[18]), 2) if isinstance(datos[18], float) else datos[18]
        datos[20] = round(float(datos[20]), 2) if isinstance(datos[20], float) else datos[20]
    except:
        return[1]
    
    # observaciones
    obs_rows = df.iloc[:, 1]

    # buscar donde se encuentra la palabra Observaciones
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break
    index += 1
    val = 22

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 1]):
            break
        ob1 = df.iloc[index, 1]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 2]):
            datos[val] = ob1
        else:
            ob2 = df.iloc[index, 2]
            datos[val] = ob1 + ' ' + ob2

        # aumentar contadores
        index += 1
        val += 1
    
    return datos

def extraer_info_DRX(df):
    datos_base = [''] * 7
    try:
        datos_base[0] = df.iloc[1, 7]	# h2 Titulo informe
        datos_base[1] = df.iloc[3, 8]	# i4 proyecto
        datos_base[2] = df.iloc[4, 8]	# i5 num informe
        datos_base[3] = df.iloc[5, 8]	# i6 orden de trabajo
        datos_base[4] = df.iloc[6, 8]	# i7 muestra
        datos_base[5] = df.iloc[7, 8]	# i8 fecha inicio
        datos_base[6] = df.iloc[8, 8]	# i9 fecha termino
    except:
        return[1]
    
    # revisar si hay datos incorrectos
    for dato in datos_base:
        if pd.isnull(dato):
            return [0]
        
    # Formatear datos
    datos_base[5] = datos_base[5].strftime("%d-%m-%Y") 
    datos_base[6] = datos_base[6].strftime("%d-%m-%Y")
    
    datos = [0] * 37
    suma = 0

    index = 10
    while not pd.isnull(df.iloc[index, 2]):
        if (df.iloc[index, 2].lower()) == "lau":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[0] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "qz":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[1] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "pl":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[2] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "fls":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[3] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "mica":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[4] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "chl":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[5] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "anh":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[6] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "cal":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[7] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "ttn":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[8] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "cpx":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[9] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "anl":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[10] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "ep":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[11] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "hm" or df.iloc[index, 2].lower() == "hem":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[12] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "arc-sm":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[13] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "prh":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[14] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "prs":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[15] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "mag":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[16] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "act":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[17] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "crs":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[18] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "heu":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[19] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "stb":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[20] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "py":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[21] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "cnp" or df.iloc[index, 2].lower() == "clp":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[22] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "wai":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[23] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "ilm":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[24] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "phl":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[25] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "opal":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[26] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "cha":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[27] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "dol":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[28] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "bar":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[29] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "gp":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[30] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "dat":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[31] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "thm":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[32] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "px":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[33] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "tre-act":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[34] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "ame":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[35] = val
            suma += val
        if (df.iloc[index, 2].lower()) == "cln":
            val = 0 if pd.isnull(df.iloc[index, 4]) else df.iloc[index, 4]
            datos[36] = val
            suma += val
    
        index += 1

    datos_obs = [""] * 4

    # observaciones
    obs_rows = df.iloc[:, 7]

    index = 0

    # buscar donde se encuentra la palabra Observaciones
    for index, row in obs_rows.to_frame().iterrows():
        if obs_rows[index] == 'Observaciones':
            break
    
    index += 1
    val = 0

    # buscar observaciones
    while index < len(df.index):
        if pd.isnull(df.iloc[index, 7]):
            break
        ob1 = df.iloc[index, 7]
        # revisar si la celda adyacente tiene contenido
        if pd.isnull(df.iloc[index, 8]):
            datos_obs[val] = ob1
        else:
            ob2 = df.iloc[index, 8]
            datos_obs[val] = ob1 + ' ' + ob2
        
        # aumentar contadores
        index += 1
        val += 1

    datos_suma = [suma]
    datos_final = datos_base
    datos_final.extend(datos)
    datos_final.extend(datos_suma)
    datos_final.extend(datos_obs)
    
    return datos_final

# -------------- Extraer info Word -------------- #
def procesar_claves_con_parentesis(data_dict):
    pattern = re.compile(r"(.*)\s+\((.*?)\)")
    new_data = {}
    for key, value in data_dict.items():
        match = pattern.match(key)
        if match and "Grado de meteorización" not in key:
            new_key = match.group(1).strip()
            new_value = f"{value} ({match.group(2)})"
            new_data[new_key] = new_value
        else:
            new_data[key] = value
    return new_data

def extraer_informacion_especial(doc):
    texto_acumulado_1 = ""
    encontrado_1 = False
    info = {"Alteración": None, "Nombre de la Roca": None}
    
    comp1 = {}
    comp2 = {}
    componentes_primera_categoria = ['Cristaloclastos', 'Vitroclastos', 'Litoclastos', 'Matriz', 'Fenocristales', 'Masa fundamental']
    
    componentes_segunda_categoria = ['Juveniles', 'Líticos', 'Piroxenos (Px)', 'Piroxeno (Px)', 'Plagioclasa (Plg)', 
                                     'Plagioclasas (Plg)', 'Cuarzo', 'Feldespato potásico (Fd-k)', 'Feldespatos potásicos (Fd-k)', 
                                     'Opacos', 'Opaco', 'Vidrios', 'Vidrio']
    
    for paragraph in doc.paragraphs:
        if 'Imágenes microscópicas' in paragraph.text:
            encontrado_1 = True
            break
        texto_acumulado_1 += " " + paragraph.text

        words = paragraph.text.split()
        
        if words:
            if (len(words)==2):
                componente = words[0].rstrip(':')
                porcentaje = words[1].rstrip(':%')
                
            elif (len(words)==3):
                componente = words[0] + ' ' + words[1]
                porcentaje = words[2].rstrip(':%')
                
            elif (len(words)==4):
                componente = words[0] + ' ' + words[1] + ' ' + words[2]
                porcentaje = words[3].rstrip(':%')
            else:
                componente = None
            
        if componente in componentes_segunda_categoria:            
                
            componente = 'a_' + componente
            if componente in comp1:
                componente = 'b_' + componente
            comp1[componente] = porcentaje
        
        if (componente in componentes_primera_categoria) and (componente not in comp1):
            comp1[componente] = porcentaje

        if len(words)>0 and ("Ceniza" in words[0]):
            categoria = ' '.join(words[:-1]).strip(":")
            info[categoria] = words[-1]
        elif len(words)>0 and ("Lapilli" in words[0]):
            categoria = ' '.join(words[:-1]).strip(":")
            info[categoria] = words[-1]
        elif len(words)>0 and ("Bloques" in words[0]):
            categoria = ' '.join(words[:-1]).strip(":")
            info[categoria] = words[-1]      
            
    if not encontrado_1:
        return "El texto imágenes microscópicas no fue encontrado"
        
    alteracion_matches = re.findall(r"Alteración:\s*(.*?)\.", texto_acumulado_1)    
    roca_matches = re.findall(r"(Nombre de la roca|Clasificación):\s*(.*?)\.", texto_acumulado_1)
        
    #patron_categoria = re.findall("\b(Ceniza|Lapilli|Bloques|Bombas)\b(?:\s*\(<\d+\s*mm\))?:\s*\d+%", texto_acumulado_2)
    patron_porcentaje = re.compile(r'\d+\s*%')
    #categorias = patron_categoria.findall(texto_acumulado_2)

    if alteracion_matches:
        info["Alteración"] = alteracion_matches[-1].strip()
    if roca_matches:
        info["Nombre de la Roca"] = roca_matches[-1][-1]

    info.update(comp1)
        
    return info

def extraer_informacion_archivo(ruta_archivo):
    doc = Document(ruta_archivo)
    archivo = ruta_archivo.split('\\')[-1]
    titulo_principal = doc.paragraphs[0].text.strip()
    informacion = {
        'Nombre del Archivo': archivo,
        'Ensayo': titulo_principal}
    
    data_dict = {}
    for tabla in doc.tables:
        for row in tabla.rows:
            campo = row.cells[0].text.strip()
            campo = re.sub(r'\s+', ' ', campo).strip()
            if campo == "Info. Ensaye":
                campo = "N° Informe"
            valor = row.cells[1].text.strip()
            data_dict[campo] = valor
            
    data_dict = procesar_claves_con_parentesis(data_dict)    
    informacion.update(data_dict)    
   
    info_especial = extraer_informacion_especial(doc)
    informacion.update(info_especial)

    return informacion
# -------------- Manejo archivos -------------- #
# Agrega una linea de texto al final del archivo excel
def to_excel(dir, line, sheet):
    try:
        wb = load_workbook(dir)
        ws = wb.worksheets[sheet]
    except FileNotFoundError:
        print("No se ha encontrado el archivo")
    ws.append(line)
    wb.save(dir)

# Cambia el valor de una casilla
def update_excel(dir, r, c, val):
    wb = load_workbook(filename=dir)
    ws = wb.worksheets[0]
    ws.cell(row=r, column=c).value = val
    wb.save(dir)

# Crea o agrega a un txt las hojas con errores
def errores(dir, error_logs, informe):
    dir_errores = dir.split("/")
    dir_errores = "/".join(dir_errores[:-1])
    # crear/abrir txt
    file_name = dir_errores + "/" + informe + ".txt"
    f = open(file_name, "a+")

    for line in error_logs:
        f.write(line + "\n")
    
    f.close()

# crea una lista de archivos por directorio y subdirectorios segun extension
def list_files_by_extension(directory):
    excel_files = []
    # Guarda todos los xlsx
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith('xlsx'):
                excel_files.append(os.path.join(root, file))
    # guarda todos los xlsm
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith('xlsm'):
                excel_files.append(os.path.join(root, file))
    # Agrega los archivos terminados en xls
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith('xls'):
                excel_files.append(os.path.join(root, file))
    return excel_files

# crea una lista de archivos .docx por directorio y subdirectorios segun extension
def list_docx_by_extension(directory):
    word_files = []
    # Guarda todos los xlsx
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith('.docx'):
                word_files.append(os.path.join(root, file))
    return word_files

# -------------- Manejo errores -------------- #
# revisa que errores son los que faltan
def manejo_errores(datos, informe):
    lista_pendientes = [0]
    if informe == name_ucs:
        if pd.isnull(datos[0]):
            lista_pendientes.append(f"\t- Falta Titulo del Informe")
        if pd.isnull(datos[1]):
            lista_pendientes.append(f"\t- Falta Proyecto")
        if pd.isnull(datos[2]):
            lista_pendientes.append(f"\t- Falta N° de Informe")
        if pd.isnull(datos[3]):
            lista_pendientes.append(f"\t- Falta Orden de trabajo")
        if pd.isnull(datos[4]):
            lista_pendientes.append(f"\t- Falta Fecha de Inicio")
        if pd.isnull(datos[5]):
            lista_pendientes.append(f"\t- Falta Fecha de Término")
        if pd.isnull(datos[6]):
            lista_pendientes.append(f"\t- Falta Muestra")
        if pd.isnull(datos[7]):
            lista_pendientes.append(f"\t- Falta Tipo de Roca")
        if pd.isnull(datos[8]):
            lista_pendientes.append(f"\t- Falta Fracturas")
        if pd.isnull(datos[9]):
            lista_pendientes.append(f"\t- Falta Alteraciones")
        if pd.isnull(datos[10]):
            lista_pendientes.append(f"\t- Falta Observaciones 1")
        if pd.isnull(datos[11]):
            lista_pendientes.append(f"\t- Falta Diametro Promedio")
        if pd.isnull(datos[12]):
            lista_pendientes.append(f"\t- Falta Altura Promedio")
        if pd.isnull(datos[13]):
            lista_pendientes.append(f"\t- Falta Densidad Humeda")
        if pd.isnull(datos[14]):
            lista_pendientes.append(f"\t- Falta Densidad seca")
        if pd.isnull(datos[15]):
            lista_pendientes.append(f"\t- Falta Contenido de humedad")
        if pd.isnull(datos[16]):
            lista_pendientes.append(f"\t- Falta Resistenia Max D50")
        if pd.isnull(datos[17]):
            lista_pendientes.append(f"\t- Falta Resistencia max D61")
    elif informe == name_hl:
        if pd.isnull(datos[0]):
            lista_pendientes.append(f"\t- Falta Titulo del Informe")
        if pd.isnull(datos[1]):
            lista_pendientes.append(f"\t- Falta Proyecto")
        if pd.isnull(datos[2]):
            lista_pendientes.append(f"\t- Falta N° de Informe")
        if pd.isnull(datos[3]):
            lista_pendientes.append(f"\t- Falta Orden de trabajo")
        if pd.isnull(datos[4]):
            lista_pendientes.append(f"\t- Falta Fecha de Inicio")
        if pd.isnull(datos[5]):
            lista_pendientes.append(f"\t- Falta Fecha de Término")
        if pd.isnull(datos[6]):
            lista_pendientes.append(f"\t- Falta Muestra")
        if pd.isnull(datos[7]):
            lista_pendientes.append(f"\t- Falta Volumen inicial muestra seca")
        if pd.isnull(datos[8]):
            lista_pendientes.append(f"\t- Falta Volumen final muestra sedimentidad")
        if pd.isnull(datos[9]):
            lista_pendientes.append(f"\t- Falta Indice Hinchiemto libre")
    elif informe == name_ch:
        if pd.isnull(datos[0]):
            lista_pendientes.append(f"\t- Falta Titulo del Informe")
        if pd.isnull(datos[1]):
            lista_pendientes.append(f"\t- Falta Proyecto")
        if pd.isnull(datos[2]):
            lista_pendientes.append(f"\t- Falta N° de Informe")
        if pd.isnull(datos[3]):
            lista_pendientes.append(f"\t- Falta Orden de trabajo")
        if pd.isnull(datos[4]):
            lista_pendientes.append(f"\t- Falta Fecha de Inicio")
        if pd.isnull(datos[5]):
            lista_pendientes.append(f"\t- Falta Fecha de Término")
        if pd.isnull(datos[6]):
            lista_pendientes.append(f"\t- Falta Muestra")
        if pd.isnull(datos[7]):
            lista_pendientes.append(f"\t- Falta Tipo de Roca")
        if pd.isnull(datos[8]):
            lista_pendientes.append(f"\t- Falta Fracturas")
        if pd.isnull(datos[9]):
            lista_pendientes.append(f"\t- Falta Alteraciones")
        if pd.isnull(datos[10]):
            lista_pendientes.append(f"\t- Falta Area muestra menor")
        if pd.isnull(datos[11]):
            lista_pendientes.append(f"\t- Falta Tensión Normal")
        if pd.isnull(datos[12]):
            lista_pendientes.append(f"\t- Falta Fuerza de carga normal")
        if pd.isnull(datos[13]):
            lista_pendientes.append(f"\t- Falta Diámetro pistón de carga")
        if pd.isnull(datos[14]):
            lista_pendientes.append(f"\t- Falta Área pistón de carga")
        if pd.isnull(datos[15]):
            lista_pendientes.append(f"\t- Falta Presión pistón de carga")
    elif informe== 'PF':
        lista_pendientes.append(f"Error")
    elif informe == 'TX':
        if pd.isnull(datos[0]):
            lista_pendientes.append(f"\t- Falta Titulo del Informe")
        if pd.isnull(datos[1]):
            lista_pendientes.append(f"\t- Falta Proyecto")
        if pd.isnull(datos[2]):
            lista_pendientes.append(f"\t- Falta N° de Informe")
        if pd.isnull(datos[3]):
            lista_pendientes.append(f"\t- Falta Orden de trabajo")
        if pd.isnull(datos[4]):
            lista_pendientes.append(f"\t- Falta Fecha de Inicio")
        if pd.isnull(datos[5]):
            lista_pendientes.append(f"\t- Falta Fecha de Término")
        if pd.isnull(datos[6]):
            lista_pendientes.append(f"\t- Falta Muestra")
        if pd.isnull(datos[7]):
            lista_pendientes.append(f"\t- Falta Tipo de Roca")
        if pd.isnull(datos[8]):
            lista_pendientes.append(f"\t- Falta Fracturas")
        if pd.isnull(datos[9]):
            lista_pendientes.append(f"\t- Falta Alteraciones")
        if pd.isnull(datos[10]):
            lista_pendientes.append(f"\t- Falta Observaciones 1")
        if pd.isnull(datos[11]):
            lista_pendientes.append(f"\t- Falta Diametro Probeta")
        if pd.isnull(datos[12]):
            lista_pendientes.append(f"\t- Falta Altura Probeta")
        if pd.isnull(datos[13]):
            lista_pendientes.append(f"\t- Falta Peso")
        if pd.isnull(datos[14]):
            lista_pendientes.append(f"\t- Falta Densidad Humeda")
        if pd.isnull(datos[15]):
            lista_pendientes.append(f"\t- Falta Densidad Seca")
        if pd.isnull(datos[16]):
            lista_pendientes.append(f"\t- Falta Humedad")
        if pd.isnull(datos[17]):
            lista_pendientes.append(f"\t- Falta Tensión de Confinamiento")
        if pd.isnull(datos[18]):
            lista_pendientes.append(f"\t- Falta Resistencia Max. D50")
        if pd.isnull(datos[19]):
            lista_pendientes.append(f"\t- Falta Módulo de Deformación")
        if pd.isnull(datos[20]):
            lista_pendientes.append(f"\t- Falta Razón de Poisson")
        if pd.isnull(datos[21]):
            lista_pendientes.append(f"\t- Falta Muestra Falla por")
        #if pd.isnull(datos[22]):
        #    lista_pendientes.append(f"\t- Falta Velocidad Onda de Compresión")
        #if pd.isnull(datos[23]):
        #    lista_pendientes.append(f"\t- Falta Valocidad Onda de Cizalle")
    elif informe == "TI":
        if pd.isnull(datos[0]):
            lista_pendientes.append(f"\t- Falta Titulo del Informe")
        if pd.isnull(datos[1]):
            lista_pendientes.append(f"\t- Falta Proyecto")
        if pd.isnull(datos[2]):
            lista_pendientes.append(f"\t- Falta N° de Informe")
        if pd.isnull(datos[3]):
            lista_pendientes.append(f"\t- Falta Orden de trabajo")
        if pd.isnull(datos[4]):
            lista_pendientes.append(f"\t- Falta Fecha de Inicio")
        if pd.isnull(datos[5]):
            lista_pendientes.append(f"\t- Falta Fecha de Término")
        if pd.isnull(datos[6]):
            lista_pendientes.append(f"\t- Falta Muestra")
        if pd.isnull(datos[7]):
            lista_pendientes.append(f"\t- Falta Tipo de Roca")
        if pd.isnull(datos[8]):
            lista_pendientes.append(f"\t- Falta Fracturas")
        if pd.isnull(datos[9]):
            lista_pendientes.append(f"\t- Falta Alteraciones")
        if pd.isnull(datos[10]):
            lista_pendientes.append(f"\t- Falta Diámetro 1")
        if pd.isnull(datos[11]):
            lista_pendientes.append(f"\t- Falta Altura 1")
        if pd.isnull(datos[12]):
            lista_pendientes.append(f"\t- Falta Fuerza de Ruptura 1")
        if pd.isnull(datos[13]):
            lista_pendientes.append(f"\t- Falta Resistensia a la tensión 1")
        if pd.isnull(datos[14]):
            lista_pendientes.append(f"\t- Falta Tipo de Falla 1")
        if pd.isnull(datos[15]):
            lista_pendientes.append(f"\t- Falta Diámetro 2")
        if pd.isnull(datos[16]):
            lista_pendientes.append(f"\t- Falta Altura 2")
        if pd.isnull(datos[17]):
            lista_pendientes.append(f"\t- Falta Fuerza de Ruptura 2")
        if pd.isnull(datos[18]):
            lista_pendientes.append(f"\t- Falta Resistencia a la tensión 2")
        if pd.isnull(datos[19]):
            lista_pendientes.append(f"\t- Falta Tipo de Falla 2")
        if pd.isnull(datos[20]):
            lista_pendientes.append(f"\t- Falta Diámetro 3")
        if pd.isnull(datos[21]):
            lista_pendientes.append(f"\t- Falta Altura 3")
        if pd.isnull(datos[22]):
            lista_pendientes.append(f"\t- Falta Fuerza de Ruptura 3")
        if pd.isnull(datos[23]):
            lista_pendientes.append(f"\t- Falta Resistencia a la tensión 3")
        if pd.isnull(datos[24]):
            lista_pendientes.append(f"\t- Falta Tipo de Falla 3")
        if pd.isnull(datos[25]):
            lista_pendientes.append(f"\t- Falta Promedio Resistencia a la tensión")
    elif informe=="PLT diametral":
        if pd.isnull(datos[0]):
            lista_pendientes.append(f"\t- Falta Titulo del Informe")
        if pd.isnull(datos[1]):
            lista_pendientes.append(f"\t- Falta Proyecto")
        if pd.isnull(datos[2]):
            lista_pendientes.append(f"\t- Falta N° de Informe")
        if pd.isnull(datos[3]):
            lista_pendientes.append(f"\t- Falta Orden de trabajo")
        if pd.isnull(datos[4]):
            lista_pendientes.append(f"\t- Falta Fecha de Inicio")
        if pd.isnull(datos[5]):
            lista_pendientes.append(f"\t- Falta Fecha de Término")
        if pd.isnull(datos[6]):
            lista_pendientes.append(f"\t- Falta Muestra")
        if pd.isnull(datos[7]):
            lista_pendientes.append(f"\t- Falta Tipo de Roca")
        if pd.isnull(datos[8]):
            lista_pendientes.append(f"\t- Falta Fracturas")
        if pd.isnull(datos[9]):
            lista_pendientes.append(f"\t- Falta Alteraciones")
        if pd.isnull(datos[10]):
            lista_pendientes.append(f"\t- Falta Diámetro 1")
        if pd.isnull(datos[11]):
            lista_pendientes.append(f"\t- Falta Largo 1")
        if pd.isnull(datos[12]):
            lista_pendientes.append(f"\t- Falta Fuerza de Ruptura 1")
        if pd.isnull(datos[13]):
            lista_pendientes.append(f"\t- Falta Indice de Carga Puntual ls 1")
        if pd.isnull(datos[14]):
            lista_pendientes.append(f"\t- Falta Factor de Corrección 1")
        if pd.isnull(datos[15]):
            lista_pendientes.append(f"\t- Falta Índice de Carga Puntual ls50 1")
        if pd.isnull(datos[16]):
            lista_pendientes.append(f"\t- Falta Muestra Falla 1")
        if pd.isnull(datos[17]):
            lista_pendientes.append(f"\t- Falta Diámetro 2")
        if pd.isnull(datos[18]):
            lista_pendientes.append(f"\t- Falta Largo 2")
        if pd.isnull(datos[19]):
            lista_pendientes.append(f"\t- Falta Fuerza de Ruptura 2")
        if pd.isnull(datos[20]):
            lista_pendientes.append(f"\t- Falta Indice de Carga Puntual ls 2")
        if pd.isnull(datos[21]):
            lista_pendientes.append(f"\t- Falta Factor de Corrección 2")
        if pd.isnull(datos[22]):
            lista_pendientes.append(f"\t- Falta Indice de Carga Puntual ls50 2")
        if pd.isnull(datos[23]):
            lista_pendientes.append(f"\t- Falta Muestra falla 2")
        if pd.isnull(datos[24]):
            lista_pendientes.append(f"\t- Falta Diametro 3")
        if pd.isnull(datos[25]):
            lista_pendientes.append(f"\t- Falta Largo 3")
        if pd.isnull(datos[26]):
            lista_pendientes.append(f"\t- Falta Fuerza de Ruptura 3")
        if pd.isnull(datos[27]): 
            lista_pendientes.append(f"\t- Falta Indice de Carga Puntual ls 3")
        if pd.isnull(datos[28]):
            lista_pendientes.append(f"\t- Falta Factor de Corrección 3")
        if pd.isnull(datos[29]):
            lista_pendientes.append(f"\t- Falta Indice de Carga puntual ls50 3")
        if pd.isnull(datos[30]):
            lista_pendientes.append(f"\t- Falta Muestra Falla 3")
        if pd.isnull(datos[31]):
            lista_pendientes.append(f"\t- Falta Contenido de Humedad")
        if pd.isnull(datos[32]):
            lista_pendientes.append(f"\t- Falta Indice de resistencia para diametro 50mm")
        if pd.isnull(datos[33]):
            lista_pendientes.append(f"\t- Falta Clasificacion de resistencia")

    return lista_pendientes

# -------------- Formateo -------------- #
# formate a y agrega las filas correspondientes al archivo final
def format_excel(dir_inicial, dir_final):
    # Chequear si las entradas son validas
    if dir_inicial == "":
        message_label.config(text= "Por favor seleccione una carpeta")
    elif dir_final == "":
        message_label.config(text="Por favor seleccione algun archivo")
    elif not dir_final.endswith(".xlsx"):
        message_label.config(text="Por favor seleccione un archivo válido")
    else: 
        # mostrar mensaje en el programa
        message_label.config(text=f'\t\t\t\tActualizando\t\t\t\t')
        root.update_idletasks()
        # abrir archivo errores
        dir_errores = dir_final.split("/")
        dir_errores = "/".join(dir_errores[:-1])
        dir_errores_final = dir_errores + "/" + error_name + ".txt"
        f = open(dir_errores_final, "a+")

        # crear lista con los archivos
        try:
            archivos = list_files_by_extension(dir_inicial)
        except:
            error = ["Algo inesperado ocurrio tratando de abrir la carpeta"]
            for line in error:
                f.write(line + "\n")

        dir = dir_final.split("\\")
        dir_informes = "\\".join(dir[:-1])
        
        # crear listado archivos existentes
        archivos_existentes = pd.ExcelFile(dir_final)

        # extraer info de archivos existentes
        ucs_existente = pd.read_excel(archivos_existentes, name_ucs)
        ucs_m_existente = pd.read_excel(archivos_existentes, name_ucs_m)
        plt_d_existente = pd.read_excel(archivos_existentes, name_plt_d)
        plt_a_existente = pd.read_excel(archivos_existentes, name_plt_a)
        plt_c_existente = pd.read_excel(archivos_existentes, name_plt_c)
        hl_existente = pd.read_excel(archivos_existentes, name_hl)
        ch_existente = pd.read_excel(archivos_existentes, name_ch)
        pf_existente = pd.read_excel(archivos_existentes, name_pf)
        ti_existente = pd.read_excel(archivos_existentes, name_ti)
        tx_existente = pd.read_excel(archivos_existentes, name_tx)
        slake_existente = pd.read_excel(archivos_existentes, name_slake)
        adm_existente = pd.read_excel(archivos_existentes, name_adm)
        eg_existente = pd.read_excel(archivos_existentes, name_eg)
        drx_existente = pd.read_excel(archivos_existentes, name_drx)
        res_existente = pd.read_excel(archivos_existentes, name_res, skiprows=1)
        doc_existente = pd.read_excel(archivos_existentes, name_mic)

        # cambiar fila
        ucs_existente["Nombre de la Hoja"] = ucs_existente["Nombre de la Hoja"].astype(str)
        ucs_m_existente["Nombre de la Hoja"] = ucs_m_existente["Nombre de la Hoja"].astype(str)
        plt_d_existente["Nombre de la Hoja"] = plt_d_existente["Nombre de la Hoja"].astype(str)
        plt_a_existente["Nombre de la Hoja"] = plt_a_existente["Nombre de la Hoja"].astype(str)
        plt_c_existente["Nombre de la Hoja"] = plt_c_existente["Nombre de la Hoja"].astype(str)
        hl_existente["Nombre de la Hoja"] = hl_existente["Nombre de la Hoja"].astype(str)
        ch_existente["Nombre de la Hoja"] = ch_existente["Nombre de la Hoja"].astype(str)
        pf_existente["Nombre de la Hoja"] = pf_existente["Nombre de la Hoja"].astype(str)
        ti_existente["Nombre de la Hoja"] = ti_existente["Nombre de la Hoja"].astype(str)
        tx_existente["Nombre de la Hoja"] = tx_existente["Nombre de la Hoja"].astype(str)
        slake_existente["Nombre de la Hoja"] = slake_existente["Nombre de la Hoja"].astype(str)
        adm_existente["Nombre de la Hoja"] = adm_existente["Nombre de la Hoja"].astype(str)
        eg_existente["Nombre de la Hoja"] = eg_existente["Nombre de la Hoja"].astype(str)
        drx_existente["Nombre de la Hoja"] = drx_existente["Nombre de la Hoja"].astype(str)
        res_existente["N° Muestra"] = res_existente["N° Muestra"].astype(str)
        res_existente["Ingreso"] = res_existente["Ingreso"].astype(str)

        count = 0

        for file in archivos:
            if '~$' in file:
                continue
            
            # Abrir excel y sacar nombres de archivos
            nombre_archivo = file.split("\\")[-1]
            try:
                num_ingreso = re.search(r"\d\d\d\d", nombre_archivo).group(0)
            except:
                num_ingreso = '-'
            excel = pd.ExcelFile(file)
            sheets = excel.sheet_names

            check_nombre = nombre_archivo.lower()

            for sheet in sheets:
                count += 1
                print(f'Revisando informe: "{nombre_archivo}"\tHoja: "{sheet}"')
                # la primera hoja es de configuracion y no contiene informacion relevante
                if sheet.lower() == "ingreso datos":
                    continue

                if sheet == "Hoja1":
                    continue
                
                # El numero de muestra se encuentra generalmente en la hoja
                try:
                    num_muestra = sheet.split("-")[-1]
                    num_muestra = re.search(r"\d+", str(num_muestra)).group(0)
                except:
                    num_muestra = sheet
                
                time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                # crear la primera parte de la data final
                final_data = [time, nombre_archivo, sheet, num_ingreso, num_muestra]

                # extraer el resto de informacion de la pagina
                page = pd.read_excel(excel, sheet, header=None)
                
                # Chequear si es un informe de UCS-M
                if "ucs-m" in check_nombre:
                    if nombre_archivo in ucs_m_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = ucs_m_existente[ucs_m_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sheet in df_nombre_archivo[["Nombre de la Hoja"]].values or sh in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    # extraer info ucs
                    info = extraer_info_UCS_M(page)
                    if info[0] == 0:
                        error = [f"Datos faltantes - UCS-M\t{nombre_archivo}\t{sheet}"]
                        # errores(dir_informes, error, error_name)
                        for line in error:
                            f.write(line + "\n")

                    elif info[0] == 1:
                        error = f"Ocurrio algo inesperado tratando de abrir {nombre_archivo} - {sheet}"
                        f.write(error + "\n")
                    else:
                        final_data.extend(info)
                        to_excel(dir_final, final_data, pag_ucs_m)

                        # chequear si existen los archivos
                        ingreso = str(int(final_data[3])) if final_data[3].isnumeric() else str(final_data[3])
                        muestra = str(int(final_data[4])) if final_data[4].isnumeric() else str(final_data[4])
                        nom_muestra = str(final_data[11]).lstrip().rstrip()
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & (res_existente["N° Muestra"]==muestra) & (res_existente["Nombre muestra"]==nom_muestra)]
                        if len(res.index)==0:
                            # agregar fila nueva
                            line = [''] * 127
                            line[0] = str(final_data[3])
                            line[1] = str(final_data[4])
                            line[2] = str(final_data[3]) + '-' + str(final_data[4])
                            line[3] = final_data[11]
                            line[4] = final_data[6]
                            cotas = final_data[11].split("(")[1] if "(" in str(final_data[11]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(')', '')
                            except:
                                line[6] = ''
                            # datos
                            line[44] = final_data[21] # d50
                            line[45] = final_data[22] # e
                            line[46] = final_data[23] # v
                            line[47] = final_data[24] # tipo falla
                            to_excel(dir_final, line, pag_res)

                        else:
                            # updatear datos
                            idx = res.index[0] + 3

                            wb = load_workbook(dir_final)
                            ws = wb.worksheets[pag_res]
                            ws.cell(row=idx, column=45).value = final_data[21]
                            ws.cell(row=idx, column=46).value = final_data[22]
                            ws.cell(row=idx, column=47).value = final_data[23]
                            ws.cell(row=idx, column=48).value = final_data[24]
                            wb.save(dir_final)

                # Chequear si es un informe de UCS
                elif "ucs" in check_nombre or "compresion uniaxial simple" in check_nombre:
                    if nombre_archivo in ucs_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = ucs_existente[ucs_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sheet in df_nombre_archivo[["Nombre de la Hoja"]].values or sh in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    # extraer info ucs
                    info = extraer_info_UCS(page)
                    if info[0] == 0:
                        error = [f"Datos faltantes - UCS\t{nombre_archivo}\t{sheet}"]
                        # errores(dir_informes, error, error_name)
                        for line in error:
                            f.write(line + "\n")

                    elif info[0] == 1:
                        error = f"Ocurrio algo inesperado tratando de abrir {nombre_archivo} - {sheet}"
                        f.write(error + "\n")

                    else:
                        final_data.extend(info)
                        to_excel(dir_final, final_data, pag_ucs)

                        # chequear si existen los archivos
                        ingreso = str(int(final_data[3])) if final_data[3].isnumeric() else str(final_data[3])
                        muestra = str(int(final_data[4])) if final_data[4].isnumeric() else str(final_data[4])
                        nom_muestra = str(final_data[11]).lstrip().rstrip()
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & (res_existente["N° Muestra"]==muestra) & (res_existente["Nombre muestra"]==nom_muestra)]
                        if len(res.index)==0:
                            # agregar fila nueva
                            line = [''] * 127
                            line[0] = str(final_data[3])
                            line[1] = str(final_data[4])
                            line[2] = str(final_data[3]) + '-' + str(final_data[4])
                            line[3] = final_data[11]
                            line[4] = final_data[6]
                            cotas = final_data[11].split("(")[1] if "(" in str(final_data[11]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(')', '')
                            except:
                                line[6] = ''
                            # datos
                            line[42] = final_data[21]
                            line[43] = final_data[-3]
                            to_excel(dir_final, line, pag_res)

                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            d50 = 43
                            fal = 44

                            update_excel(dir_final, idx, d50, final_data[21])
                            update_excel(dir_final, idx, fal, final_data[-3])

                # Chequear si es un informe de HL
                elif "h.l" in check_nombre or "hinchamiento libre" in check_nombre or "hl" in check_nombre:
                    if nombre_archivo in hl_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = hl_existente[hl_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheet in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    info = extraer_info_HL(page)
                    if info[0]==0:
                        error = [f"Datos Faltantes - HL\t{nombre_archivo}\t{sheet}"]
                        # errores(dir_informes, error, error_name)
                        for line in error:
                            f.write(line + "\n")

                    elif info[0] == 1:
                        error = f"Ocurrio algo inesperado tratando de abrir {nombre_archivo} - {sheet}"
                        f.write(error + "\n")

                    else:
                        final_data.extend(info)
                        to_excel(dir_final, final_data, pag_hl)

                        # chequear si existen los archivos
                        ingreso = str(int(final_data[3])) if final_data[3].isnumeric() else str(final_data[3])
                        muestra = str(int(final_data[4])) if final_data[4].isnumeric() else str(final_data[4])
                        nom_muestra = str(final_data[11]).lstrip().rstrip()
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & (res_existente["N° Muestra"]==muestra) & (res_existente["Nombre muestra"]==nom_muestra)]
                        if len(res.index)==0:
                            # agregar fila nueva
                            line = [''] * 127
                            line[0] = str(final_data[3])
                            line[1] = str(final_data[4])
                            line[2] = str(final_data[3]) + '-' + str(final_data[4])
                            line[3] = final_data[11]
                            line[4] = final_data[6]
                            cotas = final_data[11].split("(")[1] if "(" in str(final_data[11]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(')', '')
                            except:
                                line[6] = ''
                            # datos
                            line[35] = final_data[14]
                            to_excel(dir_final, line, pag_res)

                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            hl = 36

                            update_excel(dir_final, idx, hl, final_data[14])
                            
                
                # chequear si es un informe de ch
                elif "ch " in check_nombre or "corte hoek" in check_nombre:
                    if nombre_archivo in ch_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = ch_existente[ch_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheet in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    info = extraer_info_CH(page)
                    if info[0] == 0:
                        error = [f"Datos Faltantes - CH\t{nombre_archivo}\t{sheet}"]
                        # errores(dir_informes, error, "Informe")
                        for line in error:
                            f.write(line + "\n")

                    elif info[0] == 1:
                        error = f"Ocurrio algo inesperado tratando de abrir {nombre_archivo} - {sheet}"
                        f.write(error + "\n")

                    else:
                        final_data.extend(info)
                        to_excel(dir_final, final_data, pag_ch)

                        # chequear si existen los archivos
                        ingreso = str(int(final_data[3])) if final_data[3].isnumeric() else str(final_data[3])
                        muestra = str(int(final_data[4])) if final_data[4].isnumeric() else str(final_data[4])
                        nom_muestra = str(final_data[11]).lstrip().rstrip()
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & (res_existente["N° Muestra"]==muestra) & (res_existente["Nombre muestra"]==nom_muestra)]
                        if len(res.index)==0:
                            # agregar fila nueva
                            line = [''] * 127
                            line[0] = str(final_data[3])
                            line[1] = str(final_data[4])
                            line[2] = str(final_data[3]) + '-' + str(final_data[4])
                            line[3] = final_data[11]
                            line[4] = final_data[6]
                            cotas = final_data[11].split("(")[1] if "(" in str(final_data[11]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(')', '')
                            except:
                                line[6] = ''
                            # datos
                            for index in range(8):
                                line[66 + index] = final_data[15 + index]
                            to_excel(dir_final, line, pag_res)

                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            wb = load_workbook(dir_final)
                            ws = wb.worksheets[pag_res]
                            for index in range(8):
                                ws.cell(row=idx, column=index+67).value = final_data[15 + index]
                            wb.save(dir_final)

                # chequear si es un informe de pf
                elif "pf" in check_nombre or "propiedades físicas" in check_nombre:
                    if nombre_archivo in pf_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = pf_existente[pf_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheet in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    info = extraer_info_PF(page)
                    if info[0] == 0:
                        error = [f"Datos Faltantes - PF\t{nombre_archivo}\t{sheet}"]
                        # errores(dir_informes, error, error_name)
                        for line in error:
                            f.write(line + "\n")

                    elif info[0] == 1:
                        error = f"Ocurrio algo inesperado tratando de abrir {nombre_archivo} - {sheet}"
                        f.write(error + "\n")

                    else:
                        for row in info:
                            data = []
                            data.extend(final_data)
                            data.extend(row[1:])
                            data[4] = row[0]
                            to_excel(dir_final, data, pag_pf)
                            
                            # chequear si existen los archivos
                            ingreso = str(int(data[3])) if data[3].isnumeric() else str(data[3])
                            muestra = str(int(data[4])) if data[4].isnumeric() else str(data[4])
                            nombre_muestra = str(data[11]) + ' (' + str(data[12]) + ')'
                            res = res_existente[(res_existente["Ingreso"]==ingreso) & (res_existente["N° Muestra"]==muestra) & (res_existente["Nombre muestra"]==nombre_muestra)]
                            if len(res.index)==0:
                                # agregar fila nueva
                                line = [''] * 127
                                line[0] = str(data[3])
                                line[1] = str(data[4])
                                line[2] = str(data[3]) + '-' + str(data[4])
                                line[3] = data[11] + ' ' + data[12]
                                line[4] = data[6] 
                                cotas = data[12].split("-") 
                                line[5] = cotas[0].replace("(", "")
                                line[6] = cotas[1].replace(")", "")
                                line[5] = str(float(line[5])) if isinstance(line[5], float) else line[5]
                                line[6] = str(float(line[6])) if isinstance(line[6], float) else line[6]
                                line[29] = data[13] # humedad
                                line[30] = data[14] # absorcion
                                line[31] = data[15] # densidad
                                line[32] = data[16] # porosidad
                                to_excel(dir_final, line, pag_res)

                            else:
                                # updatear datos
                                idx = res.index[0] + 3
                                hum = 30
                                abs = 31
                                den = 32
                                por = 33
                                wb = load_workbook(dir_final)
                                ws = wb.worksheets[pag_res]
                                ws.cell(row=idx, column=hum).value = data[13]
                                ws.cell(row=idx, column=abs).value = data[14]
                                ws.cell(row=idx, column=den).value = data[15]
                                ws.cell(row=idx, column=por).value = data[16]
                                wb.save(dir_final)

                # chequear si es un informe plt
                # check*
                elif "plt" in check_nombre or "carga puntual" in check_nombre:
                    if "diametral" in check_nombre:
                        if nombre_archivo in plt_d_existente["Nombre del Archivo"].values:
                            df_nombre_archivo = plt_d_existente[plt_d_existente["Nombre del Archivo"]==nombre_archivo]
                            try:
                                sh = str(int(sheet))
                            except:
                                sh = sheet
                            if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheet in df_nombre_archivo[["Nombre de la Hoja"]].values:
                                continue
                        info = extraer_info_PLT_diametral(page)
                        if info[0] == 0:
                            error = [f"Datos Faltantes - PLT Diametral\t{nombre_archivo}\t{sheet}"]
                            # errores(dir_informes, error, error_name)
                            for line in error:
                                f.write(line + "\n")

                        elif info[0] == 1:
                            error = f"Ocurrio algo inesperado tratando de abrir {nombre_archivo} - {sheet}"
                            f.write(error + "\n")

                        else:
                            final_data.extend(info)
                            to_excel(dir_final, final_data, pag_plt_d)

                            # chequear si existen los archivos
                            ingreso = str(int(final_data[3])) if final_data[3].isnumeric() else str(final_data[3])
                            muestra = str(int(final_data[4])) if final_data[4].isnumeric() else str(final_data[4])
                            nom_muestra = str(final_data[11]).lstrip().rstrip()
                            res = res_existente[(res_existente["Ingreso"]==ingreso) & (res_existente["N° Muestra"]==muestra) & (res_existente["Nombre muestra"]==nom_muestra)]
                            if len(res.index)==0:
                                # agregar fila nueva
                                line = [''] * 127
                                line[0] = str(final_data[3])
                                line[1] = str(final_data[4])
                                line[2] = str(final_data[3]) + '-' + str(final_data[4])
                                line[3] = final_data[11]
                                line[4] = final_data[6]
                                cotas = final_data[11].split("(")[1] if "(" in str(final_data[11]) else "-"
                                cotas = cotas.split("-")
                                line[5] = cotas[0]
                                try:
                                    line[6] = cotas[1].replace(')', '')
                                except:
                                    line[6] = ''
                                # datos
                                line[41] = final_data[37] # resist a la tension
                                to_excel(dir_final, line, pag_res)

                            else:
                                # updatear datos
                                idx = res.index[0] + 3
                                rt = 42

                                update_excel(dir_final, idx, rt, final_data[37])
                    
                    elif "axial" in check_nombre:
                        if nombre_archivo in plt_a_existente["Nombre del Archivo"].values:
                            df_nombre_archivo = plt_a_existente[plt_a_existente["Nombre del Archivo"]==nombre_archivo]
                            try:
                                sh = str(int(sheet))
                            except:
                                sh = sheet
                            if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheet in df_nombre_archivo[["Nombre de la Hoja"]].values:
                                continue
                        info = extraer_info_PLT_axial(page)
                        
                        if info[0] == 0:
                            error = [f"Datos Faltantes - PLT Axial\t{nombre_archivo}\t{sheet}"]
                            # errores(dir_informes, error, error_name)
                            for line in error:
                                f.write(line + "\n")
                        
                        elif info[0] == 1:
                            error = f"Ocurrio algo inesperado tratando de abrir {nombre_archivo} - {sheet}"
                            f.write(error + "\n")

                        else:
                            final_data.extend(info)
                            to_excel(dir_final, final_data, pag_plt_a)
                           
                            # chequear si existen los archivos
                            ingreso = str(int(final_data[3])) if final_data[3].isnumeric() else str(final_data[3])
                            muestra = str(int(final_data[4])) if final_data[4].isnumeric() else str(final_data[4])
                            nom_muestra = str(final_data[11]).lstrip().rstrip()
                            res = res_existente[(res_existente["Ingreso"]==ingreso) & (res_existente["N° Muestra"]==muestra) & (res_existente["Nombre muestra"]==nom_muestra)]
                            if len(res.index)==0:
                                # agregar fila nueva
                                line = [''] * 127
                                line[0] = str(final_data[3])
                                line[1] = str(final_data[4])
                                line[2] = str(final_data[3]) + '-' + str(final_data[4])
                                line[3] = final_data[11]
                                line[4] = final_data[6]
                                cotas = final_data[11].split("(")[1] if "(" in final_data[11] else "-"
                                cotas = cotas.split("-")
                                line[5] = cotas[0]
                                try:
                                    line[6] = cotas[1].replace(')', '')
                                except:
                                    line[6] = ''
                                # datos
                                line[39] = final_data[43] # resist a la tension
                                to_excel(dir_final, line, pag_res)

                            else:
                                # updatear datos
                                idx = res.index[0] + 3
                                rt = 40

                                update_excel(dir_final, idx, rt, final_data[43])
                    
                    elif "colpa" in check_nombre:
                        if nombre_archivo in plt_c_existente["Nombre del Archivo"].values:
                            df_nombre_archivo = plt_c_existente[plt_c_existente["Nombre del Archivo"]==nombre_archivo]
                            try:
                                sh = str(int(sheet))
                            except:
                                sh = sheet
                            if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheet in df_nombre_archivo[["Nombre de la Hoja"]].values:
                                continue
                        info = extraer_info_PLT_colpa(page)

                        if info[0] == 0:
                            error = [f"Datos Faltantes - PLT Colpa\t{nombre_archivo}\t{sheet}"]
                            # errores(dir_informes, errores, error_name)
                            for line in error:
                                f.write(line + "\n")

                        elif info[0] == 1:
                            error = f"Ocurrio algo inesperado tratando de abrir {nombre_archivo} - {sheet}"
                            f.write(error + "\n")

                        else:
                            final_data.extend(info)
                            to_excel(dir_final, final_data, pag_plt_c)

                            # chequear si existen los archivos
                            ingreso = str(int(final_data[3])) if final_data[3].isnumeric() else str(final_data[3])
                            muestra = str(int(final_data[4])) if final_data[4].isnumeric() else str(final_data[4])
                            nom_muestra = str(final_data[11]).lstrip().rstrip()
                            res = res_existente[(res_existente["Ingreso"]==ingreso) & (res_existente["N° Muestra"]==muestra) & (res_existente["Nombre muestra"]==nom_muestra)]
                            if len(res.index)==0:
                                # agregar fila nueva
                                line = [''] * 127
                                line[0] = str(final_data[3])
                                line[1] = str(final_data[4])
                                line[2] = str(final_data[3]) + '-' + str(final_data[4])
                                line[3] = final_data[11]
                                line[4] = final_data[6] 
                                cotas = final_data[11].split("(")[1] if "(" in str(final_data[11]) else "-"
                                cotas = cotas.split("-")
                                line[5] = cotas[0]
                                try:
                                    line[6] = cotas[1].replace(')', '')
                                except:
                                    line[6] = ''

                                # datos
                                line[40] = final_data[143]
                                to_excel(dir_final, line, pag_res)

                            else:
                                # updatear datos
                                idx = res.index[0] + 3
                                rt = 41

                                update_excel(dir_final, idx, rt, final_data[143])
                    else:
                        error = [f"Informe no identificado en los tipos\t{nombre_archivo}\t{sheet}"]
                        # errores(dir_final, error, error_name)
                        for line in error:
                            f.write(line + "\n")

                # chequear si es un informe de ti
                elif "ti " in check_nombre or "traccion indirecta" in check_nombre:
                    if nombre_archivo in ti_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = ti_existente[ti_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheet in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    info = extraer_info_TI(page)
                    if info[0] == 0:
                        error = [f"Datos Faltantes - TI\t{nombre_archivo}\t{sheet}"]
                        # errores(dir_informes, error, error_name)
                        for line in error:
                            f.write(line + "\n")

                    elif info[0] == 1:
                        error = f"Ocurrio algo inesperado tratando de abrir {nombre_archivo} - {sheet}"
                        f.write(error + "\n")

                    else:
                        final_data.extend(info)
                        to_excel(dir_final, final_data, pag_ti)

                        # chequear si existen los archivos
                        ingreso = str(int(final_data[3])) if final_data[3].isnumeric() else str(final_data[3])
                        muestra = str(int(final_data[4])) if final_data[4].isnumeric() else str(final_data[4])
                        nom_muestra = str(final_data[11]).lstrip().rstrip()
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & (res_existente["N° Muestra"]==muestra) & (res_existente["Nombre muestra"]==nom_muestra)]
                        if len(res.index)==0:
                            # agregar fila nueva
                            line = [''] * 127
                            line[0] = str(final_data[3])
                            line[1] = str(final_data[4])
                            line[2] = str(final_data[3]) + '-' + str(final_data[4])
                            line[3] = final_data[11]
                            line[4] = final_data[6]
                            cotas = final_data[11].split("(")[1] if "(" in str(final_data[11]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(')', '')
                            except:
                                line[6] = ''
                            # datos
                            line[65] = final_data[30] # resist a la tension
                            to_excel(dir_final, line, pag_res)

                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            rt = 66

                            update_excel(dir_final, idx, rt, final_data[30])
                
                # chequear si es un informe de tx
                elif "tx" in check_nombre or "tx-m" in check_nombre or "triaxial" in check_nombre:
                    if nombre_archivo in tx_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = tx_existente[tx_existente["Nombre del Archivo"]==nombre_archivo]
                        # df_nombre_archivo["Nombre de la Hoja"] = df_nombre_archivo["Nombre de la Hoja"].astype(str)
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheet in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    info = extraer_info_TX(page)
                    if info[0] == 0:
                        error = [f"Datos Faltantes - TX\t{nombre_archivo}\t{sheet}"]
                        # errores(dir_informes, error, error_name)
                        for line in error:
                            f.write(line + "\n")

                    elif info[0] == 1:
                        error = f"Ocurrio algo inesperado tratando de abrir {nombre_archivo} - {sheet}"
                        f.write(error + "\n")

                    else:
                        final_data.extend(info)
                        to_excel(dir_final, final_data, pag_tx)
                        
                        # chequear si existen los archivos
                        ingreso = str(int(final_data[3])) if final_data[3].isnumeric() else str(final_data[3])
                        muestra = str(int(final_data[4]))if final_data[4].isnumeric() else str(final_data[4])
                        nom_muestra = str(final_data[11]).lstrip().rstrip()
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & (res_existente["N° Muestra"]==muestra) & (res_existente["Nombre muestra"]==nom_muestra)]
                        if len(res.index)==0:
                            # agregar fila nueva
                            line = [''] * 127
                            line[0] = str(final_data[3])
                            line[1] = str(final_data[4])
                            line[2] = str(final_data[3]) + '-' + str(final_data[4])
                            line[3] = final_data[11]
                            line[4] = final_data[6]
                            cotas = final_data[11].split("(")[1] if "(" in str(final_data[11]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(')', '')
                            except:
                                line[6] = ''
                            # datos
                            if sheet[-1] == "A":
                                line[50] = final_data[22] # mpa
                                line[51] = final_data[23] # max d50
                                line[52] = final_data[24] # e mpa
                                line[53] = final_data[25] # v
                                line[54] = final_data[26] # tipo de falla
                                to_excel(dir_final, line, pag_res)
                            elif sheet[-1] == "B":
                                line[55] = final_data[22]
                                line[56] = final_data[23]
                                line[57] = final_data[24]
                                line[58] = final_data[25]
                                line[59] = final_data[26]
                                to_excel(dir_final, line, pag_res)
                            elif sheet[-1] == "C":
                                line[60] = final_data[22]
                                line[61] = final_data[23]
                                line[62] = final_data[24]
                                line[63] = final_data[25]
                                line[64] = final_data[26]
                                to_excel(dir_final, line, pag_res)

                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            wb = load_workbook(dir_final)
                            ws = wb.worksheets[pag_res]
                            if sheet[-1] == "A":
                                ws.cell(row=idx, column=51).value = final_data[22]
                                ws.cell(row=idx, column=52).value = final_data[23]
                                ws.cell(row=idx, column=53).value = final_data[24]
                                ws.cell(row=idx, column=54).value = final_data[25]
                                ws.cell(row=idx, column=55).value = final_data[26]
                            elif sheet[-1] == "B":
                                ws.cell(row=idx, column=56).value = final_data[22]
                                ws.cell(row=idx, column=57).value = final_data[23]
                                ws.cell(row=idx, column=58).value = final_data[24]
                                ws.cell(row=idx, column=59).value = final_data[25]
                                ws.cell(row=idx, column=60).value = final_data[26]
                            elif sheet[-1] == "C":
                                ws.cell(row=idx, column=61).value = final_data[22]
                                ws.cell(row=idx, column=62).value = final_data[23]
                                ws.cell(row=idx, column=63).value = final_data[24]
                                ws.cell(row=idx, column=64).value = final_data[25]
                                ws.cell(row=idx, column=65).value = final_data[26]
                                
                            wb.save(dir_final)
                
                # chequear si es un informe slake
                elif "slake" in check_nombre:
                    if nombre_archivo in slake_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = slake_existente[slake_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheets in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    info = extraer_info_slake(page)
                    if info[0] == 0:
                        error = [f"Datos Faltantes - Slake\t{nombre_archivo}\t{sheet}"]
                        # errores(dir_informes, error, error_name)
                        for line in error:
                            f.write(line + "\n")

                    elif info[0] == 1:
                        error = f"Ocurrio algo inesperado tratando de abrir {nombre_archivo} - {sheet}"
                        f.write(error + "\n")

                    else:
                        final_data.extend(info)
                        to_excel(dir_final, final_data, pag_slake)

                        # chequear si existen los archivos
                        ingreso = str(int(final_data[3])) if final_data[3].isnumeric() else str(final_data[3])
                        muestra = str(int(final_data[4])) if final_data[4].isnumeric() else str(final_data[4])
                        nom_muestra = str(final_data[11]).lstrip().rstrip()
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & (res_existente["N° Muestra"]==muestra) & (res_existente["Nombre muestra"]==nom_muestra)]
                        if len(res.index)==0:
                            # agregar fila nueva
                            line = [''] * 127
                            line[0] = str(final_data[3])
                            line[1] = str(final_data[4])
                            line[2] = str(final_data[3]) + '-' + str(final_data[4])
                            line[3] = final_data[11]
                            line[4] = final_data[6]
                            cotas = final_data[11].split("(")[1] if "(" in str(final_data[11]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(')', '')
                            except:
                                line[6] = ''
                            # datos
                            line[74] = final_data[27] # ciclo 1
                            line[75] = final_data[29] # ciclo 1
                            to_excel(dir_final, line, pag_res)

                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            c1 = 75
                            c2 = 76

                            update_excel(dir_final, idx, c1, final_data[27])
                            update_excel(dir_final, idx, c2, final_data[29])

                # chequear si es un informe de azul de metileno
                elif "azul de metileno" in check_nombre:
                    if nombre_archivo in adm_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = adm_existente[adm_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheets in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    info = extraer_info_azul_de_metileno(page)
                    if info[0] == 0:
                        error = [f"Datos Faltantes - Azul de Metileno\t{nombre_archivo}\t{sheet}"]
                        # errores(dir_informes, error, error_name)
                        for line in error:
                            f.write(line + "\n")

                    elif info[0] == 1:
                        error = f"Ocurrio algo inesperado tratando de abrir {nombre_archivo} - {sheet}"
                        f.write(error + "\n")

                    else:
                        final_data.extend(info)
                        to_excel(dir_final, final_data, pag_adm)

                        # chequear si existen los archivos
                        ingreso = str(int(final_data[3])) if final_data[3].isnumeric() else str(final_data[3])
                        muestra = str(int(final_data[4])) if final_data[4].isnumeric() else str(final_data[4])
                        nom_muestra = str(final_data[11]).lstrip().rstrip()
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & (res_existente["N° Muestra"]==muestra) & (res_existente["Nombre muestra"]==nom_muestra)]
                        if len(res.index)==0:
                            # agregar fila nueva
                            line = [''] * 127
                            line[0] = str(final_data[3])
                            line[1] = str(final_data[4])
                            line[2] = str(final_data[3]) + '-' + str(final_data[4])
                            line[3] = final_data[11]
                            line[4] = final_data[6]
                            cotas = final_data[11].split("(")[1] if "(" in str(final_data[11]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(')', '')
                            except:
                                line[6] = ''
                            
                            # datos
                            line[33] = final_data[16] # adicionado
                            line[34] = final_data[18] # MBV
                            to_excel(dir_final, line, pag_res)

                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            adi = 34
                            mbv = 35
                            update_excel(dir_final, idx, adi, final_data[16])
                            update_excel(dir_final, idx, mbv, final_data[18])
                
                # chequear si es un informe de etilenglicol
                elif "etilenglicol" in check_nombre or "eg" in check_nombre:
                    if nombre_archivo in eg_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = eg_existente[eg_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheets in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    info = extraer_info_EG(page)
                    if info[0] == 0:
                        error = [f"Datos Faltantes - Etilenglicol\t{nombre_archivo}\t{sheet}"]
                        for line in error:
                            f.write(line + "\n")

                    elif info[0] == 1:
                        error = f"Ocurrio algo inesperado tratando de abrir {nombre_archivo} - {sheet}"
                        f.write(error + "\n")

                    else:
                        final_data.extend(info)
                        to_excel(dir_final, final_data, pag_eg)

                        # chequear si existen los archivos
                        ingreso = str(int(final_data[3])) if final_data[3].isnumeric() else str(final_data[3])
                        muestra = str(int(final_data[4])) if final_data[4].isnumeric() else str(final_data[4])
                        nom_muestra = str(final_data[11]).lstrip().rstrip()
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & (res_existente["N° Muestra"]==muestra) & (res_existente["Nombre muestra"]==nom_muestra)]
                        if len(res.index)==0:
                            # agregar fila nueva
                            line = [''] * 127
                            line[0] = str(final_data[3])
                            line[1] = str(final_data[4])
                            line[2] = str(final_data[3]) + '-' + str(final_data[4])
                            line[3] = final_data[11]
                            line[4] = final_data[6]
                            cotas = final_data[11].split("(")[1] if "(" in final_data[11] else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(')', '')
                            except:
                                line[6] = ''
                            # datos
                            for index in range(12):
                                line[76 + index] = final_data[15 + index]
                            to_excel(dir_final, line, pag_res)

                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            wb = load_workbook(dir_final)
                            ws = wb.worksheets[pag_res]
                            for index in range(12):
                                ws.cell(row=idx, column=index+77).value = final_data[15 + index]
                            wb.save(dir_final)

                # chequear si es un informe de drx
                elif "drx" in check_nombre or "difracciones rayos x" in check_nombre:
                    if sheet.lower() == "datos iniciales":
                        continue
                    if nombre_archivo in drx_existente["Nombre del Archivo"].values:
                        df_nombre_archivo = drx_existente[drx_existente["Nombre del Archivo"]==nombre_archivo]
                        try:
                            sh = str(int(sheet))
                        except:
                            sh = sheet
                        if sh in df_nombre_archivo[["Nombre de la Hoja"]].values or sheets in df_nombre_archivo[["Nombre de la Hoja"]].values:
                            continue
                    info = extraer_info_DRX(page)
                    if info[0] == 0:
                        error = [f"Datos Faltantes - DRX\t{nombre_archivo}\t{sheet}"]
                        for line in error:
                            f.write(line + "\n")
                    
                    elif info[0] == 1:
                        error = f"Ocurrio algo inesperado tratando de abrir {nombre_archivo} - {sheet}"
                        f.write(error + "\n")

                    else:
                        final_data.extend(info)
                        to_excel(dir_final, final_data, pag_drx)

                        # chequear si existen los archivos
                        ingreso = str(int(final_data[3])) if final_data[3].isnumeric() else str(final_data[3])
                        muestra = str(int(final_data[4])) if final_data[4].isnumeric() else str(final_data[4])
                        nom_muestra = str(final_data[9]).lstrip().rstrip()
                        res = res_existente[(res_existente["Ingreso"]==ingreso) & (res_existente["N° Muestra"]==muestra) & (res_existente["Nombre muestra"]==nom_muestra)]
                        if len(res.index)==0:
                            # agregar fila nueva
                            line = [''] * 127
                            line[0] = str(final_data[3])
                            line[1] = str(final_data[4])
                            line[2] = str(final_data[3]) + '-' + str(final_data[4])
                            line[3] = final_data[9]
                            line[4] = final_data[6]
                            cotas = final_data[9].split("(")[1] if "(" in str(final_data[9]) else "-"
                            cotas = cotas.split("-")
                            line[5] = cotas[0]
                            try:
                                line[6] = cotas[1].replace(')', '')
                            except:
                                line[6] = ''
                            # datos
                            for index in range(37):
                                line[90 + index] = final_data[12 + index]
                            to_excel(dir_final, line, pag_res)

                        else:
                            # updatear datos
                            idx = res.index[0] + 3
                            wb = load_workbook(dir_final)
                            ws = wb.worksheets[pag_res]
                            for index in range(37):
                                ws.cell(row=idx, column=index+91).value = final_data[12 + index]
                            wb.save(dir_final)
                
                # buscar informes que no calzan con las descripciones
                else:
                    error = [f"Informe no identificado en los tipos\t{nombre_archivo}\t{sheet}"]
                    # errores(dir_final, error, error_name)
                    for line in error:
                            f.write(line + "\n")

    print(f"Se han analizado {count} hojas")
    # Formatear archivos docx
    # crear lista con los archivos
    try:
        archivos = list_docx_by_extension(dir_inicial)
    except:
        error = ["Algo inesperado ocurrio tratando de abrir la carpeta"]
        for line in error:
            f.write(line + "\n")

    dir = dir_final.split("\\")
    # nombre_archivo = dir[-1]

    # crear listado de archivos existentes
    archivos_existentes = pd.ExcelFile(dir_final)

    # Extraer info de archivos existentes
    mic_existente = pd.read_excel(archivos_existentes, name_mic)

    # cambiar fila
    # mic_existente["Nombre de la Hoja"] = mic_existente["Nombre de la Hoja"].astype(str)

    for file in archivos:
        count += 1
        if file[:2] == '~$':
            continue
        

        nombre_archivo = file.split("\\")[-1]

        if nombre_archivo in doc_existente["Nombre del Archivo"].values:
            continue

        print(f"Revisando Informe {nombre_archivo}")

        info_archivos = [""] * 32
        line = extraer_informacion_archivo(file)
        for key in line.keys():
            if key == "Nombre del Archivo":
                info_archivos[0] = line[key]
            if key == "Ensayo":
                info_archivos[1] = line[key]
            if key == "Nombre Proyecto":
                info_archivos[2] = line[key]
            if key == "N° Informe":
                info_archivos[3] = line[key]
            if key == "Fecha Inicio":
                info_archivos[4] = line[key]
            if key == "Fecha Término":
                info_archivos[5] = line[key]
            if key == "Muestra":
                info_archivos[6] = line[key]
            if key == "1. Textura":
                info_archivos[7] = line[key]
            if key == "2. Grado de cristalinidad":
                info_archivos[8] = line[key]
            if key == "3. Tamaño relativo de cristales":
                info_archivos[9] = line[key]
            if key == "4. Tamaño absoluto de los cristales":
                info_archivos[10] = line[key]
            if key == "5. Tamaño del grano":
                info_archivos[11] = line[key]
            if key == "6. Forma de los cristales":
                info_archivos[12] = line[key]
            if key == "7. Estructura":
                info_archivos[13] = line[key]
            if key == "8. Fábrica":
                info_archivos[14] = line[key]
            if key == "9. Morfología especial":
                info_archivos[15] = line[key]
            if key == "10. Índice de color":
                info_archivos[16] = line[key]
            if key == "11. Grado de meteorización (ISRM, 1981)":
                info_archivos[17] = line[key]
            if key == "Alteración":
                info_archivos[18] = line[key]
            if key == "Nombre de la Roca":
                info_archivos[19] = line[key]
            if key == "Ceniza (<2 mm)":
                info_archivos[20] = line[key]
            if key == "Lapilli (2 - 64 mm)":
                info_archivos[21] = line[key]
            if key == "Bloques y Bombas (>64 mm)":
                info_archivos[22] = line[key]
            if key == "Cristaloclastos":
                info_archivos[23] = line[key]
            if key == "a_Plagioclasas (Plg)":
                info_archivos[24] = line[key]
            if key == "a_Piroxenos (Px)":
                info_archivos[25] = line[key]
            if key == "Litoclastos":
                info_archivos[26] = line[key]
            if key == "Vitroclastos":
                info_archivos[27] = line[key]
            if key == "Matriz":
                info_archivos[28] = line[key]
            if key == "Fenocristales":
                info_archivos[29] = line[key]
            if key == "Masa fundamental":
                info_archivos[30] = line[key]

        # print(keys)
        to_excel(dir_final, info_archivos, pag_mic)

        # chequear si existen los archivos
        ingreso = re.search(r"\d\d\d\d", nombre_archivo).group(0).rstrip().lstrip()
        muestra = re.search(r"-\d([A-Za-z]|\d)*", nombre_archivo).group(0).replace("-", "")
        muestra = str(int(muestra)) if isinstance(muestra, int) else muestra
        muestra = muestra[1:] if muestra[0] == '0' else muestra
        muestra = muestra.lstrip().rstrip()
        res = res_existente[(res_existente["Ingreso"]==ingreso) & (res_existente["N° Muestra"]==muestra) & (res_existente["Nombre muestra"]==info_archivos[6])]
        if len(res.index)==0:
            # agregar fila nueva
            line = [''] * 127
            line[0] = ingreso
            line[1] = muestra
            line[2] = ingreso + '-' + muestra
            line[3] = info_archivos[6]
            line[4] = info_archivos[2]
            cotas = info_archivos[6].split("(")[1] if "(" in info_archivos[6] else "-"
            cotas = cotas.split("-")
            line[5] = cotas[0]
            try:
                line[6] = cotas[1].replace(')', '')
            except:
                line[6] = ''

            final_data = info_archivos
            # datos
            line[7] = final_data[19]
            line[8] = final_data[18]
            line[9] = final_data[7]
            line[10] = final_data[8]
            line[11] = final_data[9]
            line[12] = final_data[10]
            line[13] = final_data[11]
            line[14] = final_data[12]
            line[15] = final_data[13]
            line[16] = final_data[14]
            line[17] = final_data[15]
            line[18] = final_data[16]
            line[19] = final_data[17]
            line[20] = final_data[20]
            line[21] = final_data[21]
            line[22] = final_data[22]
            line[23] = final_data[23]
            line[24] = final_data[26]
            line[25] = final_data[27]
            line[26] = final_data[28]
            line[27] = final_data[29]
            line[28] = final_data[30]
            to_excel(dir_final, line, pag_res)
        #elif len(res.index) > 1:
        #    print(f"3 - {res.index}")
        
        else:
            # updatear datos
            final_data = info_archivos
            idx = res.index[0] + 3
            wb = load_workbook(dir_final)
            ws = wb.worksheets[pag_res]
            ws.cell(row=idx, column=8).value = final_data[19]
            ws.cell(row=idx, column=9).value = final_data[18]
            ws.cell(row=idx, column=10).value = final_data[7]
            ws.cell(row=idx, column=11).value = final_data[8]
            ws.cell(row=idx, column=12).value = final_data[9]
            ws.cell(row=idx, column=13).value = final_data[10]
            ws.cell(row=idx, column=14).value = final_data[11]
            ws.cell(row=idx, column=15).value = final_data[12]
            ws.cell(row=idx, column=16).value = final_data[13]
            ws.cell(row=idx, column=17).value = final_data[14]
            ws.cell(row=idx, column=18).value = final_data[15]
            ws.cell(row=idx, column=19).value = final_data[16]
            ws.cell(row=idx, column=20).value = final_data[17]
            ws.cell(row=idx, column=21).value = final_data[20]
            ws.cell(row=idx, column=22).value = final_data[21]
            ws.cell(row=idx, column=23).value = final_data[22]
            ws.cell(row=idx, column=24).value = final_data[23]
            ws.cell(row=idx, column=25).value = final_data[26]
            ws.cell(row=idx, column=26).value = final_data[27]
            ws.cell(row=idx, column=27).value = final_data[28]
            ws.cell(row=idx, column=28).value = final_data[29]
            ws.cell(row=idx, column=29).value = final_data[30]
            wb.save(dir_final)


    error = [f'---------------Fecha {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}---------------']
    for line in error:
        f.write(line + "\n")
    message_label.config(text="Se ha terminado la actualizacion")
    print(f"Se han analizado {count} hojas")

# -------------- Funciones Programa -------------- #
# Seleccionar carpeta
def select_folder():
    folder_path = filedialog.askdirectory()
    if folder_path:
        print(f"Selected folder: {folder_path}")
        folder_selected.set(folder_path)

# Seleccionar archivo
def select_file():
    file_path = filedialog.askopenfilename()
    if file_path:
        print(f"Selected file: {file_path}")
        file_selected.set(file_path)

# Funcion para actualizar
def update():
    folder_path = select_folder()
    file_path = select_file()
    if folder_path and file_path:
        format_excel(folder_path, file_path)


# -------------- Programa -------------- #
root = tk.Tk()
root.title("Actualización de Informes")

# Variables para guardar los paths
folder_selected = tk.StringVar()
file_selected = tk.StringVar()

# Mostrar status de la funcion
message_label = tk.Label(root, text="Seleccionar carpeta y archivo, luego presionar 'ejecutar datos'.")
message_label.grid(row=3, columnspan=2, padx=10, pady=10)

# Boton de carpeta
folder_button = tk.Button(root, text="Seleccionar Carpeta", command=select_folder)
folder_button.grid(row=0, column=0, padx=10, pady=10)

# Entry to display selected folder path
folder_label = tk.Label(root, text="Seleccione la carpeta donde se encuentran los archivos que se van a actualizar")
folder_label.grid(row=0, column=1, padx=10, pady=10)

# Boton para archivo
file_button = tk.Button(root, text="Seleccionar archivo", command=select_file)
file_button.grid(row=1, column=0, padx=10, pady=10)

# Entry to display selected file path
file_label = tk.Label(root, text="Seleccione el archivo de Excel donde se encuentran los resumenes de informes")
file_label.grid(row=1, column=1, padx=10, pady=10) 

# Boton para la funcion
run_button = tk.Button(root, text="Actualizar datos", command=lambda: format_excel(folder_selected.get(), file_selected.get()))
run_button.grid(row=2, columnspan=2, padx=10, pady=10)

# Ejecutar Tkinter
root.mainloop()